import json
from typing import Dict, Any
from io import BytesIO
from datetime import datetime


def generate_pdf(title: str, content: Dict[str, Any]) -> bytes:
    """Generate PDF report (basic implementation)."""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=24,
            textColor="#1a1a1a",
            spaceAfter=30,
            alignment=TA_CENTER,
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.3 * inch))

        # Content sections
        if "sections" in content:
            for section in content["sections"]:
                section_title = ParagraphStyle(
                    "SectionTitle",
                    parent=styles["Heading2"],
                    fontSize=14,
                    textColor="#333333",
                    spaceAfter=12,
                )
                story.append(Paragraph(section.get("name", ""), section_title))

                section_content = ParagraphStyle(
                    "SectionContent",
                    parent=styles["BodyText"],
                    fontSize=10,
                    spaceAfter=12,
                )
                text = section.get("content", "").replace("\n", "<br/>")
                story.append(Paragraph(text, section_content))
                story.append(Spacer(1, 0.2 * inch))

        # Footer
        footer_style = ParagraphStyle(
            "Footer",
            parent=styles["Normal"],
            fontSize=8,
            textColor="#666666",
            spaceAfter=6,
        )
        story.append(Spacer(1, 0.3 * inch))
        story.append(
            Paragraph(
                f"Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}",
                footer_style,
            )
        )

        doc.build(story)
        return buffer.getvalue()

    except ImportError:
        return _generate_pdf_fallback(title, content)


def _generate_pdf_fallback(title: str, content: Dict[str, Any]) -> bytes:
    """Fallback PDF generation using simple text-based approach."""
    pdf_content = f"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< >>
stream
BT
/F1 24 Tf
50 750 Td
({title}) Tj
ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000214 00000 n
0000000317 00000 n
trailer
<< /Size 6 /Root 1 0 R >>
startxref
396
%%EOF"""
    return pdf_content.encode("utf-8")


def generate_docx(title: str, content: Dict[str, Any]) -> bytes:
    """Generate DOCX report (basic implementation)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sections
        if "sections" in content:
            for section in content["sections"]:
                doc.add_heading(section.get("name", ""), level=1)
                section_text = section.get("content", "")
                doc.add_paragraph(section_text)

        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph(
            f"Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        footer_para.runs[0].font.size = Pt(8)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except ImportError:
        return _generate_docx_fallback(title, content)


def _generate_docx_fallback(title: str, content: Dict[str, Any]) -> bytes:
    """Fallback DOCX generation using basic XML structure."""
    xml_content = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:r>
        <w:t>{title}</w:t>
      </w:r>
    </w:p>
  </w:body>
</w:document>"""
    return xml_content.encode("utf-8")


def generate_html(title: str, content: Dict[str, Any]) -> str:
    """Generate HTML report."""
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .header {{
            background-color: #1a1a1a;
            color: white;
            padding: 40px 20px;
            text-align: center;
            margin-bottom: 30px;
            border-radius: 5px;
        }}
        h1 {{
            margin: 0;
            font-size: 2.5em;
        }}
        h2 {{
            color: #1a1a1a;
            border-bottom: 2px solid #007bff;
            padding-bottom: 10px;
            margin-top: 30px;
            margin-bottom: 15px;
        }}
        .section {{
            background-color: white;
            padding: 20px;
            margin-bottom: 20px;
            border-radius: 5px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        .footer {{
            text-align: center;
            color: #666;
            font-size: 0.9em;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 10px;
            border-radius: 3px;
            overflow-x: auto;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{title}</h1>
    </div>
"""

    if "sections" in content:
        for section in content["sections"]:
            html += f"""    <div class="section">
        <h2>{section.get("name", "")}</h2>
        <p>{section.get("content", "").replace(chr(10), "<br>")}</p>
    </div>
"""

    html += f"""    <div class="footer">
        <p>Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
</body>
</html>"""

    return html


def generate_json_report(title: str, content: Dict[str, Any]) -> bytes:
    """Generate JSON report."""
    report = {
        "title": title,
        "generated_at": datetime.utcnow().isoformat(),
        "content": content,
    }
    return json.dumps(report, indent=2, ensure_ascii=False).encode("utf-8")
